import copy
from io import BytesIO
import json
from pathlib import Path
from tempfile import TemporaryDirectory
import unittest
from unittest.mock import patch
from zipfile import ZipFile

import generate_report as gr
import report_tool.llm as llm
import report_tool.ui_helpers as ui_helpers
from docx import Document


class StructuredHelperTests(unittest.TestCase):
    def _validated_finding(self, **overrides):
        finding = {
            "id": "VAPT-001",
            "name": "Deprecated TLS Protocol",
            "severity": "Medium",
            "cvss": "5.3",
            "cve": "[INSUFFICIENT DATA]",
            "affected_assets": "10.10.0.12:443, 10.10.0.13:443",
            "category": "Network - 10.10.0.0/24",
            "description": "Legacy TLS protocols remain enabled on the exposed service.",
            "business_impact": "Attackers may downgrade or weaken the transport security posture.",
            "proof_of_concept": "nmap --script ssl-enum-ciphers 10.10.0.12 -p 443",
            "remediation": "Disable deprecated TLS protocols and weak cipher support.",
            "control_objective": "Ensure only approved transport encryption standards are exposed.",
            "control_name": "TLS Configuration Hardening",
            "audit_requirement": "Verify only supported TLS versions and cipher suites remain enabled.",
            "reference": "OWASP Web Security Top 10, SANS25",
            "observation": "New",
            "remediation_status": "Open",
            "risk_status": "Open",
        }
        finding.update(overrides)
        return finding

    def _validated_report(self, findings=None, **overrides):
        resolved_findings = (
            [self._validated_finding()] if findings is None else findings
        )
        report = {
            "client_name": "Example Client",
            "report_date": "15-Apr-2026",
            "engagement_type": "Internal VAPT",
            "assessment_phase": "Phase 1",
            "assessor_firm": "Example Assessor",
            "scope_summary": "As a part of the internal VAPT engagement, the assessment covered the validated target environment.",
            "executive_summary": "The validated findings indicate a moderate risk posture requiring prioritised remediation.",
            "total_critical": 0,
            "total_high": 0,
            "total_medium": 1,
            "total_low": 0,
            "total_findings": len(resolved_findings),
            "conclusion": "[PRESERVE_ORIGINAL]",
            "methodology": "OWASP / PTES / NIST",
            "findings": resolved_findings,
        }
        report.update(overrides)
        return report

    def _front_matter_resolver(self):
        resolver = getattr(gr, "_resolve_front_matter_text", None)
        self.assertIsNotNone(
            resolver,
            "_resolve_front_matter_text helper is missing",
        )
        return resolver

    def _write_front_matter_template(
        self,
        template_path: Path,
        include_objectives=True,
        include_cover_marker=True,
    ):
        doc = Document()
        doc.add_paragraph("Internal VAPT Report", style="Title")
        if include_cover_marker:
            doc.add_paragraph("For", style="Title")
        doc.add_paragraph("", style="Title")
        doc.add_paragraph("Executive Summary", style="Heading 1")
        doc.add_paragraph("[PLACEHOLDER]")
        doc.add_paragraph("Introduction", style="Heading 1")
        doc.add_paragraph("[PLACEHOLDER]")
        if include_objectives:
            doc.add_paragraph("Objectives", style="Heading 1")
            doc.add_paragraph("[PLACEHOLDER]")
        doc.add_paragraph("Detailed Observation", style="Heading 1")
        doc.save(template_path)

    def _write_renderer_template(
        self,
        template_path: Path,
        *,
        include_summary_table: bool,
        include_detail_placeholder_table: bool = True,
    ):
        doc = Document()
        doc.add_paragraph("Internal VAPT Report", style="Title")
        doc.add_paragraph("For", style="Title")
        doc.add_paragraph("", style="Title")

        meta_table = doc.add_table(rows=1, cols=2)
        meta_table.rows[0].cells[0].text = "Report Release Date"
        meta_table.rows[0].cells[1].text = "[PLACEHOLDER]"

        company_table = doc.add_table(rows=8, cols=2)
        for row in company_table.rows:
            for cell in row.cells:
                cell.text = "[PLACEHOLDER]"

        doc.add_paragraph("Executive Summary", style="Heading 1")
        doc.add_paragraph("[PLACEHOLDER]")
        doc.add_paragraph("Introduction", style="Heading 1")
        doc.add_paragraph("[PLACEHOLDER]")
        doc.add_paragraph("Objectives", style="Heading 1")
        doc.add_paragraph("[PLACEHOLDER]")

        if include_summary_table:
            summary_table = doc.add_table(rows=2, cols=7)
            headers = [
                "Sr. No",
                "Category",
                "Severity",
                "Vulnerability",
                "Status",
                "Observation",
                "Affected Assets",
            ]
            for cell, text in zip(summary_table.rows[0].cells, headers):
                cell.text = text
            for cell in summary_table.rows[1].cells:
                cell.text = "[PLACEHOLDER]"

        doc.add_paragraph("Detailed Observation", style="Heading 1")
        if include_detail_placeholder_table:
            placeholder_table = doc.add_table(rows=1, cols=1)
            placeholder_table.rows[0].cells[0].text = "[PLACEHOLDER]"
        doc.save(template_path)

    def test_build_payload_requests_json_object_output(self):
        # Cloud providers (openrouter) get response_format so the model stays in JSON mode.
        payload = llm._build_payload(
            "",
            {
                "llm": {
                    "model": "test-model",
                    "provider": "openrouter",
                    "temperature": 0.1,
                    "max_tokens": 512,
                }
            },
            system_prompt="Return JSON.",
            user_content="Return one JSON object.",
        )
        self.assertEqual(payload.get("response_format"), {"type": "json_object"})

    def test_build_payload_local_provider_omits_response_format(self):
        # Local providers (LM Studio, Ollama) must NOT receive response_format
        # unless explicitly opted in via json_mode — many models reject the field.
        payload = llm._build_payload(
            "",
            {
                "llm": {
                    "model": "local-model",
                    "provider": "local",
                    "temperature": 0.1,
                    "max_tokens": 512,
                }
            },
            system_prompt="Return JSON.",
            user_content="Return one JSON object.",
        )
        self.assertNotIn("response_format", payload)

    def test_build_payload_local_json_mode_flag_adds_response_format(self):
        # Opt-in via json_mode flag should still work for local providers.
        payload = llm._build_payload(
            "",
            {
                "llm": {
                    "model": "local-model",
                    "provider": "local",
                    "json_mode": True,
                    "temperature": 0.1,
                    "max_tokens": 512,
                }
            },
            system_prompt="Return JSON.",
            user_content="Return one JSON object.",
        )
        self.assertEqual(payload.get("response_format"), {"type": "json_object"})

    def test_build_payload_requires_an_explicit_model(self):
        with self.assertRaises(ValueError):
            llm._build_payload(
                "",
                {
                    "llm": {
                        "provider": "local",
                        "temperature": 0.1,
                        "max_tokens": 512,
                    }
                },
                system_prompt="Return JSON.",
                user_content="Return one JSON object.",
            )

    def _paragraph_after_heading(self, doc: Document, heading_text: str) -> str:
        for index, paragraph in enumerate(doc.paragraphs):
            if paragraph.text.strip() != heading_text:
                continue
            if index + 1 < len(doc.paragraphs):
                return doc.paragraphs[index + 1].text
        self.fail(f"Heading not found in rendered document: {heading_text}")

    def _finding_table_rows(self, table_xml):
        rows = []
        for row in table_xml.findall(gr.qn("w:tr")):
            cells = []
            for cell in row.findall(gr.qn("w:tc")):
                cells.append(
                    "".join(
                        text_node.text or "" for text_node in cell.iter(gr.qn("w:t"))
                    )
                )
            rows.append(cells)
        return rows

    def _paragraph_texts_after_heading(
        self, doc: Document, heading_text: str
    ) -> list[str]:
        texts = [
            paragraph.text.strip()
            for paragraph in doc.paragraphs
            if paragraph.text.strip()
        ]
        try:
            start_index = texts.index(heading_text)
        except ValueError as exc:
            self.fail(f"Heading not found in rendered document: {heading_text}")
        return texts[start_index + 1 :]

    def _summary_table_rows(self, doc: Document) -> list[list[str]]:
        for table in doc.tables:
            rows = [[cell.text for cell in row.cells] for row in table.rows]
            if not rows:
                continue
            header_cells = rows[0][:4]
            if "Severity" in header_cells and any(
                "Vulnerability" in cell for cell in header_cells
            ):
                return rows
        self.fail("Summary table not found in rendered document")

    def _summary_table(self, doc: Document):
        for table in doc.tables:
            rows = [[cell.text for cell in row.cells] for row in table.rows]
            if not rows:
                continue
            header_cells = rows[0][:4]
            if "Severity" in header_cells and any(
                "Vulnerability" in cell for cell in header_cells
            ):
                return table
        self.fail("Summary table not found in rendered document")

    def _rendered_finding_tables(self, doc: Document) -> list[list[list[str]]]:
        finding_tables = []
        for table in doc.tables:
            rows = self._finding_table_rows(table._tbl)
            if (
                len(rows) == 12
                and len(rows[0]) == 3
                and rows[0][1] == "Affected URL /IP"
            ):
                finding_tables.append(rows)
        return finding_tables

    def _body_chart_and_summary_indexes(self, docx_path: Path) -> tuple[int, int]:
        from lxml import etree as ET

        with ZipFile(docx_path) as archive:
            root = ET.fromstring(archive.read("word/document.xml"))

        ns = {
            "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
            "c": "http://schemas.openxmlformats.org/drawingml/2006/chart",
        }
        body = root.find("w:body", ns)
        self.assertIsNotNone(body)

        chart_idx = -1
        summary_idx = -1
        for idx, child in enumerate(body):
            tag = ET.QName(child).localname
            if tag == "p" and child.xpath(".//c:chart", namespaces=ns):
                chart_idx = idx
            if tag == "tbl":
                headers = [
                    text.strip()
                    for text in child.xpath(
                        "./w:tr[1]/w:tc//w:t/text()",
                        namespaces=ns,
                    )
                ]
                if "Severity" in headers and any(
                    "Vulnerability" in header for header in headers
                ):
                    summary_idx = idx
        self.assertNotEqual(chart_idx, -1, "Chart paragraph not found in rendered DOCX")
        self.assertNotEqual(summary_idx, -1, "Summary table not found in rendered DOCX body")
        return chart_idx, summary_idx

    def _body_text_index(self, docx_path: Path, expected_text: str) -> int:
        from lxml import etree as ET

        with ZipFile(docx_path) as archive:
            root = ET.fromstring(archive.read("word/document.xml"))

        ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
        body = root.find("w:body", ns)
        self.assertIsNotNone(body)

        expected = expected_text.strip()
        for idx, child in enumerate(body):
            texts = [
                text.strip()
                for text in child.xpath(".//w:t/text()", namespaces=ns)
                if text.strip()
            ]
            if " ".join(texts) == expected:
                return idx

        self.fail(f"Body text not found in rendered DOCX: {expected_text}")

    def test_build_finding_preserves_all_assets_raw(self):
        assets = [f"10.0.0.{idx}:443" for idx in range(1, 13)]

        result = gr._enrich_finding_presentation(
            {
                "name": "TLS Service Exposure",
                "affected_assets": ", ".join(assets),
            }
        )

        self.assertEqual(result["affected_assets_raw"], assets)

    def test_build_finding_adds_short_assets_and_full_trace_block(self):
        assets = [f"app{idx}.example.com:8443" for idx in range(1, 13)]
        affected_assets = ", ".join(assets)

        result = gr._enrich_finding_presentation(
            {
                "name": "Exposed Administrative Interface",
                "affected_assets": affected_assets,
            }
        )

        self.assertEqual(
            result["affected_assets_short"], gr._truncate_assets(affected_assets)
        )
        self.assertNotEqual(result["affected_assets_short"], affected_assets)
        for asset in assets:
            self.assertIn(asset, result["asset_trace_block"])

    def test_build_finding_keeps_list_valued_assets_clean(self):
        assets = ["10.10.10.10:443", "10.10.10.11:443"]

        result = gr._enrich_finding_presentation(
            {
                "name": "Load Balancer Admin Interface",
                "affected_assets": assets,
            }
        )

        self.assertEqual(result["affected_assets_raw"], assets)
        self.assertEqual(result["affected_assets"], ", ".join(assets))
        self.assertEqual(
            result["asset_trace_block"],
            "Affected assets:\n- 10.10.10.10:443\n- 10.10.10.11:443",
        )

    def test_normalize_affected_assets_keeps_tuple_inputs_clean(self):
        assets = (
            " 10.10.10.10:443 ",
            "10.10.10.11:443, 10.10.10.12:443",
            "10.10.10.10:443",
            "[PLACEHOLDER]",
            None,
        )

        self.assertEqual(
            gr._normalize_affected_assets(assets),
            [
                "10.10.10.10:443",
                "10.10.10.11:443",
                "10.10.10.12:443",
            ],
        )

    def test_normalize_affected_assets_makes_set_inputs_deterministic(self):
        assets = {
            "host-e.example.com:443",
            "host-c.example.com:443",
            "host-a.example.com:443",
            "host-d.example.com:443",
            "host-b.example.com:443",
        }

        self.assertEqual(
            gr._normalize_affected_assets(assets),
            [
                "host-a.example.com:443",
                "host-b.example.com:443",
                "host-c.example.com:443",
                "host-d.example.com:443",
                "host-e.example.com:443",
            ],
        )

    def test_build_finding_display_title_prefers_explicit_name(self):
        result = gr._enrich_finding_presentation(
            {
                "name": "Explicit Finding Name",
                "short_name": "Fallback Short Name",
                "title": "Fallback Title",
                "affected_assets": "10.10.10.10:443",
            }
        )

        self.assertEqual(result["display_title"], "Explicit Finding Name")

    def test_build_finding_display_title_reuses_structured_title_heuristics(self):
        result = gr._enrich_finding_presentation(
            {
                "name": "hsts missing from https server",
                "affected_assets": "10.10.10.10:443",
            }
        )

        self.assertEqual(
            result["display_title"],
            "HSTS Missing from HTTPS Server (RFC 6797)",
        )

    def test_dedupe_findings_output_includes_resolved_asset_fields(self):
        assets = [f"172.16.0.{idx}:22" for idx in range(1, 13)]

        deduped = gr.dedupe_findings(
            [
                {
                    "id": "VAPT-001",
                    "vuln_id": "1",
                    "name": "SSH Service Exposure",
                    "description": "SSH is exposed on multiple assets.",
                    "affected_assets": ", ".join(assets),
                }
            ]
        )

        self.assertEqual(deduped[0]["affected_assets_raw"], assets)
        self.assertEqual(
            deduped[0]["affected_assets_short"],
            gr._truncate_assets(deduped[0]["affected_assets"]),
        )
        self.assertEqual(deduped[0]["display_title"], "SSH Service Exposure")

    def test_dedupe_findings_merges_unique_assets_for_duplicate_vuln_id(self):
        expected_assets = [
            "192.168.1.10:443",
            "192.168.1.11:443",
            "192.168.1.12:443",
        ]

        deduped = gr.dedupe_findings(
            [
                {
                    "id": "VAPT-001",
                    "vuln_id": "7",
                    "name": "TLS Weak Cipher Suites",
                    "description": "Weak ciphers remain exposed.",
                    "affected_assets": [
                        "192.168.1.10:443",
                        "192.168.1.11:443",
                    ],
                },
                {
                    "id": "VAPT-002",
                    "vuln_id": "7",
                    "name": "TLS Weak Cipher Suites",
                    "description": "Weak ciphers remain exposed on an additional host.",
                    "affected_assets": "192.168.1.11:443, 192.168.1.12:443",
                },
            ]
        )

        self.assertEqual(len(deduped), 1)
        self.assertEqual(deduped[0]["affected_assets_raw"], expected_assets)
        self.assertEqual(deduped[0]["affected_assets"], ", ".join(expected_assets))
        self.assertEqual(
            deduped[0]["affected_assets_short"],
            gr._truncate_assets(", ".join(expected_assets)),
        )
        self.assertEqual(
            deduped[0]["asset_trace_block"],
            "Affected assets:\n- 192.168.1.10:443\n- 192.168.1.11:443\n- 192.168.1.12:443",
        )

    def test_dedupe_findings_preserves_distinct_status_families_for_same_vuln_id(self):
        deduped = gr.dedupe_findings(
            [
                {
                    "id": "VAPT-401",
                    "vuln_id": "401",
                    "name": "SSL Version 2 and 3 Protocol Detection",
                    "description": "Legacy protocol support requires validation.",
                    "affected_assets": "10.40.0.10:443",
                    "severity": "Medium",
                    "observation": "New",
                    "remediation_status": "Open",
                    "risk_status": "Open",
                },
                {
                    "id": "VAPT-402",
                    "vuln_id": "401",
                    "name": "SSL Version 2 and 3 Protocol Detection",
                    "description": "Legacy protocol support requires validation.",
                    "affected_assets": "10.40.0.10:443",
                    "severity": "Medium",
                    "observation": "New",
                    "remediation_status": "Open",
                    "risk_status": "False Positive Check",
                },
            ]
        )

        self.assertEqual(len(deduped), 2)
        self.assertEqual(
            [finding["id"] for finding in deduped],
            ["VAPT-401", "VAPT-402"],
        )
        self.assertEqual(
            [
                gr._structured_status_family(
                    finding.get("observation", ""),
                    finding.get("remediation_status", ""),
                    finding.get("risk_status", ""),
                )
                for finding in deduped
            ],
            ["open", "fp_review"],
        )

    def test_compose_report_narrative_adds_report_level_fields_with_safe_defaults(self):
        composed = gr.compose_report_narrative(
            self._validated_report(
                findings=[
                    self._validated_finding(
                        category="[INSUFFICIENT DATA]",
                        affected_assets="host-a.example.com:443",
                    )
                ]
            )
        )

        self.assertIn("narrative_slots", composed)
        self.assertIsInstance(composed["narrative_slots"], dict)
        self.assertIn("objectives", composed)
        self.assertIsInstance(composed["objectives"], list)
        self.assertTrue(composed["objectives"])
        self.assertIn("outline_groups", composed)
        self.assertIsInstance(composed["outline_groups"], list)
        self.assertTrue(composed["outline_groups"])
        self.assertIn("introduction_overview", composed)
        self.assertIsInstance(composed["introduction_overview"], str)
        self.assertIn("introduction_scope_bridge", composed)
        self.assertIsInstance(composed["introduction_scope_bridge"], str)

        finding = composed["findings"][0]
        self.assertEqual(finding["taxonomy_family"], "general")
        self.assertEqual(finding["subnet_key"], "host-specific")
        self.assertEqual(finding["subnet_label"], "Host-specific findings")

    def test_compose_report_narrative_prefers_findings_summary_over_raw_executive_summary(
        self,
    ):
        raw_executive_summary = "Raw executive summary fallback should not win."

        composed = gr.compose_report_narrative(
            self._validated_report(executive_summary=raw_executive_summary)
        )

        self.assertEqual(
            composed["narrative_slots"]["summary_context"],
            "The validated assessment identified 1 confirmed finding across 1 grouped remediation workstream.",
        )
        self.assertNotIn(raw_executive_summary, composed["introduction_overview"])
        self.assertIn(
            composed["narrative_slots"]["summary_key_drivers"],
            composed["introduction_overview"],
        )

    def test_compose_report_narrative_prefers_derived_scope_bridge_over_raw_scope_summary(
        self,
    ):
        raw_scope_summary = "Raw scope summary fallback should not win."

        composed = gr.compose_report_narrative(
            self._validated_report(scope_summary=raw_scope_summary)
        )

        expected_intro_scope = (
            "The validated scope covered 10.10.0.0/24 within the assessed environment"
        )
        self.assertEqual(
            composed["narrative_slots"]["intro_scope"], expected_intro_scope
        )
        self.assertTrue(
            composed["introduction_scope_bridge"].startswith(expected_intro_scope)
        )
        self.assertNotIn(raw_scope_summary, composed["introduction_scope_bridge"])

    def test_compose_report_narrative_falls_back_to_raw_front_matter_when_findings_are_empty(
        self,
    ):
        raw_executive_summary = (
            "No confirmed findings were carried into the report draft."
        )
        raw_scope_summary = "The assessment scope covered the validated target environment from the extracted metadata."

        composed = gr.compose_report_narrative(
            self._validated_report(
                findings=[],
                total_findings=0,
                total_medium=0,
                executive_summary=raw_executive_summary,
                scope_summary=raw_scope_summary,
            )
        )

        self.assertEqual(
            composed["narrative_slots"]["summary_context"], raw_executive_summary
        )
        self.assertEqual(composed["narrative_slots"]["intro_scope"], raw_scope_summary)
        self.assertIn(raw_executive_summary, composed["introduction_overview"])
        self.assertIn(raw_scope_summary, composed["introduction_scope_bridge"])

    def test_compose_report_narrative_preserves_existing_asset_lookup(self):
        result = gr._enrich_finding_presentation(
            self._validated_finding(
                affected_assets="10.20.30.40:443, 10.20.30.41:443, 10.20.30.42:443"
            )
        )

        composed = gr.compose_report_narrative(
            self._validated_report(findings=[result])
        )

        self.assertEqual(
            composed["findings"][0]["affected_assets_raw"],
            result["affected_assets_raw"],
        )
        self.assertEqual(
            composed["findings"][0]["asset_trace_block"], result["asset_trace_block"]
        )

    def test_compose_report_narrative_preserves_original_affected_assets_value(self):
        original_assets = " host-a.example.com:443 , host-b.example.com:443 "

        composed = gr.compose_report_narrative(
            self._validated_report(
                findings=[
                    self._validated_finding(
                        affected_assets=original_assets,
                    )
                ]
            )
        )

        finding = composed["findings"][0]
        self.assertEqual(finding["affected_assets"], original_assets)
        self.assertEqual(
            finding["affected_assets_raw"],
            ["host-a.example.com:443", "host-b.example.com:443"],
        )

    def test_build_composed_preview_recomputes_fields_without_mutating_input(self):
        assets = [f"host-{idx}.example.com:443" for idx in range(1, 13)]
        source = self._validated_report(
            findings=[
                self._validated_finding(
                    name="hsts missing from https server",
                    affected_assets=", ".join(assets),
                    control_name="[PLACEHOLDER]",
                    control_objective="[PLACEHOLDER]",
                    display_title="Stale title",
                    affected_assets_short="stale short assets",
                    asset_trace_block="stale trace block",
                    display_control_name="stale control name",
                    display_control_objective="stale control objective",
                )
            ],
            objectives=["stale objective"],
            introduction_overview="stale overview",
            introduction_scope_bridge="stale bridge",
        )
        original = copy.deepcopy(source)

        preview = ui_helpers.build_composed_preview(source)

        self.assertEqual(source, original)
        self.assertTrue(preview["objectives"])
        self.assertNotEqual(preview["objectives"], ["stale objective"])
        self.assertNotEqual(preview["introduction_overview"], "stale overview")
        self.assertNotEqual(
            preview["introduction_scope_bridge"],
            "stale bridge",
        )

        finding = preview["findings"][0]
        self.assertEqual(
            finding["display_title"],
            "HSTS Missing from HTTPS Server (RFC 6797)",
        )
        self.assertEqual(
            finding["affected_assets_short"],
            gr._truncate_assets(", ".join(assets)),
        )
        self.assertIn("Affected assets:", finding["asset_trace_block"])
        self.assertIn("host-1.example.com:443", finding["asset_trace_block"])
        self.assertEqual(
            finding["display_control_name"],
            "Vulnerability Remediation",
        )
        self.assertEqual(
            finding["display_control_objective"],
            "Identify and remediate the vulnerability to reduce the attack surface.",
        )

    def test_apply_report_level_overrides_uses_only_meaningful_values(self):
        composed = self._validated_report(
            findings=[],
            objectives=["Auto objective"],
            introduction_overview="Auto overview.",
            introduction_scope_bridge="Auto scope bridge.",
            narrative_slots={"objectives": ["Auto objective"]},
        )
        override_source = {
            "front_matter_overrides": {
                "introduction_overview": "User-refined overview.",
                "introduction_scope_bridge": "   ",
                "objectives": "- Validate externally exposed services\n- Prioritise remediation sequencing",
            }
        }

        original_composed = copy.deepcopy(composed)
        original_source = copy.deepcopy(override_source)

        overridden = ui_helpers.apply_report_level_overrides(
            composed,
            override_source,
        )

        self.assertEqual(composed, original_composed)
        self.assertEqual(override_source, original_source)
        self.assertEqual(overridden["introduction_overview"], "User-refined overview.")
        self.assertEqual(
            overridden["introduction_scope_bridge"],
            "Auto scope bridge.",
        )
        self.assertEqual(
            overridden["objectives"],
            [
                "Validate externally exposed services",
                "Prioritise remediation sequencing",
            ],
        )
        self.assertEqual(
            overridden["narrative_slots"]["objectives"],
            [
                "Validate externally exposed services",
                "Prioritise remediation sequencing",
            ],
        )

    def test_build_composed_preview_recomputes_then_applies_report_level_overrides(
        self,
    ):
        source = self._validated_report(
            findings=[
                self._validated_finding(
                    name="hsts missing from https server",
                    affected_assets="host-1.example.com:443",
                )
            ],
            objectives=["stale objective"],
            introduction_overview="stale overview",
            introduction_scope_bridge="stale bridge",
            front_matter_overrides={
                "introduction_overview": "User-refined overview.",
                "introduction_scope_bridge": "User-refined introduction.",
                "objectives": "Validate externally exposed services\nPrioritise remediation sequencing",
            },
        )

        preview = ui_helpers.build_composed_preview(source)

        self.assertEqual(preview["introduction_overview"], "User-refined overview.")
        self.assertEqual(
            preview["introduction_scope_bridge"],
            "User-refined introduction.",
        )
        self.assertEqual(
            preview["objectives"],
            [
                "Validate externally exposed services",
                "Prioritise remediation sequencing",
            ],
        )
        self.assertIn(
            "Primary remediation drivers include",
            preview["narrative_slots"]["summary_key_drivers"],
        )

    def test_ui_helpers_loads_run_log_files(self):
        with TemporaryDirectory() as tmp_dir:
            log_path = Path(tmp_dir) / "20260422_120000_000001_run_log.json"
            log_payload = {
                "log_version": 1,
                "created_at": "2026-04-22T12:00:00",
                "events": [
                    {
                        "timestamp": "2026-04-22T12:00:01",
                        "type": "llm_interaction",
                        "data": {"response_restored": '{"ok": true}'},
                    }
                ],
            }
            log_path.write_text(json.dumps(log_payload), encoding="utf-8")

            listed = ui_helpers.list_log_files(tmp_dir)
            loaded = ui_helpers.load_log_data(str(log_path))

        self.assertEqual(len(listed), 1)
        self.assertEqual(listed[0]["name"], log_path.name)
        self.assertEqual(loaded["parsed"]["log_version"], 1)
        self.assertEqual(loaded["events"][0]["type"], "llm_interaction")

    def test_compose_risk_posture_mentions_dominant_taxonomy_themes(self):
        posture = gr._compose_risk_posture(
            [
                self._validated_finding(
                    severity="High",
                    taxonomy_family="application",
                    taxonomy_label="Application Security Findings",
                    category="Application - 10.20.0.0/24",
                ),
                self._validated_finding(
                    id="VAPT-002",
                    severity="High",
                    taxonomy_family="application",
                    taxonomy_label="Application Security Findings",
                    category="Application - 10.20.1.0/24",
                ),
                self._validated_finding(
                    id="VAPT-003",
                    severity="Medium",
                    taxonomy_family="network",
                    taxonomy_label="Network Security Findings",
                    category="Network - 10.10.0.0/24",
                ),
            ]
        )

        self.assertEqual(
            posture,
            "The current risk posture remains elevated because 2 high-severity findings require prioritised remediation, concentrated in Application Security Findings and Network Security Findings.",
        )

    def test_compose_report_objectives_vary_with_severity_and_theme_data(self):
        objectives = gr._compose_report_objectives(
            {},
            [
                self._validated_finding(
                    severity="Critical",
                    taxonomy_family="network",
                    taxonomy_label="Network Security Findings",
                    subnet_key="10.10.0.0/24",
                    subnet_label="Subnet 10.10.0.0/24",
                    category="Network - 10.10.0.0/24",
                    affected_assets="10.10.0.10:443",
                ),
                self._validated_finding(
                    id="VAPT-002",
                    severity="High",
                    taxonomy_family="application",
                    taxonomy_label="Application Security Findings",
                    subnet_key="10.20.0.0/24",
                    subnet_label="Subnet 10.20.0.0/24",
                    category="Application - 10.20.0.0/24",
                    affected_assets="10.20.0.50:8443",
                ),
            ],
        )

        self.assertEqual(
            objectives,
            [
                "Contain and remediate the 1 critical finding driving risk in Network Security Findings and Application Security Findings.",
                "Coordinate remediation workstreams around Network Security Findings and Application Security Findings while tracking progress across 2 taxonomy families and 2 subnet groups.",
                "Preserve full asset traceability for each confirmed finding while maintaining compact display fields for report-facing views.",
            ],
        )

    def test_compose_report_objectives_theme_order_is_deterministic(self):
        findings = [
            self._validated_finding(
                severity="High",
                taxonomy_family="network",
                taxonomy_label="Network Security Findings",
                subnet_key="10.10.0.0/24",
                subnet_label="Subnet 10.10.0.0/24",
                category="Network - 10.10.0.0/24",
                affected_assets="10.10.10.10:443",
            ),
            self._validated_finding(
                id="VAPT-002",
                severity="High",
                taxonomy_family="application",
                taxonomy_label="Application Security Findings",
                subnet_key="10.20.0.0/24",
                subnet_label="Subnet 10.20.0.0/24",
                category="Application - 10.20.0.0/24",
                affected_assets="10.20.0.50:8443",
            ),
        ]

        objectives = gr._compose_report_objectives({}, findings)
        reversed_objectives = gr._compose_report_objectives(
            {}, list(reversed(findings))
        )

        self.assertEqual(objectives, reversed_objectives)
        self.assertIn(
            "Application Security Findings and Network Security Findings",
            objectives[0],
        )

    def test_build_finding_removes_placeholder_like_asset_tokens_from_derived_fields(
        self,
    ):
        result = gr._enrich_finding_presentation(
            {
                "name": "Administrative Interface Exposure",
                "affected_assets": [
                    "[placeholder]",
                    "nan",
                    "None",
                    "Affected assets: [INSUFFICIENT DATA]",
                    None,
                    "host-a.example.com:443",
                    "host-b.example.com:443",
                ],
            }
        )

        self.assertEqual(
            result["affected_assets_raw"],
            ["host-a.example.com:443", "host-b.example.com:443"],
        )
        self.assertEqual(
            result["affected_assets_short"],
            "host-a.example.com:443, host-b.example.com:443",
        )
        self.assertEqual(
            result["asset_trace_block"],
            "Affected assets:\n- host-a.example.com:443\n- host-b.example.com:443",
        )

    def test_dedupe_findings_collapses_missing_vuln_id_duplicates_by_fingerprint(self):
        deduped = gr.dedupe_findings(
            [
                {
                    "id": "VAPT-001",
                    "name": "TLS Weak Cipher Suites",
                    "description": "Weak ciphers remain exposed on the internet-facing service.",
                    "affected_assets": "192.168.1.10:443",
                },
                {
                    "id": "VAPT-002",
                    "name": "TLS Weak Cipher Suites",
                    "description": "Weak ciphers remain exposed on the internet-facing service.",
                    "affected_assets": ["192.168.1.10:443"],
                },
            ]
        )

        self.assertEqual(len(deduped), 1)
        self.assertEqual(deduped[0]["affected_assets_raw"], ["192.168.1.10:443"])

    def test_build_finding_table_prefers_resolved_visible_fields(self):
        assets = [f"admin{idx}.example.com:8443" for idx in range(1, 13)]
        finding = gr._enrich_finding_presentation(
            self._validated_finding(
                id="VAPT-777",
                name="Raw Scanner Title",
                affected_assets=", ".join(assets),
                control_objective="Raw control objective",
                control_name="Raw control name",
            )
        )
        finding["display_title"] = "Validated Administrative Interface Exposure"
        finding["display_control_objective"] = (
            "Restrict administrative interfaces to approved operators."
        )
        finding["display_control_name"] = "Administrative Interface Review"

        rows = self._finding_table_rows(gr._build_finding_table(finding, 7))

        self.assertEqual(len(rows), 12)
        self.assertEqual(
            rows[0],
            ["7", "Affected URL /IP", finding["asset_trace_block"]],
        )
        self.assertNotIn("...", rows[0][2])
        self.assertEqual(
            rows[1],
            [
                "Vulnerability title / Observation",
                "Validated Administrative Interface Exposure (VAPT-777)",
            ],
        )
        self.assertEqual(
            rows[6],
            [
                "Control Objective",
                "Raw control objective",
            ],
        )
        self.assertEqual(rows[7], ["Control Name", "Raw control name"])

    def test_build_finding_table_falls_back_to_legacy_visible_fields(self):
        finding = self._validated_finding(
            id="",
            name="Legacy Finding Title",
            affected_assets="host-a.example.com:443, host-b.example.com:443",
            control_objective="Legacy control objective",
            control_name="Legacy control name",
        )

        rows = self._finding_table_rows(gr._build_finding_table(finding, 3))

        self.assertEqual(len(rows), 12)
        self.assertEqual(
            rows[0],
            ["3", "Affected URL /IP", gr._truncate_assets(finding["affected_assets"])],
        )
        self.assertEqual(
            rows[1],
            ["Vulnerability title / Observation", "Legacy Finding Title"],
        )
        self.assertEqual(rows[6], ["Control Objective", "Legacy control objective"])
        self.assertEqual(rows[7], ["Control Name", "Legacy control name"])

    def test_render_report_finding_table_falls_back_from_placeholder_like_asset_trace_block(
        self,
    ):
        with TemporaryDirectory() as tmp_dir:
            template_path = Path(tmp_dir) / "template.docx"
            output_path = Path(tmp_dir) / "rendered.docx"
            self._write_renderer_template(
                template_path,
                include_summary_table=False,
            )

            finding = self._validated_finding(
                id="VAPT-741",
                name="Placeholder Trace Finding",
                affected_assets="10.80.0.10:443, 10.80.0.11:443",
                asset_trace_block="Affected assets:\n- [INSUFFICIENT DATA]",
            )

            gr.render_report(
                self._validated_report(findings=[finding]),
                str(template_path),
                str(output_path),
            )

            rendered = Document(output_path)

        finding_tables = self._rendered_finding_tables(rendered)
        self.assertEqual(len(finding_tables), 1)
        self.assertEqual(
            finding_tables[0][0][2],
            "10.80.0.10:443, 10.80.0.11:443",
        )
        self.assertNotIn("[INSUFFICIENT DATA]", finding_tables[0][0][2])

    def test_compose_report_narrative_fills_display_control_fields(self):
        composed = gr.compose_report_narrative(
            self._validated_report(
                findings=[
                    self._validated_finding(
                        control_name="Access Control Review",
                        control_objective="Ensure administrative interfaces are restricted to authorised operators.",
                    )
                ]
            )
        )

        finding = composed["findings"][0]
        self.assertEqual(finding["display_control_name"], "Access Control Review")
        self.assertEqual(
            finding["display_control_objective"],
            "Ensure administrative interfaces are restricted to authorised operators.",
        )

    def test_compose_report_narrative_is_idempotent(self):
        report = self._validated_report()

        once = gr.compose_report_narrative(report)
        twice = gr.compose_report_narrative(once)

        self.assertEqual(once, twice)

    def test_compose_report_narrative_refresh_recomputes_derived_fields_after_edits(
        self,
    ):
        original = gr.compose_report_narrative(
            self._validated_report(
                findings=[
                    self._validated_finding(
                        severity="Medium",
                        category="Network - 10.10.0.0/24",
                        affected_assets="10.10.0.12:443",
                    )
                ]
            )
        )

        edited = copy.deepcopy(original)
        edited_finding = edited["findings"][0]
        edited_finding["severity"] = "Critical"
        edited_finding["category"] = "Application - 10.20.0.0/24"
        edited_finding["affected_assets"] = "10.20.0.50:8443"

        stale = gr.compose_report_narrative(copy.deepcopy(edited))
        refreshed = gr.compose_report_narrative(copy.deepcopy(edited), refresh=True)

        self.assertEqual(
            stale["findings"][0]["affected_assets_raw"], ["10.10.0.12:443"]
        )
        self.assertEqual(stale["findings"][0]["taxonomy_family"], "network")
        self.assertEqual(stale["findings"][0]["subnet_key"], "10.10.0.0/24")
        self.assertEqual(
            stale["introduction_overview"], original["introduction_overview"]
        )

        self.assertEqual(
            refreshed["findings"][0]["affected_assets_raw"], ["10.20.0.50:8443"]
        )
        self.assertEqual(refreshed["findings"][0]["taxonomy_family"], "application")
        self.assertEqual(refreshed["findings"][0]["subnet_key"], "10.20.0.0/24")
        self.assertEqual(
            refreshed["narrative_slots"]["summary_key_drivers"],
            "Primary remediation drivers include Application",
        )
        self.assertIn(
            "1 critical finding require immediate remediation",
            refreshed["introduction_overview"],
        )

    def test_resolve_front_matter_text_prefers_composed_fields_and_joins_objectives(
        self,
    ):
        resolved = self._front_matter_resolver()(
            self._validated_report(
                findings=[],
                executive_summary="Legacy executive summary.",
                introduction_overview="Composed executive summary.",
                scope_summary="Legacy introduction.",
                introduction_scope_bridge="Composed introduction.",
                objectives=[
                    "Validate assets.",
                    "Prioritise remediation.",
                ],
            )
        )

        self.assertEqual(resolved["executive_summary"], "Composed executive summary.")
        self.assertEqual(resolved["introduction"], "Composed introduction.")
        self.assertEqual(
            resolved["objectives"],
            "Validate assets. Prioritise remediation.",
        )

    def test_resolve_front_matter_text_falls_back_and_suppresses_placeholders(self):
        resolved = self._front_matter_resolver()(
            self._validated_report(
                findings=[],
                executive_summary="Legacy executive summary.",
                introduction_overview="[PLACEHOLDER]",
                scope_summary="Legacy introduction.",
                introduction_scope_bridge="nan",
                objectives=["[PLACEHOLDER]", " ", None, "nan"],
            )
        )

        self.assertEqual(resolved["executive_summary"], "Legacy executive summary.")
        self.assertEqual(resolved["introduction"], "Legacy introduction.")
        self.assertEqual(resolved["objectives"], "")

    def test_resolve_front_matter_text_uses_narrative_objectives_when_top_level_objectives_empty(
        self,
    ):
        resolved = self._front_matter_resolver()(
            self._validated_report(
                findings=[],
                objectives=[],
                narrative_slots={
                    "objectives": [
                        "Validate assets.",
                        "Prioritise remediation.",
                    ]
                },
            )
        )

        self.assertEqual(
            resolved["objectives"],
            "Validate assets. Prioritise remediation.",
        )

    def test_render_report_prefers_composed_front_matter_and_renders_objectives(self):
        with TemporaryDirectory() as tmp_dir:
            template_path = Path(tmp_dir) / "template.docx"
            output_path = Path(tmp_dir) / "rendered.docx"
            self._write_front_matter_template(template_path)

            gr.render_report(
                self._validated_report(
                    findings=[],
                    executive_summary="Legacy executive summary.",
                    introduction_overview="Composed executive summary.",
                    scope_summary="Legacy introduction.",
                    introduction_scope_bridge="Composed introduction.",
                    objectives=[
                        "Validate assets.",
                        "Prioritise remediation.",
                    ],
                ),
                str(template_path),
                str(output_path),
            )

            rendered = Document(output_path)

        self.assertEqual(
            self._paragraph_after_heading(rendered, "Executive Summary"),
            "Composed executive summary.",
        )
        self.assertEqual(
            self._paragraph_after_heading(rendered, "Introduction"),
            "Composed introduction.",
        )
        self.assertEqual(
            self._paragraph_after_heading(rendered, "Objectives"),
            "Validate assets. Prioritise remediation.",
        )

    def test_render_report_uses_blank_title_slot_when_cover_marker_is_missing(self):
        with TemporaryDirectory() as tmp_dir:
            template_path = Path(tmp_dir) / "template.docx"
            output_path = Path(tmp_dir) / "rendered.docx"
            self._write_front_matter_template(
                template_path,
                include_objectives=False,
                include_cover_marker=False,
            )

            gr.render_report(
                self._validated_report(
                    findings=[],
                    executive_summary="Composed executive summary.",
                    introduction_scope_bridge="Composed introduction.",
                ),
                str(template_path),
                str(output_path),
            )

            rendered = Document(output_path)

        self.assertEqual(rendered.paragraphs[1].text, "Example Client")
        self.assertEqual(
            self._paragraph_after_heading(rendered, "Executive Summary"),
            "Composed executive summary.",
        )
        self.assertEqual(
            self._paragraph_after_heading(rendered, "Introduction"),
            "Composed introduction.",
        )

    def test_render_report_skips_objectives_when_template_has_no_objectives_heading(
        self,
    ):
        with TemporaryDirectory() as tmp_dir:
            template_path = Path(tmp_dir) / "template.docx"
            output_path = Path(tmp_dir) / "rendered.docx"
            self._write_front_matter_template(template_path, include_objectives=False)

            gr.render_report(
                self._validated_report(
                    findings=[],
                    introduction_overview="Composed executive summary.",
                    introduction_scope_bridge="Composed introduction.",
                    objectives=[
                        "Validate assets.",
                        "Prioritise remediation.",
                    ],
                ),
                str(template_path),
                str(output_path),
            )

            rendered = Document(output_path)

        self.assertEqual(
            self._paragraph_after_heading(rendered, "Executive Summary"),
            "Composed executive summary.",
        )
        self.assertEqual(
            self._paragraph_after_heading(rendered, "Introduction"),
            "Composed introduction.",
        )
        self.assertNotIn(
            "Validate assets. Prioritise remediation.",
            [paragraph.text for paragraph in rendered.paragraphs],
        )

    def test_render_report_clears_objectives_placeholder_when_objectives_are_empty(
        self,
    ):
        with TemporaryDirectory() as tmp_dir:
            template_path = Path(tmp_dir) / "template.docx"
            output_path = Path(tmp_dir) / "rendered.docx"
            self._write_front_matter_template(template_path)

            gr.render_report(
                self._validated_report(
                    findings=[],
                    objectives=["[PLACEHOLDER]", " ", None, "nan"],
                ),
                str(template_path),
                str(output_path),
            )

            rendered = Document(output_path)

        self.assertEqual(self._paragraph_after_heading(rendered, "Objectives"), "")

    def test_render_report_detailed_findings_follow_composed_outline_groups_for_multi_subnet_findings(
        self,
    ):
        with TemporaryDirectory() as tmp_dir:
            template_path = Path(tmp_dir) / "template.docx"
            output_path = Path(tmp_dir) / "rendered.docx"
            self._write_renderer_template(
                template_path,
                include_summary_table=False,
            )

            finding_a = self._validated_finding(
                id="VAPT-100",
                name="Raw Scanner Title A",
                severity="Medium",
                category="[PLACEHOLDER]",
                affected_assets="10.30.0.5:443",
                taxonomy_family="network",
                taxonomy_label="Network Security Findings",
                subnet_key="10.30.0.0/24",
                subnet_label="Subnet 10.30.0.0/24",
                display_title="Validated Subnet A Exposure",
            )
            finding_b = self._validated_finding(
                id="VAPT-200",
                name="Raw Scanner Title B",
                severity="Medium",
                category="[PLACEHOLDER]",
                affected_assets="10.40.0.7:443",
                taxonomy_family="network",
                taxonomy_label="Network Security Findings",
                subnet_key="10.40.0.0/24",
                subnet_label="Subnet 10.40.0.0/24",
                display_title="Validated Subnet B Exposure",
            )

            gr.render_report(
                self._validated_report(
                    findings=[finding_b, finding_a],
                    outline_groups=[
                        {
                            "taxonomy_family": "network",
                            "taxonomy_label": "Network Security Findings",
                            "subnet_key": "10.30.0.0/24",
                            "subnet_label": "Subnet 10.30.0.0/24",
                            "highest_severity": "Medium",
                            "finding_count": 1,
                            "finding_ids": ["VAPT-100"],
                        },
                        {
                            "taxonomy_family": "network",
                            "taxonomy_label": "Network Security Findings",
                            "subnet_key": "10.40.0.0/24",
                            "subnet_label": "Subnet 10.40.0.0/24",
                            "highest_severity": "Medium",
                            "finding_count": 1,
                            "finding_ids": ["VAPT-200"],
                        },
                    ],
                ),
                str(template_path),
                str(output_path),
            )

            rendered = Document(output_path)

        detail_texts = self._paragraph_texts_after_heading(
            rendered, "Detailed Observation"
        )
        self.assertEqual(detail_texts[0], "Network Security Findings")
        self.assertEqual(detail_texts[1], "Subnet 10.30.0.0/24")
        self.assertIn("Subnet 10.40.0.0/24", detail_texts)
        self.assertNotIn("[PLACEHOLDER]", detail_texts)

        finding_tables = self._rendered_finding_tables(rendered)
        self.assertEqual(len(finding_tables), 2)
        self.assertEqual(
            finding_tables[0][1][1],
            "Validated Subnet A Exposure (VAPT-100)",
        )
        self.assertEqual(
            finding_tables[1][1][1],
            "Validated Subnet B Exposure (VAPT-200)",
        )

    def test_render_report_summary_table_uses_composed_display_fields_when_present(
        self,
    ):
        with TemporaryDirectory() as tmp_dir:
            template_path = Path(tmp_dir) / "template.docx"
            output_path = Path(tmp_dir) / "rendered.docx"
            self._write_renderer_template(
                template_path,
                include_summary_table=True,
            )

            finding = self._validated_finding(
                id="VAPT-321",
                name="Raw Scanner Title",
                category="Legacy Category",
                affected_assets=", ".join(
                    f"admin{index}.example.com:8443" for index in range(1, 13)
                ),
                display_title="Validated Administrative Interface Exposure",
                affected_assets_short="admin1.example.com:8443, admin2.example.com:8443, plus 10 more",
            )

            gr.render_report(
                self._validated_report(findings=[finding]),
                str(template_path),
                str(output_path),
                include_summary_table=True,
            )

            rendered = Document(output_path)

        summary_rows = self._summary_table_rows(rendered)
        self.assertEqual(len(summary_rows), 2)
        self.assertEqual(len(summary_rows[1]), 7)
        self.assertEqual(
            summary_rows[1][3], "Validated Administrative Interface Exposure"
        )
        self.assertEqual(
            summary_rows[1][6],
            "admin1.example.com:8443, admin2.example.com:8443, plus 10 more",
        )

    def test_render_report_handles_list_valued_asset_fallbacks(self):
        with TemporaryDirectory() as tmp_dir:
            template_path = Path(tmp_dir) / "template.docx"
            output_path = Path(tmp_dir) / "rendered.docx"
            self._write_renderer_template(
                template_path,
                include_summary_table=True,
            )

            finding = self._validated_finding(
                id="VAPT-654",
                name="List-Valued Asset Fallback",
                affected_assets=[
                    "10.99.0.10:443",
                    "10.99.0.11:443",
                ],
            )

            gr.render_report(
                self._validated_report(findings=[finding]),
                str(template_path),
                str(output_path),
                include_summary_table=True,
            )

            rendered = Document(output_path)

        summary_rows = self._summary_table_rows(rendered)
        self.assertEqual(summary_rows[1][6], "10.99.0.10:443, 10.99.0.11:443")

        finding_tables = self._rendered_finding_tables(rendered)
        self.assertEqual(len(finding_tables), 1)
        self.assertEqual(
            finding_tables[0][0][2],
            "10.99.0.10:443, 10.99.0.11:443",
        )

    def test_render_report_summary_table_prefers_structured_group_labels(self):
        with TemporaryDirectory() as tmp_dir:
            template_path = Path(tmp_dir) / "template.docx"
            output_path = Path(tmp_dir) / "rendered.docx"
            self._write_renderer_template(
                template_path,
                include_summary_table=True,
            )

            finding = self._validated_finding(
                id="VAPT-777",
                name="Structured Group Label",
                category="[PLACEHOLDER]",
                taxonomy_label="Network Security Findings",
                subnet_label="Subnet 10.30.0.0/24",
            )

            gr.render_report(
                self._validated_report(findings=[finding]),
                str(template_path),
                str(output_path),
                include_summary_table=True,
            )

            rendered = Document(output_path)

        summary_rows = self._summary_table_rows(rendered)
        self.assertEqual(
            summary_rows[1][1],
            "Network Security Findings / Subnet 10.30.0.0/24",
        )

    def test_render_report_summary_table_follows_outline_group_order(self):
        with TemporaryDirectory() as tmp_dir:
            template_path = Path(tmp_dir) / "template.docx"
            output_path = Path(tmp_dir) / "rendered.docx"
            self._write_renderer_template(
                template_path,
                include_summary_table=True,
            )

            finding_a = self._validated_finding(
                id="VAPT-110",
                name="Raw Scanner Title A",
                display_title="Validated Summary Finding A",
                category="[PLACEHOLDER]",
                affected_assets="10.10.0.10:443",
                taxonomy_family="application",
                taxonomy_label="Application Security Findings",
                subnet_key="10.10.0.0/24",
                subnet_label="Subnet 10.10.0.0/24",
            )
            finding_b = self._validated_finding(
                id="VAPT-220",
                name="Raw Scanner Title B",
                display_title="Validated Summary Finding B",
                category="[PLACEHOLDER]",
                affected_assets="10.20.0.20:443",
                taxonomy_family="network",
                taxonomy_label="Network Security Findings",
                subnet_key="10.20.0.0/24",
                subnet_label="Subnet 10.20.0.0/24",
            )

            gr.render_report(
                self._validated_report(
                    findings=[finding_b, finding_a],
                    outline_groups=[
                        {
                            "taxonomy_family": "application",
                            "taxonomy_label": "Application Security Findings",
                            "subnet_key": "10.10.0.0/24",
                            "subnet_label": "Subnet 10.10.0.0/24",
                            "highest_severity": "Medium",
                            "finding_count": 1,
                            "finding_ids": ["VAPT-110"],
                        },
                        {
                            "taxonomy_family": "network",
                            "taxonomy_label": "Network Security Findings",
                            "subnet_key": "10.20.0.0/24",
                            "subnet_label": "Subnet 10.20.0.0/24",
                            "highest_severity": "Medium",
                            "finding_count": 1,
                            "finding_ids": ["VAPT-220"],
                        },
                    ],
                ),
                str(template_path),
                str(output_path),
                include_summary_table=True,
            )

            rendered = Document(output_path)

        summary_rows = self._summary_table_rows(rendered)
        self.assertEqual(
            [summary_rows[1][3], summary_rows[2][3]],
            ["Validated Summary Finding A", "Validated Summary Finding B"],
        )

    def test_render_report_colors_summary_table_severity_cells(self):
        with TemporaryDirectory() as tmp_dir:
            template_path = Path(tmp_dir) / "template.docx"
            output_path = Path(tmp_dir) / "rendered.docx"
            self._write_renderer_template(
                template_path,
                include_summary_table=True,
            )

            finding = self._validated_finding(
                id="VAPT-999",
                severity="High",
                name="Severity Color Test",
            )

            gr.render_report(
                self._validated_report(findings=[finding]),
                str(template_path),
                str(output_path),
                include_summary_table=True,
            )

            rendered = Document(output_path)

        summary_table = self._summary_table(rendered)
        severity_cell = summary_table.rows[1].cells[2]._tc
        tc_pr = severity_cell.find(gr.qn("w:tcPr"))
        self.assertIsNotNone(tc_pr)
        shading = tc_pr.find(gr.qn("w:shd"))
        self.assertIsNotNone(shading)
        self.assertEqual(shading.get(gr.qn("w:fill")), gr._BG_SEV["High"])

    def test_render_report_preserves_template_chart_order_and_updates_chart_data(
        self,
    ):
        template_path = (
            Path("report_runtime/templates/uploaded-base-template.docx").resolve()
        )
        self.assertTrue(template_path.exists())

        findings = [
            self._validated_finding(
                id="VAPT-001",
                severity="Critical",
                affected_assets="10.10.0.10:443",
            ),
            self._validated_finding(
                id="VAPT-002",
                severity="High",
                affected_assets="10.10.0.20:443",
            ),
            self._validated_finding(
                id="VAPT-003",
                severity="Medium",
                affected_assets="10.10.0.30:443",
            ),
            self._validated_finding(
                id="VAPT-004",
                severity="Medium",
                name="Distinct Medium Exposure",
                affected_assets="10.10.0.40:443",
            ),
        ]

        with TemporaryDirectory() as tmp_dir:
            output_path = Path(tmp_dir) / "rendered.docx"
            gr.render_report(
                self._validated_report(
                    findings=findings,
                    total_critical=1,
                    total_high=1,
                    total_medium=2,
                    total_low=0,
                    total_findings=4,
                ),
                str(template_path),
                str(output_path),
                include_summary_table=True,
            )

            chart_idx, summary_idx = self._body_chart_and_summary_indexes(output_path)
            chart_caption_idx = self._body_text_index(
                output_path, "Vulnerabilities Reported Vulnerabilities Reported"
            )
            summary_heading_idx = self._body_text_index(
                output_path, "Summary of Vulnerabilities"
            )
            detail_heading_idx = self._body_text_index(
                output_path, "Detailed Observation – Proof of Concept"
            )
            self.assertLess(summary_heading_idx, summary_idx)
            self.assertLess(summary_idx, chart_idx)
            self.assertLess(summary_idx, chart_caption_idx)
            self.assertLess(chart_idx, detail_heading_idx)
            self.assertLess(chart_caption_idx, detail_heading_idx)

            from lxml import etree as ET
            from openpyxl import load_workbook

            with ZipFile(output_path) as archive:
                chart_tree = ET.fromstring(archive.read("word/charts/chart1.xml"))
                workbook = load_workbook(
                    BytesIO(
                        archive.read("word/embeddings/Microsoft_Excel_Worksheet.xlsx")
                    )
                )

            ns = {"c": "http://schemas.openxmlformats.org/drawingml/2006/chart"}
            values = [
                int(point.findtext("c:v", namespaces=ns))
                for point in chart_tree.findall(".//c:ser/c:val//c:pt", namespaces=ns)
            ]
            self.assertEqual(values, [1, 1, 2, 0])
            self.assertEqual(
                [workbook.active[f"B{row}"].value for row in range(2, 6)],
                [1, 1, 2, 0],
            )

    def test_render_report_fallback_grouping_uses_any_available_structured_label(
        self,
    ):
        with TemporaryDirectory() as tmp_dir:
            template_path = Path(tmp_dir) / "template.docx"
            output_path = Path(tmp_dir) / "rendered.docx"
            self._write_renderer_template(
                template_path,
                include_summary_table=False,
            )

            taxonomy_only = self._validated_finding(
                id="VAPT-610",
                name="Taxonomy Only Grouping",
                category="Legacy Raw Category A",
                affected_assets="10.61.0.10:443",
                taxonomy_label="Structured Taxonomy Heading",
                subnet_label="",
            )
            subnet_only = self._validated_finding(
                id="VAPT-620",
                name="Subnet Only Grouping",
                category="Legacy Raw Category B",
                affected_assets="10.62.0.10:443",
                taxonomy_label="",
                subnet_label="Subnet 10.62.0.0/24",
            )

            gr.render_report(
                self._validated_report(findings=[taxonomy_only, subnet_only]),
                str(template_path),
                str(output_path),
            )

            rendered = Document(output_path)

        detail_texts = self._paragraph_texts_after_heading(
            rendered, "Detailed Observation"
        )
        self.assertIn("Structured Taxonomy Heading", detail_texts)
        self.assertIn("Subnet 10.62.0.0/24", detail_texts)
        self.assertNotIn("Legacy Raw Category A", detail_texts)
        self.assertNotIn("Legacy Raw Category B", detail_texts)

    def test_render_report_detailed_findings_render_when_placeholder_table_is_missing(
        self,
    ):
        with TemporaryDirectory() as tmp_dir:
            template_path = Path(tmp_dir) / "template.docx"
            output_path = Path(tmp_dir) / "rendered.docx"
            self._write_renderer_template(
                template_path,
                include_summary_table=False,
                include_detail_placeholder_table=False,
            )

            finding = self._validated_finding(
                id="VAPT-888",
                name="Fallback Detail Insertion",
                category="Legacy Network - 10.70.0.0/24",
                affected_assets="10.70.0.20:443",
            )

            gr.render_report(
                self._validated_report(findings=[finding]),
                str(template_path),
                str(output_path),
            )

            rendered = Document(output_path)

        detail_texts = self._paragraph_texts_after_heading(
            rendered, "Detailed Observation"
        )
        self.assertIn("Legacy Network - 10.70.0.0/24", detail_texts)

        finding_tables = self._rendered_finding_tables(rendered)
        self.assertEqual(len(finding_tables), 1)
        self.assertEqual(
            finding_tables[0][1][1],
            "Fallback Detail Insertion (VAPT-888)",
        )

    def test_render_report_falls_back_to_raw_grouping_when_outline_groups_are_absent(
        self,
    ):
        with TemporaryDirectory() as tmp_dir:
            template_path = Path(tmp_dir) / "template.docx"
            output_path = Path(tmp_dir) / "rendered.docx"
            self._write_renderer_template(
                template_path,
                include_summary_table=False,
            )

            gr.render_report(
                self._validated_report(
                    findings=[
                        self._validated_finding(
                            id="VAPT-555",
                            name="Legacy Render Path Finding",
                            category="Legacy Network - 10.50.0.0/24",
                            affected_assets="10.50.0.10:443",
                        )
                    ]
                ),
                str(template_path),
                str(output_path),
            )

            rendered = Document(output_path)

        detail_texts = self._paragraph_texts_after_heading(
            rendered, "Detailed Observation"
        )
        self.assertIn("Legacy Network - 10.50.0.0/24", detail_texts)

        finding_tables = self._rendered_finding_tables(rendered)
        self.assertEqual(len(finding_tables), 1)
        self.assertEqual(
            finding_tables[0][1][1],
            "Legacy Render Path Finding (VAPT-555)",
        )

    def test_compose_report_narrative_outline_groups_are_deterministic(self):
        findings = [
            self._validated_finding(
                id="VAPT-003",
                name="Deprecated TLS Protocol",
                severity="Medium",
                category="Network - 10.10.0.0/24",
                affected_assets="10.10.0.12:443",
            ),
            self._validated_finding(
                id="VAPT-002",
                name="Weak Administrative Authentication",
                severity="High",
                category="Application - 10.20.0.0/24",
                affected_assets="10.20.0.50:8443",
                control_name="Administrative Access Control",
                control_objective="Restrict administrative access to authorised users.",
            ),
            self._validated_finding(
                id="VAPT-001",
                name="Critical Network Exposure",
                severity="Critical",
                category="Network - 10.10.0.0/24",
                affected_assets="10.10.0.10:443",
            ),
        ]

        composed = gr.compose_report_narrative(
            self._validated_report(findings=findings)
        )
        reversed_composed = gr.compose_report_narrative(
            self._validated_report(findings=list(reversed(copy.deepcopy(findings))))
        )

        self.assertEqual(
            composed["outline_groups"], reversed_composed["outline_groups"]
        )
        self.assertEqual(
            [
                (group["taxonomy_family"], group["subnet_key"], group["finding_ids"])
                for group in composed["outline_groups"]
            ],
            [
                ("application", "10.20.0.0/24", ["VAPT-002"]),
                ("network", "10.10.0.0/24", ["VAPT-001", "VAPT-003"]),
            ],
        )

    def test_empty_structured_severity_stays_informational(self):
        self.assertEqual(gr._structured_normalize_severity(""), "Informational")
        self.assertEqual(gr._structured_normalize_severity(None), "Informational")

    def test_clean_cve_cvss_omits_cvss_from_report_display(self):
        self.assertEqual(gr._clean_cve_cvss("CVE-1999-0517", "7.5"), "CVE-1999-0517")
        self.assertEqual(gr._clean_cve_cvss("CWE-79", "6.1"), "CWE-79")
        self.assertEqual(gr._clean_cve_cvss("", "10.0"), "N/A")

    def test_false_positive_check_is_review_not_final_disposition(self):
        finding = {
            "name": "SSL Version 2 and 3 Protocol Detection",
            "observation": "New",
            "remediation_status": "Open",
            "risk_status": "False Positive Check",
            "description": "Potential legacy protocol support requires review.",
            "severity": "Critical",
            "cvss": "10.0",
        }
        self.assertFalse(gr.is_false_positive(finding))
        self.assertEqual(
            gr._structured_status_family(
                finding["observation"],
                finding["remediation_status"],
                finding["risk_status"],
            ),
            "fp_review",
        )

    def test_structured_family_grouping_uses_handmade_titles(self):
        self.assertEqual(
            gr._structured_group_and_title(
                "Jenkins LTS < 2.492.3 / Jenkins weekly < 2.504 Multiple Vulnerabilities"
            ),
            (
                "family:jenkins-advisories",
                "Jenkins LTS <2.426.3/Jenkins weekly <2.442 Multiple Vulnerabilities",
            ),
        )
        self.assertEqual(
            gr._structured_group_and_title("Web Server Directory Enumeration"),
            ("family:web-directory-enum", "Browsable Web Directories"),
        )
        self.assertEqual(
            gr._structured_group_and_title(
                "Elasticsearch Unrestricted Access Information Disclosure"
            ),
            (
                "name:elasticsearch-open-access",
                "Unauthorized access and Information Disclosure Vulnerability",
            ),
        )

    def test_row_aware_nginx_mapping_requires_vulnerable_version(self):
        vulnerable_row = {
            "plugin output": "URL : http://host\nVersion : 1.16.1\nsource : Server: nginx/1.16.1"
        }
        safe_row = {
            "plugin output": "URL : http://host\nVersion : 1.18.0\nsource : Server: nginx/1.18.0"
        }

        self.assertEqual(
            gr._structured_group_and_title_for_row(
                vulnerable_row, "nginx HTTP Server Detection"
            ),
            ("name:nginx-version-disclosure", "nginx < 1.17.7 Information Disclosure"),
        )
        self.assertEqual(
            gr._structured_group_and_title_for_row(
                safe_row, "nginx HTTP Server Detection"
            ),
            ("name:nginx http server detection", "nginx HTTP Server Detection"),
        )

    def test_generate_per_vuln_resolves_findings_before_dedupe(self):
        finding_payload = {
            "id": "tmp-001",
            "vuln_id": 7,
            "name": "TLS Weak Cipher Suites",
            "severity": "Medium",
            "cvss": "5.3",
            "cve": "[INSUFFICIENT DATA]",
            "affected_assets": "10.10.10.10:443",
            "category": "Network",
            "description": "Weak cipher suites remain enabled on the exposed TLS service.",
            "business_impact": "Attackers may weaken the transport protection applied to sensitive traffic.",
            "proof_of_concept": "nmap --script ssl-enum-ciphers 10.10.10.10 -p 443",
            "remediation": "Disable weak cipher suites and restrict the service to approved cryptographic settings.",
            "control_objective": "[PLACEHOLDER]",
            "control_name": "[PLACEHOLDER]",
            "audit_requirement": "[INSUFFICIENT DATA]",
            "reference": "[PLACEHOLDER]",
            "observation": "New",
            "remediation_status": "Open",
            "risk_status": "Open",
        }
        llm_responses = [
            json.dumps({"client_name": "Example Client"}),
            json.dumps(
                {
                    "findings": [
                        {
                            "vuln_id": 7,
                            "short_name": "TLS Weak Cipher Suites",
                            "raw_block": "TLS service exposes weak cipher suites.",
                        }
                    ]
                }
            ),
            json.dumps(finding_payload),
        ]
        dedupe_inputs = []

        def fake_dedupe(findings):
            dedupe_inputs.extend(copy.deepcopy(findings))
            return findings

        with TemporaryDirectory() as tmp_dir:
            scan_path = Path(tmp_dir) / "scan.txt"
            scan_path.write_text("scanner output", encoding="utf-8")

            with (
                patch.object(
                    gr,
                    "load_config",
                    return_value={
                        "limits": {"max_input_words": 50},
                        "llm": {"concurrency": 1},
                        "paths": {"log_dir": tmp_dir},
                    },
                ),
                patch.object(gr, "_call_llm_generic", side_effect=llm_responses),
                patch.object(gr, "dedupe_findings", side_effect=fake_dedupe),
                patch.object(gr, "validate_json_schema", side_effect=lambda data: data),
                patch.object(
                    gr, "split_false_positives", side_effect=lambda data: (data, [])
                ),
                patch.object(
                    gr,
                    "_structured_finding_needs_llm_lookup",
                    return_value=False,
                ),
                patch.object(
                    gr, "compose_report_narrative", side_effect=lambda data: data
                ),
            ):
                data, raw_texts, fps = gr.generate_per_vuln(str(scan_path))

        self.assertEqual(len(raw_texts), 1)
        self.assertEqual(fps, [])
        self.assertEqual(data["findings"][0]["name"], "TLS Weak Cipher Suites")
        self.assertEqual(len(dedupe_inputs), 1)
        self.assertEqual(
            dedupe_inputs[0]["control_objective"],
            "Ensure only strong, up-to-date TLS protocols and cipher suites are used.",
        )
        self.assertEqual(
            dedupe_inputs[0]["control_name"],
            "TLS Configuration Hardening",
        )
        self.assertEqual(
            dedupe_inputs[0]["reference"],
            "OWASP Web Security Top 10, SANS25; CWE-326",
        )
        self.assertTrue(dedupe_inputs[0]["audit_requirement"].startswith("Verify"))
        self.assertIn("cipher suites", dedupe_inputs[0]["audit_requirement"])

    def test_generate_per_vuln_preserves_finding_when_reference_is_list(self):
        finding_payload = {
            "id": "tmp-002",
            "vuln_id": 8,
            "name": "TLS Weak Cipher Suites",
            "severity": "Medium",
            "cvss": "5.3",
            "cve": "[INSUFFICIENT DATA]",
            "affected_assets": "10.10.10.11:443",
            "category": "Network",
            "description": "Weak cipher suites remain enabled on the exposed TLS service.",
            "business_impact": "Attackers may weaken the transport protection applied to sensitive traffic.",
            "proof_of_concept": "nmap --script ssl-enum-ciphers 10.10.10.11 -p 443",
            "remediation": "Disable weak cipher suites and restrict the service to approved cryptographic settings.",
            "control_objective": "[PLACEHOLDER]",
            "control_name": "[PLACEHOLDER]",
            "audit_requirement": "[INSUFFICIENT DATA]",
            "reference": [
                "Vendor TLS Hardening Guide",
                "CWE-326",
            ],
            "observation": "New",
            "remediation_status": "Open",
            "risk_status": "Open",
        }
        llm_responses = [
            json.dumps({"client_name": "Example Client"}),
            json.dumps(
                {
                    "findings": [
                        {
                            "vuln_id": 8,
                            "short_name": "TLS Weak Cipher Suites",
                            "raw_block": "TLS service exposes weak cipher suites.",
                        }
                    ]
                }
            ),
            json.dumps(finding_payload),
        ]
        dedupe_inputs = []

        def fake_dedupe(findings):
            dedupe_inputs.extend(copy.deepcopy(findings))
            return findings

        with TemporaryDirectory() as tmp_dir:
            scan_path = Path(tmp_dir) / "scan.txt"
            scan_path.write_text("scanner output", encoding="utf-8")

            with (
                patch.object(
                    gr,
                    "load_config",
                    return_value={
                        "limits": {"max_input_words": 50},
                        "llm": {"concurrency": 1},
                        "paths": {"log_dir": tmp_dir},
                    },
                ),
                patch.object(gr, "_call_llm_generic", side_effect=llm_responses),
                patch.object(gr, "dedupe_findings", side_effect=fake_dedupe),
                patch.object(gr, "validate_json_schema", side_effect=lambda data: data),
                patch.object(
                    gr, "split_false_positives", side_effect=lambda data: (data, [])
                ),
                patch.object(
                    gr, "compose_report_narrative", side_effect=lambda data: data
                ),
            ):
                data, raw_texts, fps = gr.generate_per_vuln(str(scan_path))

        self.assertEqual(len(raw_texts), 1)
        self.assertEqual(fps, [])
        self.assertEqual(len(data["findings"]), 1)
        self.assertEqual(data["findings"][0]["name"], "TLS Weak Cipher Suites")
        self.assertEqual(len(dedupe_inputs), 1)
        self.assertEqual(
            dedupe_inputs[0]["reference"],
            "Vendor TLS Hardening Guide; CWE-326",
        )

    def test_merge_structured_lookup_result_preserves_specific_control_name(self):
        prior = self._validated_finding(control_name="TLS Configuration Hardening")

        merged = gr._merge_structured_lookup_result(
            prior,
            {
                "control_name": gr._GENERIC_CONTROL_NAME_FALLBACK,
                "reference": prior["reference"],
            },
        )

        self.assertEqual(merged["control_name"], "TLS Configuration Hardening")

    def test_merge_structured_lookup_result_preserves_specific_reference(self):
        prior = self._validated_finding(
            reference="OWASP Top 10 A02:2021 - Cryptographic Failures; CWE-326"
        )

        merged = gr._merge_structured_lookup_result(
            prior,
            {
                "control_name": prior["control_name"],
                "reference": gr._GENERIC_REFERENCE_FALLBACK,
            },
        )

        self.assertEqual(
            merged["reference"],
            "OWASP Top 10 A02:2021 - Cryptographic Failures; CWE-326",
        )

    def test_merge_structured_lookup_result_preserves_source_of_truth_fields(self):
        prior = self._validated_finding(
            affected_assets="10.10.10.10:443, 10.10.10.11:443",
            affected_assets_raw=["10.10.10.10:443", "10.10.10.11:443"],
            affected_assets_short="10.10.10.10:443, 10.10.10.11:443",
        )

        merged = gr._merge_structured_lookup_result(
            prior,
            {
                "name": "Mutated Title",
                "severity": "Low",
                "affected_assets": "1.1.1.1:443",
                "observation": "Changed",
                "remediation_status": "Closed",
                "risk_status": "Accepted",
            },
        )

        self.assertEqual(merged["name"], prior["name"])
        self.assertEqual(merged["severity"], prior["severity"])
        self.assertEqual(merged["affected_assets"], prior["affected_assets"])
        self.assertEqual(merged["observation"], prior["observation"])
        self.assertEqual(merged["remediation_status"], prior["remediation_status"])
        self.assertEqual(merged["risk_status"], prior["risk_status"])

    def test_merge_structured_lookup_result_allows_generic_values_without_better_prior(
        self,
    ):
        prior = self._validated_finding(
            control_name=gr._GENERIC_CONTROL_NAME_FALLBACK,
            reference=gr._GENERIC_REFERENCE_FALLBACK,
        )

        merged = gr._merge_structured_lookup_result(
            prior,
            {
                "control_name": gr._GENERIC_CONTROL_NAME_FALLBACK,
                "reference": gr._GENERIC_REFERENCE_FALLBACK,
            },
        )

        self.assertEqual(merged["control_name"], gr._GENERIC_CONTROL_NAME_FALLBACK)
        self.assertEqual(merged["reference"], gr._GENERIC_REFERENCE_FALLBACK)

    def test_structured_finding_needs_llm_lookup_for_missing_business_impact_and_category(
        self,
    ):
        finding = self._validated_finding(
            business_impact="[INSUFFICIENT DATA]",
            category="[INSUFFICIENT DATA]",
        )

        self.assertTrue(gr._structured_finding_needs_llm_lookup(finding))

    def test_merge_structured_lookup_result_preserves_richer_narrative_fields(
        self,
    ):
        prior = self._validated_finding(
            description=(
                "Legacy TLS 1.0 and TLS 1.1 remain enabled on the public HTTPS "
                "listener, and anonymous cipher negotiation is still exposed to "
                "unauthenticated clients."
            ),
            business_impact=(
                "Attackers may weaken negotiated transport protections, intercept "
                "sensitive sessions, and preserve a broader downgrade path on the "
                "internet-facing service."
            ),
            control_objective=(
                "Ensure internet-facing services expose only approved TLS protocol "
                "versions and hardened cipher configurations aligned to enterprise "
                "encryption standards."
            ),
            audit_requirement=(
                "Verify deprecated TLS versions and anonymous or weak cipher suites "
                "are disabled across exposed listeners, and confirm the hardened "
                "baseline remains enforced."
            ),
            remediation=(
                "Disable TLS 1.0/TLS 1.1 and remove anonymous or weak cipher suites "
                "from the exposed service configuration before the next deployment "
                "window."
            ),
        )

        merged = gr._merge_structured_lookup_result(
            prior,
            {
                "description": "TLS is weak on this service.",
                "business_impact": "Traffic could be exposed.",
                "control_objective": "Improve TLS security.",
                "audit_requirement": "Review TLS settings.",
                "remediation": "Harden TLS.",
                "control_name": prior["control_name"],
                "reference": prior["reference"],
            },
        )

        self.assertEqual(merged["description"], prior["description"])
        self.assertEqual(merged["business_impact"], prior["business_impact"])
        self.assertEqual(merged["control_objective"], prior["control_objective"])
        self.assertEqual(merged["audit_requirement"], prior["audit_requirement"])
        self.assertEqual(merged["remediation"], prior["remediation"])

    def test_merge_structured_lookup_result_allows_more_specific_narrative_fields(
        self,
    ):
        prior = self._validated_finding(
            description="Legacy TLS protocols remain enabled on the exposed service.",
            business_impact="Attackers may weaken transport protections.",
            control_objective="Ensure approved transport encryption standards are exposed.",
            audit_requirement="Verify unsupported TLS versions remain disabled.",
            remediation="Disable deprecated TLS protocols and weak cipher support.",
        )
        result = {
            "description": (
                "Legacy TLS 1.0 and TLS 1.1 remain enabled on the public HTTPS "
                "listener, and the server still allows downgradeable handshake "
                "paths from unauthenticated clients."
            ),
            "business_impact": (
                "An on-path attacker could force weaker handshakes, reduce session "
                "confidentiality for transmitted credentials, and preserve a wider "
                "internet-facing attack surface."
            ),
            "control_objective": (
                "Ensure external services restrict TLS exposure to approved protocol "
                "versions, strong cipher suites, and enterprise-approved transport "
                "hardening baselines."
            ),
            "audit_requirement": (
                "Confirm TLS 1.0/TLS 1.1 and weak cipher suites are disabled on all "
                "public listeners, and review configuration evidence showing the "
                "approved hardening baseline is enforced."
            ),
            "remediation": (
                "Disable TLS 1.0/TLS 1.1, remove anonymous and weak cipher suites, "
                "and redeploy the HTTPS service with the approved hardened policy."
            ),
            "control_name": prior["control_name"],
            "reference": prior["reference"],
        }

        merged = gr._merge_structured_lookup_result(prior, result)

        self.assertEqual(merged["description"], result["description"])
        self.assertEqual(merged["business_impact"], result["business_impact"])
        self.assertEqual(merged["control_objective"], result["control_objective"])
        self.assertEqual(merged["audit_requirement"], result["audit_requirement"])
        self.assertEqual(merged["remediation"], result["remediation"])

    def test_generate_from_structured_file_requests_specific_replacement_of_generic_control_defaults(
        self,
    ):
        enriched_payload = {
            "id": "VAPT-001",
            "vuln_id": "",
            "name": "Custom Scanner Finding",
            "severity": "Medium",
            "cvss": "[INSUFFICIENT DATA]",
            "cve": "[INSUFFICIENT DATA]",
            "affected_assets": "10.10.10.20:443",
            "category": "Infrastructure & Network",
            "description": "A custom scanner finding remains present on the exposed service.",
            "business_impact": "Attackers may leverage the exposed weakness to gain additional footholds.",
            "proof_of_concept": "[INSUFFICIENT DATA]",
            "remediation": "Restrict access to the service and harden the exposed configuration.",
            "control_objective": "Ensure the exposed service is restricted to approved access paths.",
            "control_name": "Exposure Surface Reduction",
            "audit_requirement": "Verify the service is restricted to authorized sources and hardened against unnecessary exposure.",
            "reference": "CWE-668; Internal Hardening Standard",
            "observation": "New",
            "remediation_status": "Open",
            "risk_status": "Open",
        }
        captured_prompts = {}

        def fake_llm_call(system_prompt, user_prompt, config, cancel_event, **kwargs):
            captured_prompts["system_prompt"] = system_prompt
            captured_prompts["user_prompt"] = user_prompt
            return json.dumps(enriched_payload)

        csv_body = "\n".join(
            [
                "Name,Severity,Host,Port,Description,Solution",
                (
                    "Custom Scanner Finding,Medium,10.10.10.20,443,"
                    '"A custom scanner finding remains present on the exposed service.",'
                    '"Restrict access to the service and harden the exposed configuration."'
                ),
            ]
        )

        with TemporaryDirectory() as tmp_dir:
            scan_path = Path(tmp_dir) / "structured.csv"
            scan_path.write_text(csv_body, encoding="utf-8")

            with (
                patch.dict(gr._CONTROL_MAP, {}, clear=True),
                patch.object(
                    gr,
                    "_derive_audit_requirement",
                    return_value=gr._GENERIC_AUDIT_REQUIREMENT_FALLBACK,
                ),
                patch.object(gr, "_call_llm_generic", side_effect=fake_llm_call),
                patch.object(
                    gr,
                    "_build_metadata_from_findings",
                    return_value={"client_name": "Example Client"},
                ),
                patch.object(
                    gr, "dedupe_findings", side_effect=lambda findings: findings
                ),
                patch.object(gr, "validate_json_schema", side_effect=lambda data: data),
                patch.object(
                    gr, "split_false_positives", side_effect=lambda data: (data, [])
                ),
                patch.object(
                    gr, "compose_report_narrative", side_effect=lambda data: data
                ),
            ):
                gr._generate_from_structured_file(
                    str(scan_path),
                    {"llm": {"concurrency": 1}},
                    client_context="Example Client",
                    cancel_event=None,
                    progress_callback=None,
                )

        self.assertIn(
            "TREAT GENERIC DEFAULTS AS INCOMPLETE",
            captured_prompts["system_prompt"],
        )
        self.assertIn(
            gr._GENERIC_CONTROL_NAME_FALLBACK,
            captured_prompts["system_prompt"],
        )
        self.assertIn(
            gr._GENERIC_REFERENCE_FALLBACK,
            captured_prompts["system_prompt"],
        )
        self.assertIn(
            "Replace ALL generic defaults",
            captured_prompts["user_prompt"],
        )
        self.assertIn(
            "control_objective",
            captured_prompts["user_prompt"],
        )

    def test_generate_from_structured_file_enriches_generic_control_fallbacks(self):
        enriched_payload = {
            "id": "VAPT-001",
            "vuln_id": "",
            "name": "Custom Scanner Finding",
            "severity": "Medium",
            "cvss": "[INSUFFICIENT DATA]",
            "cve": "[INSUFFICIENT DATA]",
            "affected_assets": "10.10.10.20:443",
            "category": "Infrastructure & Network",
            "description": "A custom scanner finding remains present on the exposed service.",
            "business_impact": "Attackers may leverage the exposed weakness to gain additional footholds.",
            "proof_of_concept": "[INSUFFICIENT DATA]",
            "remediation": "Restrict access to the service and harden the exposed configuration.",
            "control_objective": "Ensure the exposed service is restricted to approved access paths.",
            "control_name": "Exposure Surface Reduction",
            "audit_requirement": "Verify the service is restricted to authorized sources and hardened against unnecessary exposure.",
            "reference": "CWE-668; Internal Hardening Standard",
            "observation": "New",
            "remediation_status": "Open",
            "risk_status": "Open",
        }
        dedupe_inputs = []

        def fake_dedupe(findings):
            dedupe_inputs.extend(copy.deepcopy(findings))
            return findings

        csv_body = "\n".join(
            [
                "Name,Severity,Host,Port,Description,Solution",
                (
                    "Custom Scanner Finding,Medium,10.10.10.20,443,"
                    '"A custom scanner finding remains present on the exposed service.",'
                    '"Restrict access to the service and harden the exposed configuration."'
                ),
            ]
        )

        with TemporaryDirectory() as tmp_dir:
            scan_path = Path(tmp_dir) / "structured.csv"
            scan_path.write_text(csv_body, encoding="utf-8")

            with (
                patch.dict(gr._CONTROL_MAP, {}, clear=True),
                patch.object(
                    gr,
                    "_derive_audit_requirement",
                    return_value=(
                        "Verify the affected service is securely configured, unnecessary exposure is removed, and periodic review and monitoring are in place."
                    ),
                ),
                patch.object(
                    gr,
                    "_call_llm_generic",
                    return_value=json.dumps(enriched_payload),
                ) as llm_call,
                patch.object(
                    gr,
                    "_build_metadata_from_findings",
                    return_value={"client_name": "Example Client"},
                ),
                patch.object(gr, "dedupe_findings", side_effect=fake_dedupe),
                patch.object(gr, "validate_json_schema", side_effect=lambda data: data),
                patch.object(
                    gr, "split_false_positives", side_effect=lambda data: (data, [])
                ),
                patch.object(
                    gr, "compose_report_narrative", side_effect=lambda data: data
                ),
            ):
                data, raw_texts, fps = gr._generate_from_structured_file(
                    str(scan_path),
                    {"llm": {"concurrency": 1}},
                    client_context="",
                    cancel_event=None,
                    progress_callback=None,
                )

        self.assertEqual(llm_call.call_count, 1)
        self.assertEqual(len(raw_texts), 1)
        self.assertEqual(fps, [])
        self.assertEqual(len(data["findings"]), 1)
        self.assertEqual(len(dedupe_inputs), 1)
        self.assertEqual(
            dedupe_inputs[0]["control_objective"],
            "Ensure the exposed service is restricted to approved access paths.",
        )
        self.assertEqual(
            dedupe_inputs[0]["control_name"],
            "Exposure Surface Reduction",
        )
        self.assertEqual(
            dedupe_inputs[0]["audit_requirement"],
            "Verify the service is restricted to authorized sources and hardened against unnecessary exposure.",
        )
        self.assertEqual(
            dedupe_inputs[0]["reference"],
            "CWE-668; Internal Hardening Standard",
        )
        self.assertEqual(
            dedupe_inputs[0]["business_impact"],
            "Attackers may leverage the exposed weakness to gain additional footholds.",
        )

    def test_generate_from_structured_file_logs_derived_findings_to_single_run_log(self):
        enriched_payload = {
            "id": "VAPT-001",
            "vuln_id": "",
            "name": "Custom Scanner Finding",
            "severity": "Medium",
            "cvss": "[INSUFFICIENT DATA]",
            "cve": "[INSUFFICIENT DATA]",
            "affected_assets": "10.10.10.20:443",
            "category": "Infrastructure & Network",
            "description": "A custom scanner finding remains present on the exposed service.",
            "business_impact": "Attackers may leverage the exposed weakness to gain additional footholds.",
            "proof_of_concept": "[INSUFFICIENT DATA]",
            "remediation": "Restrict access to the service and harden the exposed configuration.",
            "control_objective": "Ensure the exposed service is restricted to approved access paths.",
            "control_name": "Exposure Surface Reduction",
            "audit_requirement": "Verify the service is restricted to authorized sources and hardened against unnecessary exposure.",
            "reference": "CWE-668; Internal Hardening Standard",
            "observation": "New",
            "remediation_status": "Open",
            "risk_status": "Open",
        }
        csv_body = "\n".join(
            [
                "Name,Severity,Host,Port,Description,Solution",
                (
                    "Custom Scanner Finding,Medium,10.10.10.20,443,"
                    '"A custom scanner finding remains present on the exposed service.",'
                    '"Restrict access to the service and harden the exposed configuration."'
                ),
            ]
        )

        with TemporaryDirectory() as tmp_dir:
            scan_path = Path(tmp_dir) / "structured.csv"
            scan_path.write_text(csv_body, encoding="utf-8")

            with (
                patch.dict(gr._CONTROL_MAP, {}, clear=True),
                patch.object(
                    gr,
                    "_call_llm_generic",
                    return_value=json.dumps(enriched_payload),
                ),
                patch.object(
                    gr,
                    "_build_metadata_from_findings",
                    return_value={"client_name": "Example Client"},
                ),
                patch.object(gr, "dedupe_findings", side_effect=lambda findings: findings),
                patch.object(gr, "validate_json_schema", side_effect=lambda data: data),
                patch.object(
                    gr, "split_false_positives", side_effect=lambda data: (data, [])
                ),
                patch.object(
                    gr,
                    "_structured_finding_needs_llm_lookup",
                    return_value=False,
                ),
                patch.object(
                    gr, "compose_report_narrative", side_effect=lambda data: data
                ),
            ):
                gr._generate_from_structured_file(
                    str(scan_path),
                    {"llm": {"concurrency": 1}, "paths": {"log_dir": tmp_dir}},
                    client_context="",
                    cancel_event=None,
                    progress_callback=None,
                )

            log_files = list(Path(tmp_dir).glob("*_run_log.json"))
            self.assertEqual(len(log_files), 1)
            payload = json.loads(log_files[0].read_text(encoding="utf-8"))

        derived_events = [
            event
            for event in payload["events"]
            if event.get("type") == "structured_findings_derived"
        ]
        parsed_events = [
            event
            for event in payload["events"]
            if event.get("type") == "nessus_findings_parsed"
        ]
        self.assertEqual(len(parsed_events), 1)
        self.assertEqual(parsed_events[0]["data"]["finding_count"], 1)
        self.assertEqual(parsed_events[0]["data"]["findings"][0]["name"], "Custom Scanner Finding")
        self.assertEqual(len(derived_events), 1)
        derived = derived_events[0]["data"]
        self.assertEqual(derived["finding_count"], 1)
        self.assertEqual(derived["findings"][0]["name"], "Custom Scanner Finding")
        self.assertIn("needs_llm_lookup", derived["findings"][0])
        self.assertEqual(derived["selected_columns"]["name"], "name")

    def test_generate_from_structured_file_skips_second_stage_after_cloud_lookup_by_default(self):
        cloud_finding = {
            "id": "VAPT-001",
            "vuln_id": "1001",
            "name": "Custom Scanner Finding",
            "severity": "Medium",
            "cvss": "[INSUFFICIENT DATA]",
            "cve": "[INSUFFICIENT DATA]",
            "affected_assets": "10.10.10.20:443, 10.10.10.21:443, 10.10.10.22:443",
            "description": "Service is exposed with a weak configuration.",
            "remediation": "Restrict service exposure.",
            "business_impact": "Attackers may leverage the exposed configuration to expand access within the environment.",
            "proof_of_concept": "Validated with direct connectivity to the exposed service.",
            "control_objective": "Ensure the service is only exposed to approved sources and hardened appropriately.",
            "control_name": "Service Exposure Hardening",
            "audit_requirement": "Verify access restrictions and service hardening settings remain enforced for the exposed endpoint.",
            "reference": "Vendor hardening guide",
            "category": "Internal Network - 10.10.10.0/24",
            "observation": "New",
            "remediation_status": "Open",
            "risk_status": "Open",
        }
        csv_body = "\n".join(
            [
                "Name,Severity,Host,Port,Description,Solution",
                (
                    "Custom Scanner Finding,Medium,10.10.10.20,443,"
                    '"Service is exposed with a weak configuration.",'
                    '"Restrict service exposure."'
                ),
            ]
        )

        with TemporaryDirectory() as tmp_dir:
            scan_path = Path(tmp_dir) / "structured.csv"
            scan_path.write_text(csv_body, encoding="utf-8")
            progress_events = []

            with (
                patch("report_tool.lookup.cloud_enrich.lookup_report", return_value={
                    "findings": [cloud_finding],
                    "_lookup_stats": {
                        "total_findings": 1,
                        "cloud_eligible_findings": 1,
                        "resolved_without_cloud": 0,
                    },
                }),
                patch.object(gr, "_call_llm_generic") as llm_call,
                patch.object(
                    gr,
                    "_build_metadata_from_findings",
                    return_value={"client_name": "Example Client"},
                ),
                patch.object(gr, "dedupe_findings", side_effect=lambda findings: findings),
                patch.object(gr, "validate_json_schema", side_effect=lambda data: data),
                patch.object(
                    gr, "split_false_positives", side_effect=lambda data: (data, [])
                ),
                patch.object(
                    gr,
                    "_structured_finding_needs_llm_lookup",
                    return_value=False,
                ),
                patch.object(
                    gr, "compose_report_narrative", side_effect=lambda data: data
                ),
            ):
                data, raw_texts, fps = gr._generate_from_structured_file(
                    str(scan_path),
                    {
                        "llm": {
                            "provider": "openrouter",
                            "model": "openai/gpt-oss-120b:free",
                            "base_url": "https://openrouter.ai/api/v1",
                            "api_key": "session-key",
                            "max_retries": 1,
                        },
                        "paths": {"log_dir": tmp_dir},
                    },
                    client_context="",
                    cancel_event=None,
                    progress_callback=lambda *args: progress_events.append(args),
                )

            llm_call.assert_not_called()
            self.assertEqual(fps, [])
            self.assertEqual(raw_texts, [])
            self.assertEqual(len(data["findings"]), 1)
            self.assertEqual(data["findings"][0]["name"], "Custom Scanner Finding")
            self.assertIn(
                (
                    "enrich",
                    1,
                    1,
                    "All 1 findings were enriched after cloud lookup — no additional second-stage LLM lookup required.",
                ),
                progress_events,
            )

    def test_generate_from_structured_file_raises_when_llm_unavailable_and_critical_fields_missing(
        self,
    ):
        csv_body = "\n".join(
            [
                "Name,Severity,Host,Port,Description,Solution",
                (
                    "Custom Scanner Finding,Medium,10.10.10.20,443,"
                    '"A custom scanner finding remains present on the exposed service.",'
                    '"Restrict access to the service and harden the exposed configuration."'
                ),
            ]
        )

        with TemporaryDirectory() as tmp_dir:
            scan_path = Path(tmp_dir) / "structured.csv"
            scan_path.write_text(csv_body, encoding="utf-8")

            with (
                patch.dict(gr._CONTROL_MAP, {}, clear=True),
                patch.object(
                    gr,
                    "_call_llm_generic",
                    side_effect=RuntimeError("LLM unavailable"),
                ) as llm_call,
                patch.object(
                    gr,
                    "_build_metadata_from_findings",
                    return_value={"client_name": "Example Client"},
                ),
                patch.object(gr, "validate_json_schema", side_effect=lambda data: data),
                patch.object(
                    gr, "split_false_positives", side_effect=lambda data: (data, [])
                ),
                patch.object(
                    gr, "compose_report_narrative", side_effect=lambda data: data
                ),
            ):
                with self.assertRaisesRegex(
                    ValueError,
                    "Structured findings missing required source-of-truth fields",
                ):
                    gr._generate_from_structured_file(
                        str(scan_path),
                        {"llm": {"concurrency": 1}},
                        client_context="",
                        cancel_event=None,
                        progress_callback=None,
                    )

        self.assertEqual(llm_call.call_count, 1)

    def test_generate_from_structured_file_retries_once_after_parse_failure(self):
        enriched_payload = {
            "id": "VAPT-001",
            "vuln_id": "",
            "name": "Custom Scanner Finding",
            "severity": "Medium",
            "cvss": "[INSUFFICIENT DATA]",
            "cve": "[INSUFFICIENT DATA]",
            "affected_assets": "10.10.10.20:443",
            "category": "Infrastructure & Network",
            "description": "A custom scanner finding remains present on the exposed service.",
            "business_impact": "Attackers may leverage the exposed weakness to gain additional footholds.",
            "proof_of_concept": "curl -vk https://10.10.10.20/health",
            "remediation": "Restrict access to the service and harden the exposed configuration.",
            "control_objective": "Ensure the exposed service is restricted to approved access paths.",
            "control_name": "Exposure Surface Reduction",
            "audit_requirement": "Verify the service is restricted to authorized sources and hardened against unnecessary exposure.",
            "reference": "CWE-668; Internal Hardening Standard",
            "observation": "New",
            "remediation_status": "Open",
            "risk_status": "Open",
        }
        dedupe_inputs = []

        def fake_dedupe(findings):
            dedupe_inputs.extend(copy.deepcopy(findings))
            return findings

        csv_body = "\n".join(
            [
                "Name,Severity,Host,Port,Description,Solution",
                (
                    "Custom Scanner Finding,Medium,10.10.10.20,443,"
                    '"A custom scanner finding remains present on the exposed service.",'
                    '"Restrict access to the service and harden the exposed configuration."'
                ),
            ]
        )

        malformed_response = '{"id": "VAPT-001", "name": "Custom Scanner Finding"'

        with TemporaryDirectory() as tmp_dir:
            scan_path = Path(tmp_dir) / "structured.csv"
            scan_path.write_text(csv_body, encoding="utf-8")

            with (
                patch.dict(gr._CONTROL_MAP, {}, clear=True),
                patch.object(
                    gr,
                    "_derive_audit_requirement",
                    return_value=(
                        "Verify the affected service is securely configured, unnecessary exposure is removed, and periodic review and monitoring are in place."
                    ),
                ),
                patch.object(
                    gr,
                    "_call_llm_generic",
                    side_effect=[malformed_response, json.dumps(enriched_payload)],
                ) as llm_call,
                patch.object(
                    gr,
                    "_build_metadata_from_findings",
                    return_value={"client_name": "Example Client"},
                ),
                patch.object(gr, "dedupe_findings", side_effect=fake_dedupe),
                patch.object(gr, "validate_json_schema", side_effect=lambda data: data),
                patch.object(
                    gr, "split_false_positives", side_effect=lambda data: (data, [])
                ),
                patch.object(
                    gr, "compose_report_narrative", side_effect=lambda data: data
                ),
            ):
                data, raw_texts, fps = gr._generate_from_structured_file(
                    str(scan_path),
                    {"llm": {"concurrency": 1}},
                    client_context="",
                    cancel_event=None,
                    progress_callback=None,
                )

        self.assertEqual(llm_call.call_count, 2)
        first_cfg = llm_call.call_args_list[0].args[2]
        retry_cfg = llm_call.call_args_list[1].args[2]
        self.assertGreater(
            retry_cfg["llm"]["max_tokens"],
            first_cfg["llm"]["max_tokens"],
        )
        self.assertEqual(fps, [])
        self.assertEqual(len(data["findings"]), 1)
        self.assertEqual(len(raw_texts), 1)
        self.assertEqual(len(dedupe_inputs), 1)
        self.assertEqual(
            dedupe_inputs[0]["control_objective"],
            "Ensure the exposed service is restricted to approved access paths.",
        )
        self.assertEqual(
            dedupe_inputs[0]["control_name"],
            "Exposure Surface Reduction",
        )
        self.assertEqual(
            dedupe_inputs[0]["reference"],
            "CWE-668; Internal Hardening Standard",
        )
        self.assertEqual(
            dedupe_inputs[0]["business_impact"],
            "Attackers may leverage the exposed weakness to gain additional footholds.",
        )

    def test_generate_from_structured_file_raises_after_parse_retry_failure_when_critical_fields_missing(
        self,
    ):
        csv_body = "\n".join(
            [
                "Name,Severity,Host,Port,Description,Solution",
                (
                    "Custom Scanner Finding,Medium,10.10.10.20,443,"
                    '"A custom scanner finding remains present on the exposed service.",'
                    '"Restrict access to the service and harden the exposed configuration."'
                ),
            ]
        )

        malformed_response = '{"id": "VAPT-001", "name": "Custom Scanner Finding"'
        malformed_retry_response = (
            '{"id": "VAPT-001", "name": "Custom Scanner Finding",'
        )

        with TemporaryDirectory() as tmp_dir:
            scan_path = Path(tmp_dir) / "structured.csv"
            scan_path.write_text(csv_body, encoding="utf-8")

            with (
                patch.dict(gr._CONTROL_MAP, {}, clear=True),
                patch.object(
                    gr,
                    "_call_llm_generic",
                    side_effect=[malformed_response, malformed_retry_response],
                ) as llm_call,
                patch.object(
                    gr,
                    "_build_metadata_from_findings",
                    return_value={"client_name": "Example Client"},
                ),
                patch.object(gr, "validate_json_schema", side_effect=lambda data: data),
                patch.object(
                    gr, "split_false_positives", side_effect=lambda data: (data, [])
                ),
                patch.object(
                    gr, "compose_report_narrative", side_effect=lambda data: data
                ),
            ):
                with self.assertRaisesRegex(
                    ValueError,
                    "Structured findings missing required source-of-truth fields",
                ):
                    gr._generate_from_structured_file(
                        str(scan_path),
                        {"llm": {"concurrency": 1}},
                        client_context="",
                        cancel_event=None,
                        progress_callback=None,
                    )

        self.assertEqual(llm_call.call_count, 2)

    def test_per_vuln_detail_prompt_requires_specific_control_guidance_examples(self):
        prompt = gr.PER_VULN_DETAIL_SYSTEM

        self.assertIn("control_objective", prompt)
        self.assertIn("audit_requirement", prompt)
        self.assertIn("reference", prompt)
        self.assertIn("Example control_objective", prompt)
        self.assertIn("Example audit_requirement", prompt)
        self.assertIn("Example reference", prompt)
        self.assertIn("Write remediation as a specific action", prompt)

    def test_apply_local_normalization_pass_normalizes_severity_and_assets(self):
        data = {
            "findings": [
                {
                    "name": "  Weak TLS Config   ",
                    "severity": "moderate",
                    "affected_assets": "[INSUFFICIENT DATA]",
                    "description": "Observed on 10.10.10.20 and app.example.local:8443.",
                    "proof_of_concept": "curl -vk https://10.10.10.20:443/health",
                }
            ]
        }

        out = gr._apply_local_normalization_pass(data, {"llm": {}, "paths": {}})

        finding = out["findings"][0]
        self.assertEqual(finding["name"], "Weak TLS Config")
        self.assertEqual(finding["severity"], "Medium")
        self.assertIn("10.10.10.20", finding["affected_assets"])
        self.assertIn("10.10.10.20:443", finding["affected_assets"])
        self.assertIn("app.example.local:8443", finding["affected_assets"])

    def test_generate_from_structured_file_applies_reference_validation_pass(self):
        enriched_payload = {
            "id": "VAPT-001",
            "vuln_id": "",
            "name": "Custom Scanner Finding",
            "severity": "Medium",
            "cvss": "5.0",
            "cve": "CVE-2024-9999",
            "affected_assets": "10.10.10.20:443",
            "category": "Infrastructure & Network",
            "description": "A custom scanner finding remains present on the exposed service.",
            "business_impact": "Attackers may leverage the exposed weakness.",
            "proof_of_concept": "curl -vk https://10.10.10.20/health",
            "remediation": "Patch and harden the exposed service.",
            "control_objective": "Ensure exposed services are hardened.",
            "control_name": "Service Hardening",
            "audit_requirement": "Verify service patch and hardening status.",
            "reference": "OWASP Web Security Top 10, SANS25",
            "observation": "New",
            "remediation_status": "Open",
            "risk_status": "Open",
        }

        csv_body = "\n".join(
            [
                "Name,Severity,Host,Port,Description,Solution",
                (
                    "Custom Scanner Finding,Medium,10.10.10.20,443,"
                    '"A custom scanner finding remains present on the exposed service.",'
                    '"Patch and harden the exposed service."'
                ),
            ]
        )

        with TemporaryDirectory() as tmp_dir:
            scan_path = Path(tmp_dir) / "structured.csv"
            scan_path.write_text(csv_body, encoding="utf-8")

            with (
                patch.dict(gr._CONTROL_MAP, {}, clear=True),
                patch.object(
                    gr,
                    "_call_llm_generic",
                    return_value=json.dumps(enriched_payload),
                ),
                patch.object(
                    gr,
                    "_build_metadata_from_findings",
                    return_value={"client_name": "Example Client"},
                ),
                patch.object(
                    gr, "dedupe_findings", side_effect=lambda findings: findings
                ),
                patch.object(gr, "validate_json_schema", side_effect=lambda data: data),
                patch.object(
                    gr, "split_false_positives", side_effect=lambda data: (data, [])
                ),
                patch.object(
                    gr, "compose_report_narrative", side_effect=lambda data: data
                ),
                patch.object(
                    gr,
                    "_apply_reference_validation_pass",
                    side_effect=lambda data, config: data,
                ) as ref_pass,
            ):
                gr._generate_from_structured_file(
                    str(scan_path),
                    {"llm": {"concurrency": 1}},
                    client_context="",
                    cancel_event=None,
                    progress_callback=None,
                )

        self.assertEqual(ref_pass.call_count, 1)


class DeterministicBackfillTests(unittest.TestCase):
    """Guarantee category + business_impact are populated before validator runs."""

    def test_derive_category_from_single_subnet(self):
        finding = {"affected_assets": "10.10.0.12:443, 10.10.0.13:443"}
        self.assertEqual(
            gr._derive_category_from_finding(finding),
            "Internal Network - 10.10.0.0/24",
        )

    def test_derive_category_from_multiple_subnets(self):
        finding = {"affected_assets": "10.1.0.5, 192.168.1.10"}
        self.assertEqual(
            gr._derive_category_from_finding(finding),
            "Internal Network - Multiple Subnets",
        )

    def test_derive_category_from_url_hostname(self):
        finding = {"affected_assets": "https://app.example.com/login"}
        self.assertEqual(
            gr._derive_category_from_finding(finding),
            "Web Application",
        )

    def test_derive_category_from_keyword_when_no_assets(self):
        finding = {
            "affected_assets": "",
            "name": "TLS 1.0 Protocol Detection",
            "description": "Server accepts deprecated TLS 1.0 cipher suites.",
        }
        self.assertEqual(
            gr._derive_category_from_finding(finding),
            "Cryptographic Configuration",
        )

    def test_derive_category_default_fallback(self):
        finding = {"affected_assets": "", "name": "Unknown Issue", "description": ""}
        self.assertEqual(
            gr._derive_category_from_finding(finding),
            "Network Infrastructure",
        )

    def test_derive_business_impact_per_severity(self):
        for severity in ("Critical", "High", "Medium", "Low", "Informational"):
            with self.subTest(severity=severity):
                text = gr._derive_business_impact_from_finding(
                    {"name": "Sample Finding", "severity": severity}
                )
                self.assertIn("Sample Finding", text)
                self.assertGreater(len(text), 80)

    def test_derive_business_impact_unknown_severity_falls_back_to_medium(self):
        text = gr._derive_business_impact_from_finding(
            {"name": "Edge Case", "severity": "Bogus"}
        )
        self.assertIn("Edge Case", text)
        self.assertIn("defence-in-depth", text)

    def test_fill_missing_fields_backfills_category_and_business_impact(self):
        finding = {
            "name": "SNMP Default Community Strings",
            "severity": "High",
            "affected_assets": "10.20.30.5",
            "category": "[INSUFFICIENT DATA]",
            "business_impact": "[INSUFFICIENT DATA]",
            "description": "SNMP service exposes default community strings.",
        }
        gr.fill_missing_fields(finding)
        self.assertEqual(finding["category"], "Internal Network - 10.20.30.0/24")
        self.assertIn("SNMP Default Community Strings", finding["business_impact"])

    def test_fill_missing_fields_preserves_existing_category_and_business_impact(self):
        finding = {
            "name": "Sample",
            "severity": "Medium",
            "affected_assets": "10.0.0.1",
            "category": "Custom Category Label",
            "business_impact": "Custom impact text already provided.",
        }
        gr.fill_missing_fields(finding)
        self.assertEqual(finding["category"], "Custom Category Label")
        self.assertEqual(
            finding["business_impact"], "Custom impact text already provided."
        )

    def test_validator_passes_after_fill_missing_fields(self):
        finding = {
            "id": "VAPT-001",
            "name": "TLS 1.0 Detected",
            "severity": "Medium",
            "affected_assets": "10.0.0.1:443",
            "description": "Service negotiates TLS 1.0 with weak ciphers and exposes "
            "downgrade paths to attackers on the network segment.",
            "remediation": (
                "Disable TLS 1.0 and 1.1 across all listeners and require TLS 1.2+.\n"
                "If immediate disablement is not feasible, restrict access via "
                "firewall to trusted hosts only.\n"
                "Monitor TLS handshake telemetry for legacy negotiations and alert."
            ),
            "control_objective": "Ensure transport security uses approved TLS versions only.",
            "control_name": "TLS Hardening",
            "audit_requirement": "Verify only TLS 1.2+ negotiates using a TLS scanner.",
            "reference": "OWASP Web Security Top 10, SANS25",
            "category": "[INSUFFICIENT DATA]",
            "business_impact": "[INSUFFICIENT DATA]",
        }
        gr.fill_missing_fields(finding)
        gr._validate_structured_source_fields([finding])


if __name__ == "__main__":
    unittest.main()
