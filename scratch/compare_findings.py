"""Compare Nessus xlsx vulns vs handmade report findings."""

import openpyxl
from docx import Document
import json, os, sys

BASE = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))

# --- 1. Parse Nessus xlsx ---
xlsx_path = os.path.join(BASE, "tests/input/AEG_Subnet1_y06m7b 1.xlsx")
wb = openpyxl.load_workbook(xlsx_path)
ws = wb[wb.sheetnames[0]]
headers = [cell.value for cell in ws[1]]
print("XLSX Headers:", headers)

# Build column index
col = {h: i for i, h in enumerate(headers) if h}

# Gather all rows
nessus_rows = []
for row in ws.iter_rows(min_row=2, max_row=ws.max_row, values_only=True):
    nessus_rows.append(list(row))

# Group by plugin name
from collections import defaultdict

by_name = defaultdict(list)
risk_col = col.get("Risk", col.get("Severity", col.get("Risk Factor")))
name_col = col.get("Name", col.get("Plugin Name"))
host_col = col.get("Host", col.get("IP", col.get("Host IP")))
desc_col = col.get("Synopsis", col.get("Description"))
solution_col = col.get("Solution")
cve_col = col.get("CVE")
cvss_col = col.get(
    "CVSS v2.0 Base Score", col.get("CVSS V2", col.get("CVSS v3.0 Base Score"))
)
plugin_output_col = col.get("Plugin Output")

for r in nessus_rows:
    name = r[name_col] if name_col is not None else None
    if name:
        by_name[name].append(r)

# Filter out None/informational
risk_map = {"Critical": 0, "High": 1, "Medium": 2, "Low": 3}
vuln_groups = {}
for name, rows in by_name.items():
    risk = rows[0][risk_col] if risk_col is not None else None
    if risk and risk in risk_map:
        hosts = sorted(set(str(r[host_col]) for r in rows if r[host_col]))
        vuln_groups[name] = {
            "risk": risk,
            "count": len(rows),
            "hosts": hosts,
            "cve": rows[0][cve_col] if cve_col is not None else None,
            "cvss": rows[0][cvss_col] if cvss_col is not None else None,
        }

print(f"\nNessus unique vulns (excl None/Info): {len(vuln_groups)}")
print("\nNessus vulns sorted by severity:")
for name in sorted(
    vuln_groups, key=lambda n: (risk_map.get(vuln_groups[n]["risk"], 99), n)
):
    v = vuln_groups[name]
    print(
        f"  [{v['risk']:8s}] {name} ({v['count']} hosts: {', '.join(v['hosts'][:5])}{'...' if len(v['hosts'])>5 else ''})"
    )

# --- 2. Parse handmade report ---
print("\n" + "=" * 80)
print("HANDMADE REPORT FINDINGS")
print("=" * 80)
handmade_path = os.path.join(BASE, "tests/input/Internal_VAPT_Report_AEG_Vision.docx")
doc = Document(handmade_path)

handmade_findings = []
for t_idx, table in enumerate(doc.tables):
    rows = []
    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells]
        rows.append(cells)
    if len(rows) >= 10:
        row0_texts = [r[0] for r in rows]
        if any("Vulnerability title" in t or "Severity" in t for t in row0_texts):
            finding = {}
            for r in rows:
                key = r[1] if len(r) > 1 else r[0]
                val = r[2] if len(r) > 2 else (r[1] if len(r) > 1 else "")
                if "Affected" in key:
                    finding["hosts"] = val
                elif "Vulnerability title" in key:
                    finding["name"] = val
                elif "Severity" in key:
                    finding["severity"] = val
                elif "CVE" in key:
                    finding["cve"] = val
                elif "Recommendation" in key:
                    finding["recommendation"] = val
                elif "Status" in key:
                    finding["status"] = val
            if finding.get("name"):
                finding["table_idx"] = t_idx
                handmade_findings.append(finding)

print(f"Total handmade findings: {len(handmade_findings)}")
for i, f in enumerate(handmade_findings, 1):
    print(f"  {i:2d}. [{f.get('severity','?'):8s}] {f['name']}")
    print(f"      Hosts: {f.get('hosts','')[:80]}...")

# --- 3. Parse generated report ---
print("\n" + "=" * 80)
print("GENERATED REPORT FINDINGS")
print("=" * 80)
gen_dir = os.path.join(BASE, "outputs")
gen_files = []
for root, dirs, files in os.walk(gen_dir):
    for f in files:
        if f.endswith(".docx"):
            gen_files.append(os.path.join(root, f))
gen_files.sort(key=lambda x: os.path.getmtime(x), reverse=True)

generated_findings = []
if gen_files:
    latest = gen_files[0]
    print(f"Using: {latest}")
    doc2 = Document(latest)
    for t_idx, table in enumerate(doc2.tables):
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(cells)
        if len(rows) >= 10:
            row0_texts = [r[0] for r in rows]
            if any("Vulnerability title" in t or "Severity" in t for t in row0_texts):
                finding = {}
                for r in rows:
                    key = r[1] if len(r) > 1 else r[0]
                    val = r[2] if len(r) > 2 else (r[1] if len(r) > 1 else "")
                    if "Affected" in key:
                        finding["hosts"] = val
                    elif "Vulnerability title" in key:
                        finding["name"] = val
                    elif "Severity" in key:
                        finding["severity"] = val
                    elif "CVE" in key:
                        finding["cve"] = val
                    elif "Recommendation" in key:
                        finding["recommendation"] = val
                    elif "Status" in key:
                        finding["status"] = val
                if finding.get("name"):
                    finding["table_idx"] = t_idx
                    generated_findings.append(finding)

    print(f"Total generated findings: {len(generated_findings)}")
    for i, f in enumerate(generated_findings, 1):
        print(f"  {i:2d}. [{f.get('severity','?'):8s}] {f['name']}")
        print(f"      Hosts: {f.get('hosts','')[:80]}...")

# --- 4. Gap analysis ---
print("\n" + "=" * 80)
print("GAP ANALYSIS: Handmade vs Generated")
print("=" * 80)

handmade_names = {f["name"].lower().strip() for f in handmade_findings}
generated_names = {f["name"].lower().strip() for f in generated_findings}

in_handmade_not_gen = handmade_names - generated_names
in_gen_not_handmade = generated_names - handmade_names

print(f"\nIn handmade but MISSING from generated ({len(in_handmade_not_gen)}):")
for n in sorted(in_handmade_not_gen):
    match = [f for f in handmade_findings if f["name"].lower().strip() == n]
    if match:
        print(f"  - [{match[0].get('severity','?')}] {match[0]['name']}")

print(f"\nIn generated but NOT in handmade ({len(in_gen_not_handmade)}):")
for n in sorted(in_gen_not_handmade):
    match = [f for f in generated_findings if f["name"].lower().strip() == n]
    if match:
        print(f"  - [{match[0].get('severity','?')}] {match[0]['name']}")

# Fuzzy match for close names
print("\nFuzzy near-matches (possible name variants):")
from difflib import SequenceMatcher

for hn in sorted(in_handmade_not_gen):
    best_score = 0
    best_match = None
    for gn in generated_names:
        score = SequenceMatcher(None, hn, gn).ratio()
        if score > best_score:
            best_score = score
            best_match = gn
    if best_score > 0.5:
        print(f"  Handmade: {hn}")
        print(f"  Generated: {best_match} (similarity: {best_score:.2f})")
        print()

# --- 5. Nessus vulns not in either report ---
print("=" * 80)
print("NESSUS VULNS NOT IN HANDMADE REPORT")
print("=" * 80)
nessus_names_lower = {n.lower().strip() for n in vuln_groups}
not_in_handmade = nessus_names_lower - handmade_names

# Fuzzy check
actually_missing = []
for nn in sorted(not_in_handmade):
    best_score = 0
    for hn in handmade_names:
        score = SequenceMatcher(None, nn, hn).ratio()
        if score > best_score:
            best_score = score
    if best_score < 0.7:
        orig = [n for n in vuln_groups if n.lower().strip() == nn][0]
        v = vuln_groups[orig]
        actually_missing.append((orig, v))

print(f"Nessus vulns with no close match in handmade ({len(actually_missing)}):")
for name, v in sorted(
    actually_missing, key=lambda x: (risk_map.get(x[1]["risk"], 99), x[0])
):
    print(f"  [{v['risk']:8s}] {name} ({v['count']} hosts)")

# --- 6. Host consolidation check ---
print("\n" + "=" * 80)
print("HOST CONSOLIDATION CHECK")
print("=" * 80)
for hf in handmade_findings:
    hn = hf["name"].lower().strip()
    # Find matching nessus vuln
    best_score = 0
    best_nessus = None
    for nn in vuln_groups:
        score = SequenceMatcher(None, hn, nn.lower()).ratio()
        if score > best_score:
            best_score = score
            best_nessus = nn
    if best_score > 0.6 and best_nessus:
        nessus_hosts = vuln_groups[best_nessus]["hosts"]
        handmade_hosts_str = hf.get("hosts", "")
        nessus_host_count = len(nessus_hosts)
        # Rough count from handmade
        print(f"\n  {hf['name'][:60]}...")
        print(
            f"    Nessus hosts ({nessus_host_count}): {', '.join(nessus_hosts[:5])}..."
        )
        print(f"    Handmade hosts: {handmade_hosts_str[:80]}...")
