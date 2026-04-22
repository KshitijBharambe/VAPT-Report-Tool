"""Extract full text content from two DOCX files for comparison."""

import os
import sys
from docx import Document
from docx.table import Table


def extract_docx_content(filepath: str) -> str:
    """Extract all text content from a DOCX file, preserving structure."""
    doc = Document(filepath)
    lines = []

    lines.append(f"File: {os.path.basename(filepath)}")
    lines.append(f"Paragraphs: {len(doc.paragraphs)}")
    lines.append(f"Tables: {len(doc.tables)}")
    lines.append(f"Sections: {len(doc.sections)}")
    lines.append("")

    # Walk through the document body in order to preserve interleaving of
    # paragraphs and tables.
    from docx.oxml.ns import qn

    body = doc.element.body
    table_idx = 0
    para_idx = 0

    tables_by_element = {tbl._element: tbl for tbl in doc.tables}
    paras_by_element = {p._element: p for p in doc.paragraphs}

    for child in body:
        tag = child.tag
        if tag == qn("w:p"):
            p = paras_by_element.get(child)
            if p is None:
                continue
            style_name = p.style.name if p.style else "(no style)"
            text = p.text.strip()
            indent = ""
            if style_name.startswith("List") or "Bullet" in style_name:
                indent = "  • "
            elif style_name.startswith("Heading"):
                level = style_name.replace("Heading", "").strip()
                indent = "#" * (int(level) if level.isdigit() else 1) + " "

            if text:
                lines.append(f"[{style_name}] {indent}{text}")
            else:
                lines.append(f"[{style_name}] (empty)")
            para_idx += 1

        elif tag == qn("w:tbl"):
            tbl = tables_by_element.get(child)
            if tbl is None:
                continue
            table_idx += 1
            lines.append("")
            lines.append(f"--- TABLE {table_idx} ---")
            for ri, row in enumerate(tbl.rows):
                cells = [cell.text.strip().replace("\n", " | ") for cell in row.cells]
                lines.append(f"  Row {ri}: {' || '.join(cells)}")
            lines.append(f"--- END TABLE {table_idx} ---")
            lines.append("")

    # Also extract headers/footers from sections
    for si, section in enumerate(doc.sections):
        for hf_name, hf_attr in [
            ("header", "header"),
            ("footer", "footer"),
            ("first_page_header", "first_page_header"),
            ("first_page_footer", "first_page_footer"),
        ]:
            try:
                hf = getattr(section, hf_attr)
                if hf and hf.paragraphs:
                    texts = [p.text.strip() for p in hf.paragraphs if p.text.strip()]
                    if texts:
                        lines.append(f"\n[Section {si} {hf_name}]")
                        for t in texts:
                            lines.append(f"  {t}")
            except Exception:
                pass

    return "\n".join(lines)


def main():
    handmade = "tests/input/Internal_VAPT_Report_AEG_Vision.docx"
    generated = "outputs/[Client_Name]_20260416_001052_VAPT_Report.docx"

    tmpdir = os.environ.get("TMPDIR", "/tmp")

    for label, path in [("HANDMADE", handmade), ("GENERATED", generated)]:
        if not os.path.isfile(path):
            print(f"ERROR: {path} not found", file=sys.stderr)
            continue

        print(f"Extracting {label}: {path} ...")
        content = extract_docx_content(path)

        outfile = os.path.join(tmpdir, f"docx_extract_{label.lower()}.txt")
        with open(outfile, "w") as f:
            f.write(f"=== {label} REPORT ===\n\n")
            f.write(content)
            f.write(f"\n\n=== END {label} REPORT ===\n")
        print(f"  Saved to {outfile} ({len(content)} chars)")

        # Print first 300 lines to terminal as preview
        preview_lines = content.split("\n")
        print(f"\n{'='*80}")
        print(f"=== {label} REPORT (first 300 lines) ===")
        print(f"{'='*80}")
        for line in preview_lines[:300]:
            print(line)
        if len(preview_lines) > 300:
            print(f"... ({len(preview_lines) - 300} more lines, see full file)")
        print()

    print(f"\nFull extracts saved in {tmpdir}/docx_extract_*.txt")
    print("To view: cat $TMPDIR/docx_extract_handmade.txt")
    print("         cat $TMPDIR/docx_extract_generated.txt")


if __name__ == "__main__":
    main()
