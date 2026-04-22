"""
generate_report.py — Core Engine for VAPT Report Generator
===========================================================
Renders using python-docx directly on the user's actual template, preserving
all original branding, styles, images, and layout.
"""

import json
import os
import re
import sys
import copy
import logging
import unicodedata
from inspect import Parameter, signature
from datetime import datetime
from io import BytesIO
from pathlib import Path

import pandas as pd

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from report_core.constants import (
    SYSTEM_PROMPT,
    REPORT_SCHEMA_REQUIRED,
    FINDING_REQUIRED_FIELDS,
    SEVERITY_COLORS,
)
from report_core.json_schema import (
    _extract_json_str,
    _repair_json,
    _strip_think_blocks,
    safe_parse_json,
    validate_json_schema as _validate_json_schema_impl,
)
from report_core.input_processing import (
    _extract_text_from_pdf,
    _extract_text_from_docx,
    preprocess_scan,
    read_scan_input,
    chunk_scan_text,
)
from report_core.privacy import (
    ClientDataInternetEgressError,
    assert_clean_client_context_for_cloud,
    is_cloud_provider,
    prepare_text_for_cloud_egress,
    sanitize_finding,
    restore_finding,
)
from report_tool.prompts import (
    FINDINGS_CHUNK_PROMPT,
    METADATA_PROMPT,
    PER_VULN_DETAIL_SYSTEM,
    PER_VULN_DETAIL_USER_TEMPLATE,
    PER_VULN_INITIAL_PROMPT,
)

CONFIG_FILE = "config.json"
# Module-level cache for config content keyed by resolved path.
_CONFIG_CACHE: dict = {}
logger = logging.getLogger(__name__)


def _callback_positional_arity(callback) -> int | None:
    """Return supported positional args, or None when callback accepts *args."""
    try:
        params = signature(callback).parameters.values()
    except (TypeError, ValueError):
        return 4
    if any(param.kind == Parameter.VAR_POSITIONAL for param in params):
        return None
    return sum(
        1
        for param in params
        if param.kind in (Parameter.POSITIONAL_ONLY, Parameter.POSITIONAL_OR_KEYWORD)
    )


def _emit_progress_callback(
    progress_callback, stage, current, total, message, detail=None
) -> None:
    if not progress_callback:
        return
    arity = _callback_positional_arity(progress_callback)
    payload = [stage, current, total, message]
    if detail is not None and (arity is None or arity >= 5):
        payload.append(detail)
    progress_callback(*payload[:arity] if arity is not None else payload)

# ── Severity inference helpers ────────────────────────────────────────────────

# CVSS v3 ranges → severity (FIRST_GE order)
_CVSS_RANGES = [
    (9.0, "Critical"),
    (7.0, "High"),
    (4.0, "Medium"),
    (0.1, "Low"),
]

# Keywords → severity (checked against name + description, first match wins)
_SEVERITY_KEYWORDS = {
    "Critical": [
        "remote code execution",
        "unauthenticated rce",
        "rce",
        "arbitrary code execution",
        "zero-day",
        "0-day",
        "critical rce",
        "pre-auth rce",
    ],
    "High": [
        "sql injection",
        "sqli",
        "command injection",
        "os command",
        "privilege escalation",
        "authentication bypass",
        "auth bypass",
        "xxe",
        "xml external entity",
        "insecure deserialization",
        "ldap injection",
        "xpath injection",
        "heap overflow",
        "buffer overflow",
        "broken authentication",
        "account takeover",
    ],
    "Medium": [
        "cross-site scripting",
        "xss",
        "reflected xss",
        "stored xss",
        "csrf",
        "cross-site request forgery",
        "ssrf",
        "open redirect",
        "directory traversal",
        "path traversal",
        "lfi",
        "rfi",
        "sensitive data exposure",
        "idor",
        "broken access control",
        "insecure direct object",
        "session fixation",
        "clickjacking",
    ],
    "Low": [
        "missing security header",
        "missing x-frame",
        "missing csp",
        "missing hsts",
        "cookie without httponly",
        "cookie without secure",
        "banner grabbing",
        "version disclosure",
        "server version",
        "software version exposed",
        "weak cipher",
        "weak tls",
        "self-signed certificate",
        "expired certificate",
    ],
}

# Vuln-type → control fields (keyed by substring that appears in name/description)
# Keys are matched case-insensitively against finding name + description.
# These are the ONLY source of deterministic control fields — no generic fallback is applied.
from report_core.finding_helpers import (
    _CONTROL_MAP,
    _PLACEHOLDER_TEXT,
    _REC_TIER_LABELS,
    _is_placeholder_text,
    _dedupe_findings,
    _try_parse_python_list_repr,
    _format_recommendation_cell,
    _format_reference_cell,
    _is_cwe_or_cve_ref,
    _normalize_report_text,
    _looks_noisy_line,
    _looks_like_evidence_text,
    _derive_audit_requirement,
    _prepare_audit_requirement,
    _prepare_proof_of_concept,
    infer_severity_from_cvss,
    infer_severity_from_keywords,
    _normalize_lookup_field_value,
    _is_empty_lookup_value,
    _structured_field_is_generic,
    fill_missing_fields,
    _derive_category_from_finding,
    _derive_business_impact_from_finding,
    _structured_field_is_incomplete,
    _structured_finding_needs_llm_lookup,
    _structured_value_is_less_specific,
    _structured_should_preserve_prior_value,
    _merge_structured_lookup_result,
    _GENERIC_CONTROL_OBJECTIVE_FALLBACK,
    _GENERIC_CONTROL_NAME_FALLBACK,
    _GENERIC_AUDIT_REQUIREMENT_FALLBACK,
    _GENERIC_REFERENCE_FALLBACK,
)


_STRUCTURED_REQUIRED_SOURCE_FIELDS = (
    "description",
    "remediation",
    "business_impact",
    "control_objective",
    "control_name",
    "audit_requirement",
    "reference",
    "category",
)


def _has_structured_source_value(value) -> bool:
    if isinstance(value, str):
        return bool(value.strip()) and not _is_placeholder_text(value)
    if isinstance(value, (list, tuple, set)):
        return bool(value)
    if isinstance(value, dict):
        return bool(value)
    return value not in (None, "")


def _validate_structured_source_fields(findings: list[dict]) -> None:
    missing: list[str] = []
    for idx, finding in enumerate(findings, start=1):
        finding_id = str(finding.get("id") or f"finding#{idx}")
        missing_fields = [
            field
            for field in _STRUCTURED_REQUIRED_SOURCE_FIELDS
            if not _has_structured_source_value(finding.get(field))
        ]
        if missing_fields:
            missing.append(f"{finding_id}: {', '.join(missing_fields)}")

    if missing:
        raise ValueError(
            "Structured findings missing required source-of-truth fields. "
            "Cloud detail lookup or structured lookup did not complete successfully. "
            "Missing fields for: " + "; ".join(missing)
        )


def load_config(path: str = None) -> dict:
    """Load config.json, searching several locations if path not given."""
    if not path:
        for p in ["config.json", Path(__file__).parent / CONFIG_FILE]:
            if Path(p).exists():
                path = str(p)
                break

    if not path:
        raise FileNotFoundError("config.json not found in default locations.")

    config_path = Path(path)
    if not config_path.exists():
        raise FileNotFoundError(f"config.json not found at '{path}'.")

    # Use resolved absolute path as cache key and reload only if file changed.
    cache_key = str(config_path.resolve())
    mtime = config_path.stat().st_mtime
    cached = _CONFIG_CACHE.get(cache_key)
    if cached and cached.get("mtime") == mtime:
        return copy.deepcopy(cached["config"])

    # Load fresh config from disk and update cache
    with open(config_path, "r", encoding="utf-8") as f:
        cfg = json.load(f)

    _CONFIG_CACHE[cache_key] = {"mtime": mtime, "config": copy.deepcopy(cfg)}
    logger.debug("Loaded config from: %s", config_path.absolute())
    return copy.deepcopy(cfg)


def _capped(config: dict, max_tokens: int) -> dict:
    """Return config copy with max_tokens capped to max_tokens."""
    c = config.copy()
    c["llm"] = {
        **config["llm"],
        "max_tokens": min(config["llm"].get("max_tokens", 8192), max_tokens),
    }
    return c


from report_tool.llm import (
    fetch_lm_studio_models,
    fetch_openrouter_models,
    CancelledError,
    TokenBudgetExceeded,
    _build_payload,
    call_llm,
    _call_llm_openai,
    _call_llm_generic,
    reset_token_budget,
)
from report_tool.run_logging import append_run_log_event, ensure_run_log


# ── LLM call ──────────────────────────────────────────────────────────────────


# ── Chunked generate pipeline ───────────────────────────────────────────────


def generate_chunked(
    scan_filepath: str,
    config_path: str = CONFIG_FILE,
    cancel_event=None,
    progress_callback=None,
    findings_per_chunk: int = None,
) -> tuple:
    """
    Chunked pipeline: processes large scans by splitting into manageable pieces.

    1. Read & preprocess scan text (no truncation — we handle size via chunks)
    2. Call LLM once for metadata (client name, exec summary, dates, etc.)
    3. Split scan text into chunks of ~findings_per_chunk vulns each
    4. Call LLM once per chunk to extract findings
    5. Merge all findings, validate, split false positives

    progress_callback(stage, current, total, message[, detail]) is called at
    each step so the UI can show meaningful progress.

    Returns (data_dict, raw_llm_texts_list, false_positive_list).
    """
    config = load_config(config_path)
    ensure_run_log(config, pipeline="chunked")
    if findings_per_chunk is None:
        findings_per_chunk = config.get("limits", {}).get("findings_per_chunk", 6)

    def _progress(stage, current, total, message, detail=None):
        _emit_progress_callback(
            progress_callback, stage, current, total, message, detail
        )

    # ── Step 1: Read scan (don't truncate — chunking handles length) ──
    _progress("read", 0, 1, "Reading scan file…")

    ext = Path(scan_filepath).suffix.lower()
    if ext == ".pdf":
        content = _extract_text_from_pdf(scan_filepath)
    elif ext in (".docx", ".doc"):
        content = _extract_text_from_docx(scan_filepath)
    else:
        with open(scan_filepath, "r", encoding="utf-8", errors="replace") as f:
            content = f.read()

    scan_text = preprocess_scan(content)
    word_count = len(scan_text.split())
    print(f"ℹ️  Preprocessed scan: {word_count} words", file=sys.stderr)

    # ── Decide: single-shot or chunked ──
    # Cloud providers (OpenRouter) have large context windows — always single-shot.
    # Chunking was designed for local models with small context limits.
    max_words = config["limits"]["max_input_words"]
    if is_cloud_provider(config) or word_count <= max_words:
        _progress("single", 0, 1, "Small scan — using single-shot analysis…")
        raw_text = call_llm(scan_text, config, cancel_event=cancel_event)
        data = safe_parse_json(raw_text)
        data = validate_json_schema(data)
        data, fps = split_false_positives(data)
        data = _apply_local_normalization_pass(data, config)
        data = _apply_reference_validation_pass(data, config)
        data = _apply_quality_gate(data, config)
        data = compose_report_narrative(data)
        _progress("done", 1, 1, "Complete")
        return data, [raw_text], fps

    # ── Step 2: Extract metadata ──
    _progress("metadata", 0, 1, "Extracting report metadata…")

    # Send a truncated version for metadata (just need headers/intro)
    meta_text = " ".join(scan_text.split()[:max_words])
    meta_raw = _call_llm_generic(
        METADATA_PROMPT,
        "Extract report metadata from this VAPT scan data:\n\n" + meta_text,
        _capped(config, 512),
        cancel_event=cancel_event,
        log_label="metadata",
        task_type="metadata",
    )
    metadata = safe_parse_json(meta_raw)
    _progress("metadata", 1, 1, "Metadata extracted")

    # ── Step 3: Split into chunks ──
    chunks = chunk_scan_text(scan_text, findings_per_chunk=findings_per_chunk)
    total_chunks = len(chunks)
    print(
        f"ℹ️  Split into {total_chunks} chunk(s) for finding extraction", file=sys.stderr
    )

    # ── Step 4: Extract findings per chunk ──
    all_findings = []
    all_raw_texts = [meta_raw]

    for i, chunk_text in enumerate(chunks):
        if cancel_event and cancel_event.is_set():
            raise CancelledError("Cancelled by user.")

        _progress(
            "findings",
            i,
            total_chunks,
            f"Extracting findings from chunk {i + 1}/{total_chunks}…",
        )

        chunk_raw = _call_llm_generic(
            FINDINGS_CHUNK_PROMPT,
            f"Extract vulnerability findings from this section of VAPT scan data "
            f"(chunk {i + 1} of {total_chunks}):\n\n" + chunk_text,
            _capped(config, 4096),
            cancel_event=cancel_event,
            log_label=f"findings_chunk_{i + 1}",
            task_type="chunk_extraction",
        )
        all_raw_texts.append(chunk_raw)

        try:
            chunk_data = safe_parse_json(chunk_raw)
            chunk_findings = chunk_data.get("findings", [])
            if not isinstance(chunk_findings, list):
                chunk_findings = []

            # Collect findings from this chunk; defer dedupe/renumbering until
            # after all chunks are processed so duplicates across chunks can
            # be consolidated.
            all_findings.extend(chunk_findings)
            print(
                f"  ✓ Chunk {i + 1}: extracted {len(chunk_findings)} finding(s)",
                file=sys.stderr,
            )
        except Exception as e:
            print(f"  ⚠️  Chunk {i + 1}: failed to parse findings: {e}", file=sys.stderr)

    _progress(
        "findings",
        total_chunks,
        total_chunks,
        f"Extracted {len(all_findings)} total findings",
    )

    # ── Step 5: Merge & validate ──
    _progress("merge", 0, 1, "Merging and validating results…")

    # Build final data dict from metadata + all findings
    # Deduplicate merged findings (LLM may repeat similar entries across chunks)
    deduped = dedupe_findings(all_findings)
    for idx, f in enumerate(deduped, start=1):
        f["id"] = f"VAPT-{idx:03d}"
    data = metadata
    data["findings"] = deduped
    data = validate_json_schema(data)
    data, fps = split_false_positives(data)
    data = _apply_local_normalization_pass(data, config)
    data = _apply_reference_validation_pass(data, config)
    data = _apply_quality_gate(data, config)
    data = compose_report_narrative(data)

    _progress(
        "done",
        1,
        1,
        f"Complete — {data['total_findings']} findings, {len(fps)} false positives",
    )

    return data, all_raw_texts, fps


def _build_per_vuln_progress_message(
    completed: int, total: int, candidate: dict
) -> str:
    """Return a user-facing progress message including the active vulnerability."""
    short_name = str(
        (candidate or {}).get("short_name")
        or (candidate or {}).get("name")
        or "[Unnamed Vulnerability]"
    ).strip()
    vuln_id = str((candidate or {}).get("vuln_id") or "?").strip()
    if short_name and len(short_name) > 90:
        short_name = short_name[:87].rstrip() + "..."
    return f"Processing vulnerability {completed}/{total}: [{vuln_id}] {short_name}"


def _resolve_structured_cloud_lookup_cfg(config: dict) -> dict | None:
    """Resolve cloud detail-lookup settings from runtime config (session-only API key, no env fallback)."""
    if not is_cloud_provider(config):
        return None
    llm_cfg = (config or {}).get("llm", {}) if isinstance(config, dict) else {}
    task_models = llm_cfg.get("task_models") or {}
    # API key must come from session/request payload only — no environment variable fallback
    api_key = str(llm_cfg.get("api_key") or "").strip()
    if not api_key:
        return None
    model_id = (
        task_models.get("lookup")
        or llm_cfg.get("lookup_model")
        or llm_cfg.get("model")
        or ""
    ).strip()
    if not model_id:
        raise ValueError(
            "OpenRouter model id missing. Choose a model in the UI or set llm.model in config.json."
        )
    return {
        "enabled": True,
        "api_key": api_key,
        "model": model_id,
    }


def _require_session_api_key_for_cloud(config: dict) -> None:
    if not is_cloud_provider(config):
        return
    llm_cfg = (config or {}).get("llm", {}) if isinstance(config, dict) else {}
    api_key = str(llm_cfg.get("api_key") or "").strip()
    if api_key:
        return
    raise ValueError(
        "OpenRouter API key missing from current session. Provide it in the UI before running analysis."
    )


def _raise_if_client_data_egress_blocked(exc: Exception) -> None:
    if isinstance(exc, ClientDataInternetEgressError):
        raise exc


def _build_structured_lookup_user_prompt(
    lookup_context: str,
    finding_payload_json: str,
) -> str:
    return (
        "Process this vulnerability finding and return the result JSON object.\n\n"
        "CRITICAL REQUIREMENTS:\n"
        "- 'remediation': EXACTLY 3 lines separated by newlines (NOT a JSON array, NOT Python list). "
        "Line 1: specific primary fix. Line 2: compensating control if primary infeasible (start with 'If...'). "
        "Line 3: detection/monitoring only (start with 'Isolate...', 'Monitor...', or 'Enable...').\n"
        "- 'business_impact': 1-2 sentences explaining the concrete attacker outcome and business impact for THIS finding.\n"
        "- 'control_objective': Specific to THIS vulnerability type. Never use the generic outdated-version "
        "objective for non-EOL findings such as cipher issues, SNMP, RDP, SMB, or service config problems.\n"
        "- 'control_name': 2-5 word noun phrase specific to the vuln category.\n"
        "- 'audit_requirement': How an auditor TESTS the control (tool + what to check), not how to fix it.\n"
        "- 'reference': 2-4 authoritative references only. Prefer NVD, MITRE CWE, OWASP, and vendor advisories. Avoid generic placeholder references.\n"
        "- TREAT GENERIC DEFAULTS AS INCOMPLETE: Replace ALL generic defaults with vulnerability-specific content.\n"
        "- BANNED outputs: generic boilerplate, Python list syntax in string fields, JSON arrays for remediation.\n\n"
        "IMPORTANT:\n"
        "- Focus on completing missing narrative fields. Do NOT waste tokens repeating large input fields verbatim.\n"
        "- If the payload contains asset summary fields, use them as context and do not attempt to expand the full asset inventory.\n\n"
        "Client Context:\n"
        + lookup_context
        + "\n\nPayload:\n"
        + finding_payload_json
    )


def _should_run_structured_second_stage(
    config: dict,
    cloud_lookup_stats: dict,
) -> bool:
    structured_cfg = (config or {}).get("structured_lookup") or {}
    explicit = structured_cfg.get("second_stage_enabled")
    if explicit is not None:
        return bool(explicit)
    cloud_eligible = int((cloud_lookup_stats or {}).get("cloud_eligible_findings") or 0)
    if is_cloud_provider(config) and cloud_eligible > 0:
        return False
    return True


def _build_structured_lookup_prompt_finding(finding: dict) -> dict:
    prompt_finding = copy.deepcopy(finding or {})
    assets_raw = _normalize_affected_assets(
        prompt_finding.get("affected_assets_raw") or prompt_finding.get("affected_assets")
    )
    if assets_raw:
        prompt_finding["affected_assets_count"] = len(assets_raw)
        prompt_finding["affected_assets_sample"] = assets_raw[:10]
        prompt_finding["affected_assets"] = _truncate_assets(assets_raw, max_items=10)
    return prompt_finding


def _preflight_structured_cloud_prompts(
    config: dict,
    system_prompt: str,
    lookup_context: str,
) -> None:
    if not is_cloud_provider(config):
        return
    prepare_text_for_cloud_egress(
        system_prompt,
        config,
        "structured lookup system prompt",
    )
    prepare_text_for_cloud_egress(
        _build_structured_lookup_user_prompt(
            lookup_context,
            "{}",
        ),
        config,
        "structured lookup prompt preflight",
    )


def generate_per_vuln(
    scan_filepath,
    config_path=CONFIG_FILE,
    client_context="",
    cancel_event=None,
    progress_callback=None,
    api_key: str = "",
):
    """
    New per-vulnerability pipeline.

    Steps:
    1. Read & preprocess the scan text.
    2. Extract metadata via METADATA_PROMPT.
    3. Ask the LLM to produce a compact list of vulnerability raw blocks (vuln_id + raw_block).
    4. For each vuln, call the LLM sequentially to expand into the full finding JSON.
    5. Merge results, validate schema, and split false positives.

    Returns (data, raw_texts_list, fps).
    """
    config = load_config(config_path)
    if api_key:
        config["llm"]["api_key"] = api_key
    _require_session_api_key_for_cloud(config)
    ensure_run_log(config, pipeline="per_vuln")
    reset_token_budget(config.get("limits", {}).get("max_report_tokens", 0))

    def _progress(stage, current, total, message, detail=None):
        _emit_progress_callback(
            progress_callback, stage, current, total, message, detail
        )

    # Read & preprocess input — support single path or list of paths
    _progress("read", 0, 1, "Reading scan file…")

    # Multi-file support: if list, check if all are structured
    if isinstance(scan_filepath, (list, tuple)):
        exts = {Path(fp).suffix.lower() for fp in scan_filepath}
        if exts <= {".csv", ".xlsx", ".xls"}:
            return _generate_from_structured_file(
                scan_filepath, config, client_context, cancel_event, progress_callback
            )
        # Mixed: process structured files together, text files separately and merge
        structured = [
            fp
            for fp in scan_filepath
            if Path(fp).suffix.lower() in (".csv", ".xlsx", ".xls")
        ]
        text_files = [fp for fp in scan_filepath if fp not in structured]
        if structured and not text_files:
            return _generate_from_structured_file(
                structured, config, client_context, cancel_event, progress_callback
            )
        # For now, use first file if mixed
        scan_filepath = text_files[0] if text_files else structured[0]

    ext = Path(scan_filepath).suffix.lower()

    # Delegate to structured parser if applicable
    if ext in (".csv", ".xlsx", ".xls"):
        return _generate_from_structured_file(
            scan_filepath, config, client_context, cancel_event, progress_callback
        )
    try:
        if ext == ".pdf":
            content = _extract_text_from_pdf(scan_filepath)
        elif ext in (".docx", ".doc"):
            content = _extract_text_from_docx(scan_filepath)
        else:
            with open(scan_filepath, "r", encoding="utf-8", errors="replace") as f:
                content = f.read()
    except Exception as e:
        raise RuntimeError(f"Failed to read scan file: {e}")

    scan_text = preprocess_scan(content)

    raw_texts_list = []

    # Step 1: Metadata
    _progress("metadata", 0, 1, "Extracting metadata…")
    try:
        meta_text = " ".join(
            scan_text.split()[
                : max(1, config.get("limits", {}).get("max_input_words", 6000))
            ]
        )
        meta_raw = _call_llm_generic(
            METADATA_PROMPT,
            "Extract report metadata from this VAPT scan data:\n\n" + meta_text,
            _capped(config, 512),
            cancel_event=cancel_event,
            log_label="metadata",
        )
        raw_texts_list.append(meta_raw)
        metadata = safe_parse_json(meta_raw)
    except CancelledError:
        raise
    except Exception as e:
        _raise_if_client_data_egress_blocked(e)
        print(f"⚠️  Metadata extraction failed: {e}", file=sys.stderr)
        metadata = {}
    _progress("metadata", 1, 1, "Metadata extracted")

    # Step 2: Initial per-vuln extraction (compact list of raw blocks)
    _progress("initial_extract", 0, 1, "Extracting per-vulnerability raw blocks…")
    try:
        initial_raw = _call_llm_generic(
            PER_VULN_INITIAL_PROMPT,
            "Extract compact vulnerability list from the following scan text:\n\n"
            + scan_text,
            _capped(config, 4096),
            cancel_event=cancel_event,
            log_label="per_vuln_initial",
            task_type="extraction",
        )
        raw_texts_list.append(initial_raw)
        try:
            initial_parsed = safe_parse_json(initial_raw)
            initial_findings = initial_parsed.get("findings", [])
            if not isinstance(initial_findings, list):
                initial_findings = []
        except Exception as e:
            print(f"⚠️  Failed to parse initial per-vuln output: {e}", file=sys.stderr)
            initial_findings = []
    except CancelledError:
        raise
    except Exception as e:
        _raise_if_client_data_egress_blocked(e)
        print(f"⚠️  Initial per-vuln LLM call failed: {e}", file=sys.stderr)
        initial_findings = []
    _progress(
        "initial_extract", 1, 1, f"Found {len(initial_findings)} candidate vuln blocks"
    )

    # Normalize and sort by numeric vuln_id
    normalized_candidates = []
    for item in initial_findings:
        try:
            vuln_id = item.get("vuln_id", None)
            if isinstance(vuln_id, str) and vuln_id.isdigit():
                vuln_id = int(vuln_id)
            elif isinstance(vuln_id, (int, float)):
                vuln_id = int(vuln_id)
            else:
                raw = str(item.get("vuln_id") or item.get("id") or "")
                digits = re.findall(r"\d+", raw)
                vuln_id = int(digits[0]) if digits else None
            short_name = str(
                item.get("short_name") or item.get("name") or "[INSUFFICIENT DATA]"
            )
            raw_block = str(item.get("raw_block") or "")
            if vuln_id is None:
                vuln_id = 10**6 + len(normalized_candidates)
            normalized_candidates.append(
                {"vuln_id": vuln_id, "short_name": short_name, "raw_block": raw_block}
            )
        except Exception as e:
            print(f"  ⚠️  Skipping malformed initial candidate: {e}", file=sys.stderr)

    normalized_candidates = sorted(
        normalized_candidates, key=lambda x: int(x["vuln_id"])
    )
    append_run_log_event(
        config,
        "initial_vulnerability_candidates",
        {
            "source_files": [str(scan_filepath)],
            "candidate_count": len(normalized_candidates),
            "candidates": normalized_candidates,
        },
    )

    # Step 3: Per-vuln detailed extraction
    per_vuln_results = []
    raw_texts_list = []
    total = len(normalized_candidates)

    # Partial-save path for crash recovery
    partial_dir = Path(config["paths"]["log_dir"])
    partial_dir.mkdir(parents=True, exist_ok=True)
    partial_path = partial_dir / "_partial_per_vuln_results.json"

    # Pre-flight connectivity check for local providers — fail fast before
    # spending time discovering a downed server one vuln at a time.
    _llm_cfg = config.get("llm", {})
    if _llm_cfg.get("provider", "local") == "local":
        _base_url = _llm_cfg.get("base_url", "http://127.0.0.1:1234/v1")
        try:
            from report_tool.llm import probe_local_endpoint

            probe_local_endpoint(_base_url)
        except RuntimeError as _probe_err:
            raise RuntimeError(str(_probe_err)) from None

    _progress("per_vuln", 0, total, f"Processing {total} vulnerabilities…")

    # Use a lower token cap for individual finding expansion — 2048 is ample for one finding.
    per_vuln_config = config.copy()
    per_vuln_config["llm"] = {**config["llm"], "max_tokens": 1500}

    _consecutive_llm_failures = 0
    _MAX_CONSECUTIVE_FAILURES = 3

    for completed, (idx, cand) in enumerate(
        enumerate(normalized_candidates, start=1), start=1
    ):
        if cancel_event and cancel_event.is_set():
            break
        vuln_id = int(cand["vuln_id"])
        _progress(
            "per_vuln",
            max(0, completed - 1),
            total,
            _build_per_vuln_progress_message(completed, total, cand),
        )
        label = f"per_vuln_{vuln_id}"
        user_payload = PER_VULN_DETAIL_USER_TEMPLATE.format(
            raw_block=cand.get("raw_block", ""),
            vuln_id=vuln_id,
        )
        try:
            raw_detail = _call_llm_generic(
                PER_VULN_DETAIL_SYSTEM,
                user_payload,
                per_vuln_config,
                cancel_event=cancel_event,
                log_label=label,
                task_type="per_vuln_extraction",
            )
            parsed = safe_parse_json(raw_detail)
            if not isinstance(parsed, dict):
                print(
                    f"  ⚠️  {label}: expected an object, got {type(parsed)}",
                    file=sys.stderr,
                )
            else:
                try:
                    parsed_vuln_id = parsed.get("vuln_id", None)
                    if isinstance(parsed_vuln_id, str) and parsed_vuln_id.isdigit():
                        parsed_vuln_id = int(parsed_vuln_id)
                    elif isinstance(parsed_vuln_id, (int, float)):
                        parsed_vuln_id = int(parsed_vuln_id)
                    else:
                        parsed_vuln_id = vuln_id
                    parsed["vuln_id"] = int(parsed_vuln_id)
                except Exception:
                    parsed["vuln_id"] = vuln_id
                fill_missing_fields(parsed)
                parsed["audit_requirement"] = _prepare_audit_requirement(parsed)
                per_vuln_results.append(parsed)
                _consecutive_llm_failures = 0  # reset on success
                # Save partial results for crash recovery
                try:
                    with open(partial_path, "w", encoding="utf-8") as _pf:
                        json.dump(per_vuln_results, _pf, indent=2)
                except Exception:
                    pass
            if raw_detail:
                raw_texts_list.append(raw_detail)
        except Exception as e:
            _raise_if_client_data_egress_blocked(e)
            print(f"  ⚠️  LLM error for vuln {vuln_id}: {e}", file=sys.stderr)
            _consecutive_llm_failures += 1
            if _consecutive_llm_failures >= _MAX_CONSECUTIVE_FAILURES:
                _lm_url = per_vuln_config.get("llm", {}).get(
                    "base_url", "the configured base_url"
                )
                raise RuntimeError(
                    f"Pipeline aborted after {_consecutive_llm_failures} consecutive LLM "
                    f"failures. Last error: {e}. "
                    f"Check that your model server is running and reachable at {_lm_url}."
                )

        _progress(
            "per_vuln",
            completed,
            total,
            _build_per_vuln_progress_message(completed, total, cand),
        )

    # Clean up partial file on successful completion
    partial_path.unlink(missing_ok=True)

    # Step 4: Merge metadata + per-vuln findings, validate and split FPs
    _progress("merge", 0, 1, "Merging findings and validating schema…")
    data = metadata if isinstance(metadata, dict) else {}
    # Deduplicate per-vuln results (LLM may repeat similar content) and
    # renumber IDs sequentially.
    deduped = dedupe_findings(per_vuln_results)
    for idx, f in enumerate(deduped, start=1):
        f["id"] = f"VAPT-{idx:03d}"
    data["findings"] = deduped
    data = validate_json_schema(data)
    data, fps = split_false_positives(data)
    data = _apply_local_normalization_pass(data, config)
    data = _apply_reference_validation_pass(data, config)
    data = _apply_quality_gate(data, config)
    data = compose_report_narrative(data)
    _progress(
        "done",
        1,
        1,
        f"Complete — {data.get('total_findings',0)} findings, {len(fps)} false positives",
    )

    return data, raw_texts_list, fps


# ── JSON extraction & validation ──────────────────────────────────────────────


def validate_json_schema(data: dict) -> dict:
    return _validate_json_schema_impl(
        data,
        report_schema_required=REPORT_SCHEMA_REQUIRED,
        finding_required_fields=FINDING_REQUIRED_FIELDS,
        severity_colors=SEVERITY_COLORS,
        infer_severity_from_cvss=infer_severity_from_cvss,
        infer_severity_from_keywords=infer_severity_from_keywords,
        fill_missing_fields=fill_missing_fields,
        prepare_audit_requirement=_prepare_audit_requirement,
        prepare_proof_of_concept=_prepare_proof_of_concept,
    )


# ── False-positive detection ──────────────────────────────────────────────────


def _fp_similarity(a: str, b: str) -> float:
    if not a and not b:
        return 1.0
    if not a or not b:
        return 0.0
    matches = sum(ca == cb for ca, cb in zip(a, b))
    return 2.0 * matches / (len(a) + len(b))


_FP_CANONICAL_TOKENS = [
    "false positive",
    "falsepositive",
    "false pos",
    "fals positiv",
    "flase positive",
    "fasle positive",
    "false postive",
    "false positiv",
    "fals positive",
    "fp - ",
    "(fp)",
    "[fp]",
    "- fp",
    "not exploitable",
    "not a vuln",
    "not a vulnerability",
    "informational only",
    "deemed informational",
]
_FP_FUZZY_THRESHOLD = 0.72
_FP_TARGET = "false positive"
_FP_REVIEW_MARKERS = ("check", "verify", "review", "pending", "possible")


def _status_text_is_fp_review(text: str) -> bool:
    combined = _normalize_text_for_key(text)
    if "false pos" not in combined and "falsepositive" not in combined:
        return False
    return any(marker in combined for marker in _FP_REVIEW_MARKERS)


def is_false_positive(finding: dict) -> bool:
    name = str(finding.get("name", ""))
    obs = str(finding.get("observation", ""))
    risk = str(finding.get("risk_status", ""))
    desc = str(finding.get("description", ""))[:200]
    combined = " ".join([name, obs, risk, desc]).lower()
    review_only = _status_text_is_fp_review(combined)

    # Keyword-based check
    if not review_only:
        for token in _FP_CANONICAL_TOKENS:
            if token in combined:
                return True

    # Fuzzy name match against "false positive"
    if not review_only:
        name_norm = re.sub(r"[^a-z ]", " ", name.lower())
        win_len = len(_FP_TARGET)
        for i in range(max(1, len(name_norm) - win_len + 1)):
            if (
                _fp_similarity(name_norm[i : i + win_len], _FP_TARGET)
                >= _FP_FUZZY_THRESHOLD
            ):
                return True

    # Accepted risk / risk-accepted status
    risk_lower = risk.strip().lower()
    if risk_lower in (
        "accepted",
        "risk accepted",
        "accept",
        "accepted risk",
        "closed - accepted",
    ):
        return True

    # CVSS score of 0 = informational / non-finding
    cvss_str = str(finding.get("cvss", "")).strip()
    if cvss_str in ("0", "0.0", "0.00"):
        return True

    # Observation explicitly "Informational"
    if obs.strip().lower() == "informational":
        return True

    # Severity "Informational" = not an actionable vulnerability
    sev = str(finding.get("severity", "")).strip().lower()
    if sev == "informational":
        return True

    return False


def split_false_positives(data: dict) -> tuple:
    real_findings, fp_findings = [], []
    for finding in data.get("findings", []):
        (fp_findings if is_false_positive(finding) else real_findings).append(finding)
    data["findings"] = real_findings
    counts = {"Critical": 0, "High": 0, "Medium": 0, "Low": 0}
    for f in real_findings:
        sev = f.get("severity", "")
        if sev in counts:
            counts[sev] += 1
    data["total_critical"] = counts["Critical"]
    data["total_high"] = counts["High"]
    data["total_medium"] = counts["Medium"]
    data["total_low"] = counts["Low"]
    data["total_findings"] = sum(counts.values())
    return data, fp_findings


def _normalize_text_for_key(s: str) -> str:
    s = str(s or "").lower()
    s = re.sub(r"\s+", " ", s).strip()
    return s


_STRUCTURED_SEVERITY_ORDER = {
    "Informational": 0,
    "Low": 1,
    "Medium": 2,
    "High": 3,
    "Critical": 4,
}

_STRUCTURED_SEVERITY_OVERRIDES = {
    "TLS Version 1.1 Deprecated Protocol": "Medium",
    "SNMP Agent Default Community Name (public)": "High",
    "Ncache Server/Service Exposed": "High",
    "nginx < 1.17.7 Information Disclosure": "Medium",
}


def _structured_cell_text(value) -> str:
    try:
        if pd.isna(value):
            return ""
    except Exception:
        pass
    text = str(value or "").strip()
    return "" if text.lower() in {"nan", "none"} else text


def _structured_normalize_severity(value) -> str:
    text = _structured_cell_text(value).strip().lower()
    if not text:
        return "Informational"
    mapping = {
        "critical": "Critical",
        "high": "High",
        "medium": "Medium",
        "low": "Low",
        "informational": "Informational",
        "info": "Informational",
        "none": "Informational",
    }
    return mapping.get(text, "Informational")


def _structured_status_family(*values) -> str:
    combined = " ".join(_structured_cell_text(value) for value in values)
    lowered = _normalize_text_for_key(combined)
    if not lowered:
        return "open"
    if _status_text_is_fp_review(lowered):
        return "fp_review"
    if any(
        token in lowered
        for token in (
            "confirmed false positive",
            "false positive",
            "falsepositive",
            "false pos",
            "fals positiv",
            "flase positive",
        )
    ):
        return "fp_final"
    if any(token in lowered for token in ("repeat", "repeated")):
        return "repeat"
    if any(token in lowered for token in ("closed", "resolved", "remediated", "fixed")):
        return "closed"
    if any(token in lowered for token in ("open", "active", "confirmed", "new")):
        return "open"
    return "other"


def _structured_group_and_title(name: str) -> tuple[str, str]:
    lowered = _normalize_text_for_key(name)

    if lowered == "web server directory enumeration":
        return ("family:web-directory-enum", "Browsable Web Directories")
    if (
        lowered == "hsts missing from https server"
        or lowered == "hsts missing from https server (rfc 6797)"
    ):
        return ("name:hsts-missing", "HSTS Missing from HTTPS Server (RFC 6797)")
    if lowered.startswith("canonical ubuntu linux seol"):
        return ("name:ubuntu-1604-eol", "Using Canonical Ubuntu 16.04 LTS")
    if lowered.startswith("microsoft message queuing rce"):
        return ("name:msmq-rce", "Microsoft Message Queueing RCE")
    if lowered == "elasticsearch unrestricted access information disclosure":
        return (
            "name:elasticsearch-open-access",
            "Unauthorized access and Information Disclosure Vulnerability",
        )
    if "jenkins" in lowered and "<" in lowered:
        return (
            "family:jenkins-advisories",
            "Jenkins LTS <2.426.3/Jenkins weekly <2.442 Multiple Vulnerabilities",
        )
    if "kibana" in lowered:
        return (
            "family:kibana-advisories",
            "Kibana ESA-2019-01, ESA-2019-02, ESA-2019-03",
        )
    if lowered in (
        "tls version 1.0 protocol detection",
        "tls version 1.1 deprecated protocol",
        "ssl version 2 and 3 protocol detection",
        "sslv3 padding oracle on downgraded legacy encryption vulnerability (poodle)",
    ):
        return ("family:tls-legacy", "TLS Version 1.1 Deprecated Protocol")
    if lowered in (
        "ssl medium strength cipher suites supported (sweet32)",
        "ssl rc4 cipher suites supported (bar mitzvah)",
        "ssl weak cipher suites supported",
        "ssl/tls diffie-hellman modulus <= 1024 bits (logjam)",
    ):
        return (
            "family:ssl-weak-ciphers",
            "SSL Medium Strength Cipher Suites Supported (SWEET32)",
        )
    if lowered in (
        "ssh weak mac algorithms enabled",
        "ssh weak key exchange algorithms enabled",
    ):
        return (
            "family:ssh-weak-algorithms",
            "SSH Weak Key Exchange Algorithms Enabled",
        )
    if (
        lowered.startswith("openssh <")
        or lowered == "ssh terrapin prefix truncation weakness (cve-2023-48795)"
    ):
        return ("family:openssh-advisories", "OpenSSH < 9.6 Multiple Vulnerabilities")
    if "snmp" in lowered and (
        "default community" in lowered
        or "community name" in lowered
        or lowered == "snmp agent default community name (public)"
    ):
        return (
            "name:snmp-default-community",
            "SNMP Agent Default Community Name (public)",
        )
    if "snmp" in lowered and "server detection" in lowered:
        return (
            "name:snmp-default-community",
            "SNMP Agent Default Community Name (public)",
        )
    if "ncache" in lowered or (
        "cache" in lowered
        and ("web manager" in lowered or "exposed" in lowered or "service" in lowered)
    ):
        return ("name:ncache-exposed", "Ncache Server/Service Exposed")

    return (f"name:{lowered}", name)


def _strip_leading_plugin_id(text: str) -> str:
    cleaned = re.sub(r"^\s*\[\d+\]\s*", "", str(text or "")).strip()
    return cleaned or str(text or "").strip()


def _parse_version_tuple(version_text: str) -> tuple[int, ...]:
    parts = [int(part) for part in re.findall(r"\d+", str(version_text or ""))]
    return tuple(parts)


def _is_version_less_than(version_text: str, minimum_version: str) -> bool:
    current = _parse_version_tuple(version_text)
    minimum = _parse_version_tuple(minimum_version)
    if not current or not minimum:
        return False

    width = max(len(current), len(minimum))
    current = current + (0,) * (width - len(current))
    minimum = minimum + (0,) * (width - len(minimum))
    return current < minimum


def _structured_group_and_title_for_row(row: dict, name: str) -> tuple[str, str]:
    lowered = _normalize_text_for_key(name)
    if lowered == "nginx http server detection":
        plugin_output = _structured_cell_text(row.get("plugin output", ""))
        version_match = re.search(r"version\s*:\s*(\d[\d.]+)", plugin_output, re.I)
        if version_match and _is_version_less_than(version_match.group(1), "1.17.7"):
            return (
                "name:nginx-version-disclosure",
                "nginx < 1.17.7 Information Disclosure",
            )
    # SNMP detection with default community string in plugin output
    if lowered == "snmp server detection":
        plugin_output = _structured_cell_text(row.get("plugin output", ""))
        if "public" in plugin_output.lower():
            return (
                "name:snmp-default-community",
                "SNMP Agent Default Community Name (public)",
            )

    return _structured_group_and_title(name)


# Informational (Risk=None) group keys that should be promoted to real findings
_PROMOTED_INFORMATIONAL_GROUPS = {
    "family:web-directory-enum",
    "name:snmp-default-community",
    "name:ncache-exposed",
    "name:nginx-version-disclosure",
}


def _structured_keep_informational(row: dict, name: str) -> bool:
    group_key, _ = _structured_group_and_title_for_row(row, name)
    return group_key in _PROMOTED_INFORMATIONAL_GROUPS


def _structured_preferred_title_score(raw_name: str, preferred_title: str) -> int:
    raw = _normalize_text_for_key(raw_name)
    preferred = _normalize_text_for_key(preferred_title)
    return int(raw == preferred)


def _structured_is_preferred_title(finding: dict, preferred_title: str) -> bool:
    return (
        _structured_preferred_title_score(
            finding.get("_raw_name", finding.get("name", "")), preferred_title
        )
        == 1
    )


def _structured_find_score(finding: dict, preferred_title: str) -> tuple:
    cvss_value = 0.0
    try:
        match = re.search(r"(\d+(?:\.\d+)?)", str(finding.get("cvss") or ""))
        if match:
            cvss_value = float(match.group(1))
    except Exception:
        cvss_value = 0.0
    return (
        _structured_preferred_title_score(
            finding.get("_raw_name", finding.get("name", "")), preferred_title
        ),
        _STRUCTURED_SEVERITY_ORDER.get(finding.get("severity", "Medium"), 2),
        int(bool(_structured_cell_text(finding.get("cve")))),
        len(_structured_cell_text(finding.get("description")))
        + len(_structured_cell_text(finding.get("remediation"))),
        cvss_value,
    )


def _structured_merge_text(existing: str, candidate: str) -> str:
    existing_text = _structured_cell_text(existing)
    candidate_text = _structured_cell_text(candidate)
    placeholders = {"", "[INSUFFICIENT DATA]", "[PLACEHOLDER]"}
    if existing_text in placeholders:
        return candidate_text or existing_text
    if candidate_text in placeholders:
        return existing_text
    return candidate_text if len(candidate_text) > len(existing_text) else existing_text


def _is_placeholder_asset_token(value) -> bool:
    text = re.sub(r"\s+", " ", str(value or "")).strip()
    if not text:
        return True
    if _is_placeholder_text(text):
        return True

    lowered = text.lower()
    if lowered.startswith("affected assets:"):
        return _is_placeholder_text(text.split(":", 1)[1].strip())

    return False


def _normalize_affected_assets(value) -> list[str]:
    """Normalize affected_assets values into a clean, unique, ordered list."""
    if value is None:
        return []

    if isinstance(value, str):
        candidates = [value]
    elif isinstance(value, set):
        candidates = sorted(value, key=lambda item: str(item))
    elif isinstance(value, (list, tuple)):
        candidates = list(value)
    else:
        candidates = [value]

    normalized = []
    for candidate in candidates:
        if candidate is None:
            continue
        parts = (
            candidate.split(",")
            if isinstance(candidate, str)
            else str(candidate).split(",")
        )
        for part in parts:
            asset = str(part).strip()
            if _is_placeholder_asset_token(asset):
                continue
            if asset not in normalized:
                normalized.append(asset)

    return normalized


def _canonicalize_severity(value) -> str:
    text = str(value or "").strip().lower()
    if not text:
        return "Medium"
    mapping = {
        "critical": "Critical",
        "high": "High",
        "medium": "Medium",
        "moderate": "Medium",
        "low": "Low",
        "informational": "Low",
        "info": "Low",
        "none": "Low",
    }
    return mapping.get(text, str(value or "Medium").strip().title() or "Medium")


def _extract_asset_candidates_from_text(text: str) -> list[str]:
    if not text:
        return []
    candidates: list[str] = []
    for match in re.findall(r"\b(?:\d{1,3}\.){3}\d{1,3}(?::\d+)?\b", text):
        candidates.append(match)
    for match in re.findall(
        r"\b(?:[a-zA-Z0-9](?:[a-zA-Z0-9\-]{0,61}[a-zA-Z0-9])?\.)+[a-zA-Z]{2,}(?::\d+)?\b",
        text,
    ):
        candidates.append(match)
    return candidates


def _normalize_finding_with_local_model(finding: dict, config: dict) -> dict:
    """Optional local LLM normalization for title/severity/assets."""
    norm_cfg = (config or {}).get("normalization") or {}
    if not norm_cfg.get("local_llm_enabled", False):
        return finding

    llm_cfg = dict((config or {}).get("llm") or {})
    if not llm_cfg:
        return finding

    local_cfg = copy.deepcopy(config or {})
    local_cfg.setdefault("llm", {})
    local_cfg.setdefault("paths", {})
    local_cfg["paths"].setdefault("log_dir", "logs")
    local_cfg["llm"] = {
        **llm_cfg,
        "provider": "local",
        "base_url": norm_cfg.get(
            "base_url", llm_cfg.get("base_url", "http://127.0.0.1:1234/v1")
        ),
        "model": norm_cfg.get(
            "model",
            (
                (llm_cfg.get("task_models") or {}).get("normalization")
                or llm_cfg.get("model", "")
            ),
        ),
        "api_key": norm_cfg.get("api_key", llm_cfg.get("api_key", "")),
        "temperature": 0,
        "max_tokens": min(int(norm_cfg.get("max_tokens", 512)), 2048),
    }

    system_prompt = (
        "Normalize vulnerability title, severity, and affected assets. "
        "Return JSON object only with keys: name, severity, affected_assets."
    )
    user_payload = json.dumps(
        {
            "name": finding.get("name", ""),
            "severity": finding.get("severity", ""),
            "affected_assets": finding.get("affected_assets", ""),
            "description": finding.get("description", ""),
            "proof_of_concept": finding.get("proof_of_concept", ""),
        },
        ensure_ascii=False,
    )
    try:
        raw = _call_llm_generic(
            system_prompt,
            user_payload,
            local_cfg,
            task_type="normalization",
            log_label="local_normalization",
        )
        parsed = safe_parse_json(raw)
        if isinstance(parsed, dict):
            merged = dict(finding)
            for key in ("name", "severity", "affected_assets"):
                val = parsed.get(key)
                if isinstance(val, str) and val.strip():
                    merged[key] = val.strip()
            return merged
    except Exception:
        return finding
    return finding


def _apply_local_normalization_pass(data: dict, config: dict) -> dict:
    """Deterministic post-pass normalization with optional local-model refinement."""
    norm_cfg = (config or {}).get("normalization") or {}
    if norm_cfg.get("enabled", True) is False:
        return data

    findings = data.get("findings") or []
    normalized_findings: list[dict] = []

    for finding in findings:
        if not isinstance(finding, dict):
            normalized_findings.append(finding)
            continue

        normalized = dict(finding)
        normalized["name"] = (
            re.sub(r"\s+", " ", str(normalized.get("name") or "")).strip()
            or str(normalized.get("name") or "").strip()
        )
        normalized["severity"] = _canonicalize_severity(normalized.get("severity"))

        assets = []
        assets.extend(_normalize_affected_assets(normalized.get("affected_assets_raw")))
        assets.extend(_normalize_affected_assets(normalized.get("affected_assets")))
        for field in ("description", "proof_of_concept", "observation"):
            assets.extend(
                _extract_asset_candidates_from_text(str(normalized.get(field) or ""))
            )

        normalized_assets = _normalize_affected_assets(assets)
        if normalized_assets:
            normalized["affected_assets_raw"] = normalized_assets
            if _is_placeholder_asset_token(normalized.get("affected_assets")):
                normalized["affected_assets"] = ", ".join(normalized_assets)

        normalized = _normalize_finding_with_local_model(normalized, config)
        normalized_findings.append(normalized)

    out = dict(data)
    out["findings"] = normalized_findings
    return out


def _apply_quality_gate(data: dict, config: dict) -> dict:
    """Score findings; attach per-finding quality metadata; log summary.

    Non-blocking gate: annotates low-score findings so downstream review surfaces them.
    Controlled by config.quality.scorer (enabled default True).
    """
    quality_cfg = (config or {}).get("quality") or {}
    scorer_cfg = quality_cfg.get("scorer") or {}
    if scorer_cfg.get("enabled", True) is False:
        return data

    try:
        from report_tool.quality.scorer import score_finding, score_report
    except Exception:
        return data

    findings = data.get("findings") or []
    min_pass = float(scorer_cfg.get("min_pass", 0.55))
    for f in findings:
        if isinstance(f, dict):
            try:
                f["_quality"] = score_finding(f)
            except Exception:
                continue

    try:
        summary = score_report(data, min_pass=min_pass)
        failing = len(summary.get("failing") or [])
        print(
            f"ℹ️  Quality gate: count={summary.get('count')} "
            f"mean={summary.get('mean')} pass_rate={summary.get('pass_rate')} "
            f"failing<{min_pass}={failing}",
            file=sys.stderr,
        )
        data["_quality_summary"] = {
            "count": summary.get("count"),
            "mean": summary.get("mean"),
            "pass_rate": summary.get("pass_rate"),
            "failing": failing,
            "min_pass": min_pass,
        }
    except Exception as exc:
        print(f"⚠️  Quality gate skipped: {exc}", file=sys.stderr)
    return data


def _apply_reference_validation_pass(data: dict, config: dict) -> dict:
    """Validate and canonicalize finding references with deterministic URL checks."""
    quality_cfg = (config or {}).get("quality") or {}
    ref_cfg = quality_cfg.get("reference_validation") or {}
    if ref_cfg.get("enabled", True) is False:
        return data

    try:
        from report_tool.quality.references import validate_report_refs
    except Exception:
        return data

    try:
        return validate_report_refs(
            data,
            add_canonical=ref_cfg.get("add_canonical_refs", True),
            max_workers=int(ref_cfg.get("max_workers", 8)),
            timeout=float(ref_cfg.get("timeout", 8.0)),
        )
    except Exception as exc:
        _raise_if_client_data_egress_blocked(exc)
        print(f"⚠️  Reference validation skipped: {exc}", file=sys.stderr)
        return data


def dedupe_findings(findings: list) -> list:
    """De-duplicate findings while preserving a stable order.

    Strategy:
        - Prefer grouping by numeric `vuln_id` plus meaningful status family when
            available and pick the most complete record per group (heuristic: count
            of non-placeholder fields
      and description length).
    - For items without a usable `vuln_id`, fall back to a fingerprint of
      (name, affected_assets, description prefix).
    - Return items in the original input order (first-seen of chosen records).
    """

    def _to_int_vid(v):
        try:
            if v is None:
                return None
            if isinstance(v, (int, float)):
                return int(v)
            s = str(v).strip()
            digits = re.findall(r"\d+", s)
            return int(digits[0]) if digits else None
        except Exception:
            return None

    groups_by_vid = {}
    no_vid_items = []
    for idx, f in enumerate(findings):
        vid = _to_int_vid(f.get("vuln_id"))
        if vid is not None:
            status_family = _structured_status_family(
                f.get("observation", ""),
                f.get("remediation_status", ""),
                f.get("risk_status", ""),
            )
            groups_by_vid.setdefault((vid, status_family), []).append((idx, f))
        else:
            no_vid_items.append((idx, f))

    chosen = {}

    def completeness_score(item):
        score = 0
        for k, v in item.items():
            if v is None:
                continue
            if isinstance(v, str) and v.strip() in (
                "[PLACEHOLDER]",
                "[INSUFFICIENT DATA]",
            ):
                continue
            if isinstance(v, (list, dict)) and len(v) == 0:
                continue
            score += 1
        desc_len = len(str(item.get("description") or ""))
        return (score, desc_len)

    for (_, _status_family), items in groups_by_vid.items():
        best = None
        best_score = (-1, -1)
        best_idx = None
        merged_assets = []
        for idx, f in items:
            sc = completeness_score(f)
            if sc > best_score:
                best_score = sc
                best = f
                best_idx = idx
            for asset in _normalize_affected_assets(f.get("affected_assets_raw")):
                if asset not in merged_assets:
                    merged_assets.append(asset)
            for asset in _normalize_affected_assets(f.get("affected_assets")):
                if asset not in merged_assets:
                    merged_assets.append(asset)
        if best is not None:
            merged = dict(best)
            if merged_assets:
                merged["affected_assets_raw"] = merged_assets
                merged["affected_assets"] = ", ".join(merged_assets)
            chosen[best_idx] = merged

    seen_fp = set()
    for idx, f in no_vid_items:
        name = _normalize_text_for_key(f.get("name", ""))
        assets = _normalize_text_for_key(
            ", ".join(_normalize_affected_assets(f.get("affected_assets_raw")))
        )
        if not assets:
            assets = _normalize_text_for_key(
                ", ".join(_normalize_affected_assets(f.get("affected_assets")))
            )
        desc = _normalize_text_for_key((f.get("description", "") or "")[:300])
        key = (name, assets, desc)
        if key in seen_fp:
            continue
        seen_fp.add(key)
        chosen[idx] = f

    out = [_build_finding_presentation(chosen[i]) for i in sorted(chosen.keys())]
    return out


# ── python-docx Renderer ──────────────────────────────────────────────────────

# ── Sizing constants from handmade template (in dxa = 1/20th pt) ──────────────
_COL_W_NUM = 617  # First narrow column (row number / section counter)
_COL_W_LABEL = 2914  # Middle column (field label)
_COL_W_VALUE = 6826  # Right column (value)
_BG_HEADER = "2E5395"  # Deep blue for row-0 number cell
_BG_SEV = {
    "Critical": "C00000",
    "High": "FF0000",
    "Medium": "FF8C00",
    "Low": "FFD700",
}


def _truncate_assets(assets, max_items: int = 10) -> str:
    """Truncate asset values to max_items entries while tolerating list-like inputs."""
    if isinstance(assets, str):
        stripped_assets = assets.strip()
        if not stripped_assets:
            return "[INSUFFICIENT DATA]"
        if stripped_assets in ("[PLACEHOLDER]", "[INSUFFICIENT DATA]"):
            return stripped_assets

    parts = _normalize_affected_assets(assets)
    if not parts:
        return "[INSUFFICIENT DATA]"
    if len(parts) <= max_items:
        return ", ".join(parts)
    shown = ", ".join(parts[:max_items])
    return f"{shown}, … and {len(parts) - max_items} more"


def _has_meaningful_asset_trace_block(value) -> bool:
    text = _normalize_report_text(value)
    if not text or _is_placeholder_text(text):
        return False

    stripped = re.sub(r"^\s*affected assets\s*:\s*", "", text, flags=re.IGNORECASE)
    for line in stripped.splitlines() or [stripped]:
        candidate_line = re.sub(r"^\s*[-*]\s*", "", line).strip()
        if not candidate_line:
            continue
        for part in candidate_line.split(","):
            token = part.strip()
            if token and not _is_placeholder_asset_token(token):
                return True
    return False


def _compose_summary_group_label(finding: dict) -> str:
    structured_label = str(finding.get("display_group_label") or "").strip()
    if _has_meaningful_text(structured_label):
        return structured_label

    taxonomy_label = str(finding.get("taxonomy_label") or "").strip()
    subnet_label = str(finding.get("subnet_label") or "").strip()
    if _has_meaningful_text(taxonomy_label) and _has_meaningful_text(subnet_label):
        if taxonomy_label == subnet_label:
            return taxonomy_label
        return f"{taxonomy_label} / {subnet_label}"
    if _has_meaningful_text(taxonomy_label):
        return taxonomy_label
    if _has_meaningful_text(subnet_label):
        return subnet_label

    category = str(finding.get("category") or "").strip()
    if _has_meaningful_text(category):
        return category
    return "General Findings"


def _build_finding_presentation(finding: dict) -> dict:
    """Add traceability-focused presentation fields without changing core finding data."""
    result = dict(finding or {})
    original_affected_assets = result.get("affected_assets")

    affected_assets_raw = _normalize_affected_assets(result.get("affected_assets_raw"))
    if not affected_assets_raw:
        affected_assets_raw = _normalize_affected_assets(result.get("affected_assets"))

    if affected_assets_raw:
        normalized_affected_assets = ", ".join(affected_assets_raw)
    else:
        normalized_affected_assets = "[INSUFFICIENT DATA]"

    if isinstance(original_affected_assets, str) and _has_meaningful_text(
        original_affected_assets
    ):
        result["affected_assets"] = original_affected_assets
    else:
        result["affected_assets"] = normalized_affected_assets
    result["affected_assets_raw"] = affected_assets_raw
    result["affected_assets_short"] = _truncate_assets(normalized_affected_assets)
    if affected_assets_raw:
        result["asset_trace_block"] = "Affected assets:\n" + "\n".join(
            f"- {asset}" for asset in affected_assets_raw
        )
    else:
        result["asset_trace_block"] = "Affected assets: [INSUFFICIENT DATA]"

    result["display_title"] = _display_title_from_finding(result)
    return result


def _enrich_finding_presentation(finding: dict) -> dict:
    """Backward-compatible alias for older call sites/tests."""
    return _build_finding_presentation(finding)


def _has_meaningful_text(value) -> bool:
    text = str(value or "").strip()
    return bool(text) and not _is_placeholder_text(text)


def _slugify_report_token(value: str) -> str:
    normalized = unicodedata.normalize("NFKD", str(value or ""))
    ascii_value = normalized.encode("ascii", "ignore").decode("ascii")
    return re.sub(r"[^a-z0-9]+", "-", ascii_value.lower()).strip("-")


def _title_from_slug(slug: str, fallback: str) -> str:
    cleaned = str(slug or "").replace("-", " ").strip()
    return cleaned.title() if cleaned else fallback


def _display_title_from_finding(finding: dict) -> str:
    raw_title = (
        str(
            finding.get("name")
            or finding.get("short_name")
            or finding.get("title")
            or "[INSUFFICIENT DATA]"
        ).strip()
        or "[INSUFFICIENT DATA]"
    )
    if raw_title == "[INSUFFICIENT DATA]":
        return raw_title
    raw_title = _strip_leading_plugin_id(raw_title)
    _, normalized_title = _structured_group_and_title_for_row(finding, raw_title)
    return str(normalized_title or raw_title).strip() or "[INSUFFICIENT DATA]"


def _ensure_sentence(text: str) -> str:
    sentence = str(text or "").strip()
    if not sentence:
        return ""
    return sentence if sentence[-1] in ".!?" else sentence + "."


def _join_sentences(*parts: str) -> str:
    sentences = []
    for part in parts:
        sentence = _ensure_sentence(part)
        if sentence and sentence not in sentences:
            sentences.append(sentence)
    return " ".join(sentences)


def _resolve_front_matter_text(data: dict) -> dict[str, str]:
    source = data or {}

    def _clean_text(value) -> str:
        normalized = _normalize_report_text(value)
        return normalized if _has_meaningful_text(normalized) else ""

    def _resolve_objectives_text(value) -> str:
        if isinstance(value, (list, tuple, set)):
            objective_parts = [_clean_text(item) for item in value]
            return _join_sentences(*(part for part in objective_parts if part))
        return _clean_text(value)

    def _pick_text(*values) -> str:
        for value in values:
            text = _clean_text(value)
            if text:
                return text
        return ""

    objectives_text = ""
    for objectives_value in (
        source.get("objectives"),
        (source.get("narrative_slots") or {}).get("objectives"),
    ):
        objectives_text = _resolve_objectives_text(objectives_value)
        if objectives_text:
            break

    return {
        "executive_summary": _pick_text(
            source.get("introduction_overview"),
            source.get("executive_summary"),
        ),
        "introduction": _pick_text(
            source.get("introduction_scope_bridge"),
            source.get("scope_summary"),
        ),
        "objectives": objectives_text,
    }


def _derive_taxonomy_family(finding: dict) -> str:
    existing = str(finding.get("taxonomy_family") or "").strip()
    if _has_meaningful_text(existing):
        return _slugify_report_token(existing) or "general"

    category = str(finding.get("category") or "")
    if _has_meaningful_text(category):
        cleaned = re.sub(r"\b(?:\d{1,3}\.){3}\d{1,3}(?:/\d{1,2})?\b", " ", category)
        cleaned = re.sub(r"[_|/]+", " ", cleaned)
        cleaned = re.sub(r"\s+-\s+", " ", cleaned)
        cleaned = re.sub(r"\s+", " ", cleaned).strip(" -")
        slug = _slugify_report_token(cleaned)
        if slug:
            return slug

    return "general"


def _derive_taxonomy_label(finding: dict, taxonomy_family: str) -> str:
    category = str(finding.get("category") or "")
    if _has_meaningful_text(category):
        cleaned = re.sub(r"\b(?:\d{1,3}\.){3}\d{1,3}(?:/\d{1,2})?\b", " ", category)
        cleaned = re.sub(r"[_|/]+", " ", cleaned)
        cleaned = re.sub(r"\s+-\s+", " ", cleaned)
        cleaned = re.sub(r"\s+", " ", cleaned).strip(" -")
        if cleaned:
            return cleaned
    if taxonomy_family == "general":
        return "General Findings"
    return _title_from_slug(taxonomy_family, "General Findings")


def _ip_to_subnet_key(ip_address: str) -> str:
    parts = str(ip_address or "").split(".")
    if len(parts) != 4:
        return ""
    try:
        octets = [int(part) for part in parts]
    except ValueError:
        return ""
    if any(octet < 0 or octet > 255 for octet in octets):
        return ""
    return f"{octets[0]}.{octets[1]}.{octets[2]}.0/24"


def _collect_subnet_candidates(finding: dict) -> list[str]:
    candidates = []
    category = str(finding.get("category") or "")
    for subnet in re.findall(r"\b(?:\d{1,3}\.){3}\d{1,3}/\d{1,2}\b", category):
        if subnet not in candidates:
            candidates.append(subnet)

    asset_values = _normalize_affected_assets(
        finding.get("affected_assets_raw") or finding.get("affected_assets")
    )
    for asset in asset_values:
        for host in re.findall(r"\b((?:\d{1,3}\.){3}\d{1,3})(?::\d+)?\b", asset):
            subnet = _ip_to_subnet_key(host)
            if subnet and subnet not in candidates:
                candidates.append(subnet)

    return candidates


def _derive_subnet_fields(finding: dict) -> tuple[str, str]:
    existing_key = str(finding.get("subnet_key") or "").strip()
    existing_label = str(finding.get("subnet_label") or "").strip()
    if _has_meaningful_text(existing_key) and _has_meaningful_text(existing_label):
        return existing_key, existing_label

    subnet_candidates = _collect_subnet_candidates(finding)
    if len(subnet_candidates) == 1:
        return subnet_candidates[0], subnet_candidates[0]
    if len(subnet_candidates) > 1:
        return "multiple-subnets", "Multiple subnets"
    return "host-specific", "Host-specific findings"


def _finding_outline_sort_key(finding: dict) -> tuple:
    return (
        -_STRUCTURED_SEVERITY_ORDER.get(finding.get("severity", "Informational"), 0),
        _normalize_text_for_key(
            finding.get("display_title") or finding.get("name") or ""
        ),
        _normalize_text_for_key(finding.get("id") or ""),
    )


def _take_unused_finding_by_id(
    findings_by_id: dict[str, list[tuple[int, dict]]],
    finding_id: str,
    used_indexes: set[int],
) -> tuple[int, dict] | None:
    for finding_index, finding in findings_by_id.get(finding_id, []):
        if finding_index in used_indexes:
            continue
        used_indexes.add(finding_index)
        return finding_index, finding
    return None


def _finding_matches_outline_group(finding: dict, group: dict) -> bool:
    taxonomy_family = str(group.get("taxonomy_family") or "").strip()
    subnet_key = str(group.get("subnet_key") or "").strip()

    if (
        taxonomy_family
        and str(finding.get("taxonomy_family") or "").strip() != taxonomy_family
    ):
        return False
    if subnet_key and str(finding.get("subnet_key") or "").strip() != subnet_key:
        return False
    return True


def _collect_outline_group_matches(
    indexed_findings: list[tuple[int, dict]],
    findings_by_id: dict[str, list[tuple[int, dict]]],
    group: dict,
    used_indexes: set[int],
) -> list[tuple[int, dict]]:
    matched = []
    for finding_id in group.get("finding_ids") or []:
        pair = _take_unused_finding_by_id(
            findings_by_id, str(finding_id or "").strip(), used_indexes
        )
        if pair is not None:
            matched.append(pair)

    if matched:
        return matched

    for finding_index, finding in indexed_findings:
        if finding_index in used_indexes:
            continue
        if not _finding_matches_outline_group(finding, group):
            continue
        matched.append((finding_index, finding))
        used_indexes.add(finding_index)
    return matched


def _order_findings_for_outline_groups(
    findings: list[dict], outline_groups: list[dict] | None
) -> list[dict]:
    if not isinstance(outline_groups, list) or not outline_groups:
        return list(findings)

    indexed_findings = list(enumerate(findings))
    findings_by_id = {}
    for finding_index, finding in indexed_findings:
        finding_id = str(finding.get("id") or "").strip()
        if finding_id:
            findings_by_id.setdefault(finding_id, []).append((finding_index, finding))

    used_indexes = set()
    ordered = []

    for group in outline_groups:
        matched = _collect_outline_group_matches(
            indexed_findings, findings_by_id, group, used_indexes
        )

        if matched:
            ordered.extend(
                finding
                for _, finding in sorted(
                    matched, key=lambda item: _finding_outline_sort_key(item[1])
                )
            )

    ordered.extend(
        finding
        for finding_index, finding in indexed_findings
        if finding_index not in used_indexes
    )
    return ordered


def _compose_outline_groups(findings: list[dict]) -> list[dict]:
    grouped = {}
    for finding in findings:
        key = (finding["taxonomy_family"], finding["subnet_key"])
        grouped.setdefault(key, []).append(finding)

    outline_groups = []
    for taxonomy_family, subnet_key in sorted(grouped.keys()):
        grouped_findings = sorted(
            grouped[(taxonomy_family, subnet_key)], key=_finding_outline_sort_key
        )
        outline_groups.append(
            {
                "taxonomy_family": taxonomy_family,
                "taxonomy_label": grouped_findings[0].get(
                    "taxonomy_label",
                    _title_from_slug(taxonomy_family, "General Findings"),
                ),
                "subnet_key": subnet_key,
                "subnet_label": grouped_findings[0].get(
                    "subnet_label", "Host-specific findings"
                ),
                "highest_severity": grouped_findings[0].get(
                    "severity", "Informational"
                ),
                "finding_count": len(grouped_findings),
                "finding_ids": [
                    finding.get("id")
                    for finding in grouped_findings
                    if finding.get("id")
                ],
            }
        )
    return outline_groups


def _narrative_severity_rank(severity: str) -> int:
    return {
        "Critical": 0,
        "High": 1,
        "Medium": 2,
        "Low": 3,
        "Informational": 4,
    }.get(str(severity or ""), 5)


def _narrative_severity_counts(findings: list[dict]) -> dict[str, int]:
    counts = {
        "Critical": 0,
        "High": 0,
        "Medium": 0,
        "Low": 0,
        "Informational": 0,
    }
    for finding in findings:
        severity = str(finding.get("severity") or "")
        if severity in counts:
            counts[severity] += 1
    return counts


def _narrative_theme_label(finding: dict) -> str:
    taxonomy_label = str(finding.get("taxonomy_label") or "").strip()
    if _has_meaningful_text(taxonomy_label):
        return taxonomy_label

    taxonomy_family = str(finding.get("taxonomy_family") or "").strip()
    if not _has_meaningful_text(taxonomy_family):
        taxonomy_family = _derive_taxonomy_family(finding)
    return _derive_taxonomy_label(finding, taxonomy_family)


def _dominant_taxonomy_themes(findings: list[dict], limit: int = 2) -> list[str]:
    theme_summary: dict[str, dict[str, int]] = {}

    for finding in findings:
        label = _narrative_theme_label(finding)
        if not _has_meaningful_text(label):
            continue

        summary = theme_summary.setdefault(
            label,
            {
                "count": 0,
                "best_rank": _narrative_severity_rank(None),
            },
        )
        summary["count"] += 1
        summary["best_rank"] = min(
            summary["best_rank"],
            _narrative_severity_rank(finding.get("severity")),
        )

    ordered = sorted(
        theme_summary.items(),
        key=lambda item: (-item[1]["count"], item[1]["best_rank"], item[0]),
    )
    return [label for label, _ in ordered[:limit]]


def _join_report_labels(labels: list[str]) -> str:
    cleaned = [str(label).strip() for label in labels if _has_meaningful_text(label)]
    if not cleaned:
        return ""
    if len(cleaned) == 1:
        return cleaned[0]
    if len(cleaned) == 2:
        return f"{cleaned[0]} and {cleaned[1]}"
    return ", ".join(cleaned[:-1]) + f", and {cleaned[-1]}"


def _narrative_subnet_key(finding: dict) -> str:
    subnet_key = str(finding.get("subnet_key") or "").strip()
    if _has_meaningful_text(subnet_key):
        return subnet_key
    derived_key, _ = _derive_subnet_fields(finding)
    return derived_key or "host-specific"


def _compose_report_objectives(data: dict, findings: list[dict]) -> list[str]:
    existing = data.get("objectives")
    if isinstance(existing, list):
        cleaned_existing = [
            str(item).strip() for item in existing if _has_meaningful_text(item)
        ]
        if cleaned_existing:
            return cleaned_existing

    if not findings:
        return []

    severity_counts = _narrative_severity_counts(findings)
    family_count = len(
        {
            str(finding.get("taxonomy_family") or _derive_taxonomy_family(finding))
            or "general"
            for finding in findings
        }
    )
    subnet_count = len({_narrative_subnet_key(finding) for finding in findings})
    total_findings = len(findings)
    dominant_theme_phrase = _join_report_labels(_dominant_taxonomy_themes(findings))

    if severity_counts["Critical"]:
        critical_count = severity_counts["Critical"]
        first_objective = f"Contain and remediate the {critical_count} critical finding{'s' if critical_count != 1 else ''} driving risk in {dominant_theme_phrase or 'the in-scope environment'}."
    elif severity_counts["High"]:
        high_count = severity_counts["High"]
        first_objective = f"Prioritise the {high_count} high-severity finding{'s' if high_count != 1 else ''} driving risk in {dominant_theme_phrase or 'the in-scope environment'}."
    else:
        first_objective = f"Sequence planned remediation for the {total_findings} confirmed medium- and low-severity finding{'s' if total_findings != 1 else ''} across {dominant_theme_phrase or 'the in-scope environment'}."

    return [
        first_objective,
        f"Coordinate remediation workstreams around {dominant_theme_phrase or 'the in-scope environment'} while tracking progress across {family_count} taxonomy {'family' if family_count == 1 else 'families'} and {subnet_count} subnet group{'s' if subnet_count != 1 else ''}.",
        "Preserve full asset traceability for each confirmed finding while maintaining compact display fields for report-facing views.",
    ]


def _compose_risk_posture(findings: list[dict]) -> str:
    severity_counts = _narrative_severity_counts(findings)
    dominant_theme_phrase = _join_report_labels(_dominant_taxonomy_themes(findings))
    dominant_theme_clause = (
        f", concentrated in {dominant_theme_phrase}" if dominant_theme_phrase else ""
    )

    if severity_counts["Critical"]:
        count = severity_counts["Critical"]
        return f"The current risk posture remains elevated because {count} critical finding{'s' if count != 1 else ''} require immediate remediation{dominant_theme_clause}."
    if severity_counts["High"]:
        count = severity_counts["High"]
        return f"The current risk posture remains elevated because {count} high-severity finding{'s' if count != 1 else ''} require prioritised remediation{dominant_theme_clause}."
    if findings:
        medium_low_clause = (
            f" with the strongest concentration in {dominant_theme_phrase}"
            if dominant_theme_phrase
            else ""
        )
        return (
            "The current risk posture is driven primarily by medium- and low-severity findings"
            + medium_low_clause
            + " that should be addressed through planned remediation."
        )
    return (
        "No confirmed findings were available for deterministic narrative composition."
    )


def _derive_scope_sentence(outline_groups: list[dict]) -> str:
    subnet_labels = [group["subnet_label"] for group in outline_groups]
    if subnet_labels:
        return (
            "The validated scope covered "
            + ", ".join(subnet_labels)
            + " within the assessed environment"
        )
    return "The validated scope covered the assessed environment"


def _compose_scope_bridge(data: dict, outline_groups: list[dict]) -> str:
    narrative_slots = dict(data.get("narrative_slots") or {})

    intro_scope = str(narrative_slots.get("intro_scope") or "").strip()
    if not _has_meaningful_text(intro_scope):
        derived_scope = _derive_scope_sentence(outline_groups) if outline_groups else ""
        if _has_meaningful_text(derived_scope):
            intro_scope = derived_scope
        else:
            scope_summary = str(data.get("scope_summary") or "").strip()
            intro_scope = (
                scope_summary
                if _has_meaningful_text(scope_summary)
                else _derive_scope_sentence(outline_groups)
            )

    intro_method = str(narrative_slots.get("intro_method") or "").strip()
    if not _has_meaningful_text(intro_method):
        methodology = str(data.get("methodology") or "").strip()
        if _has_meaningful_text(methodology):
            intro_method = f"The assessment approach followed {methodology}"
        else:
            intro_method = "The assessment approach used validated scan-derived findings and deterministic grouping to support report composition"

    return _join_sentences(intro_scope, intro_method)


_REFRESHED_FINDING_FIELDS = {
    "affected_assets_raw",
    "affected_assets_short",
    "asset_trace_block",
    "display_title",
    "taxonomy_family",
    "taxonomy_label",
    "subnet_key",
    "subnet_label",
    "display_control_name",
    "display_control_objective",
}


def _reset_finding_composition_state(finding: dict) -> dict:
    refreshed = dict(finding or {})
    for field in _REFRESHED_FINDING_FIELDS:
        refreshed.pop(field, None)
    return refreshed


def compose_report_narrative(data: dict, *, refresh: bool = False) -> dict:
    """Return a pure, deterministic report view with narrative and grouping metadata."""
    composed = dict(data or {})
    source_findings = composed.get("findings") or []

    findings = []
    for finding in source_findings:
        source_finding = (
            _reset_finding_composition_state(finding) if refresh else finding
        )
        result = _build_finding_presentation(source_finding)

        taxonomy_family = _derive_taxonomy_family(result)
        taxonomy_label = _derive_taxonomy_label(result, taxonomy_family)
        subnet_key, subnet_label = _derive_subnet_fields(result)

        result["taxonomy_family"] = taxonomy_family or "general"
        result["taxonomy_label"] = taxonomy_label or "General Findings"
        result["subnet_key"] = subnet_key or "host-specific"
        result["subnet_label"] = subnet_label or "Host-specific findings"

        if not _has_meaningful_text(result.get("display_control_name")):
            control_name = str(result.get("control_name") or "").strip()
            result["display_control_name"] = (
                control_name
                if _has_meaningful_text(control_name)
                else "Vulnerability Remediation"
            )
        if not _has_meaningful_text(result.get("display_control_objective")):
            control_objective = str(result.get("control_objective") or "").strip()
            result["display_control_objective"] = (
                control_objective
                if _has_meaningful_text(control_objective)
                else "Identify and remediate the vulnerability to reduce the attack surface."
            )

        findings.append(result)

    outline_groups = _compose_outline_groups(findings)
    objectives = _compose_report_objectives({} if refresh else composed, findings)

    narrative_slots = {} if refresh else dict(composed.get("narrative_slots") or {})
    summary_context = narrative_slots.get("summary_context")
    if not _has_meaningful_text(summary_context):
        total_findings = len(findings)
        total_groups = len(outline_groups)
        if findings or outline_groups:
            summary_context = f"The validated assessment identified {total_findings} confirmed finding{'s' if total_findings != 1 else ''} across {total_groups} grouped remediation workstream{'s' if total_groups != 1 else ''}."
    if not _has_meaningful_text(summary_context):
        summary_context = str(composed.get("executive_summary") or "").strip()
    if not _has_meaningful_text(summary_context):
        summary_context = f"The validated assessment identified {total_findings} confirmed finding{'s' if total_findings != 1 else ''} across {total_groups} grouped remediation workstream{'s' if total_groups != 1 else ''}."

    summary_risk_posture = narrative_slots.get("summary_risk_posture")
    if not _has_meaningful_text(summary_risk_posture):
        summary_risk_posture = _compose_risk_posture(findings)

    summary_key_drivers = narrative_slots.get("summary_key_drivers")
    if not _has_meaningful_text(summary_key_drivers):
        taxonomy_labels = [group["taxonomy_label"] for group in outline_groups[:3]]
        if taxonomy_labels:
            summary_key_drivers = "Primary remediation drivers include " + ", ".join(
                taxonomy_labels
            )
        else:
            summary_key_drivers = "Primary remediation drivers were not available in the validated findings"

    intro_scope = narrative_slots.get("intro_scope")
    if not _has_meaningful_text(intro_scope):
        derived_scope = _derive_scope_sentence(outline_groups) if outline_groups else ""
        if _has_meaningful_text(derived_scope):
            intro_scope = derived_scope
        else:
            scope_summary = str(composed.get("scope_summary") or "").strip()
            intro_scope = (
                scope_summary
                if _has_meaningful_text(scope_summary)
                else _derive_scope_sentence(outline_groups)
            )

    intro_method = narrative_slots.get("intro_method")
    if not _has_meaningful_text(intro_method):
        methodology = str(composed.get("methodology") or "").strip()
        if _has_meaningful_text(methodology):
            intro_method = f"The assessment approach followed {methodology}"
        else:
            intro_method = "The assessment approach used validated scan-derived findings and deterministic grouping to support report composition"

    narrative_slots["summary_context"] = summary_context
    narrative_slots["summary_risk_posture"] = summary_risk_posture
    narrative_slots["summary_key_drivers"] = summary_key_drivers
    narrative_slots["intro_scope"] = intro_scope
    narrative_slots["intro_method"] = intro_method
    narrative_slots["objectives"] = list(objectives)

    composed["findings"] = findings
    composed["narrative_slots"] = narrative_slots
    composed["objectives"] = list(objectives)
    composed["outline_groups"] = outline_groups

    introduction_overview = (
        "" if refresh else str(composed.get("introduction_overview") or "").strip()
    )
    if not _has_meaningful_text(introduction_overview):
        introduction_overview = _join_sentences(
            summary_context, summary_risk_posture, summary_key_drivers
        )
    composed["introduction_overview"] = introduction_overview

    introduction_scope_bridge = (
        "" if refresh else str(composed.get("introduction_scope_bridge") or "").strip()
    )
    if not _has_meaningful_text(introduction_scope_bridge):
        introduction_scope_bridge = _compose_scope_bridge(composed, outline_groups)
    composed["introduction_scope_bridge"] = introduction_scope_bridge

    return composed


_CVE_RE = re.compile(r"CVE[-\s]?\d{4}[-\s]?\d{4,7}", re.IGNORECASE)
_CWE_RE = re.compile(r"CWE[-\s]?\d{1,5}", re.IGNORECASE)


def _extract_cves_cwes_from_finding(finding: dict) -> tuple[list[str], list[str]]:
    """Collect CVE + CWE IDs from finding across possible storage keys + reference list."""
    haystacks: list[str] = []
    for k in ("cve", "cwe", "cwe_ids", "category", "description"):
        v = finding.get(k)
        if isinstance(v, list):
            haystacks.extend(str(x) for x in v)
        elif v:
            haystacks.append(str(v))
    refs = finding.get("reference")
    if isinstance(refs, list):
        for r in refs:
            if isinstance(r, dict):
                haystacks.append(str(r.get("title", "")))
                haystacks.append(str(r.get("url", "")))
            else:
                haystacks.append(str(r))
    blob = " ".join(haystacks)
    cves, seen_c = [], set()
    for m in _CVE_RE.findall(blob):
        norm = re.sub(r"\s+", "-", m.upper())
        if norm not in seen_c:
            seen_c.add(norm)
            cves.append(norm)
    cwes, seen_w = [], set()
    for m in _CWE_RE.findall(blob):
        norm = re.sub(r"\s+", "-", m.upper())
        if norm not in seen_w:
            seen_w.add(norm)
            cwes.append(norm)
    return cves, cwes


def _clean_cve_cvss(cve, _cvss, finding: dict | None = None) -> str:
    """Build CVE/CWE display string (CVEs + CWEs, comma-joined). CVSS not shown in table."""
    _EMPTY = {"nan", "none", "", "[placeholder]", "[insufficient data]", "n/a"}
    if finding is not None:
        cves, cwes = _extract_cves_cwes_from_finding(finding)
    else:
        raw = str(cve or "").strip()
        cves = (
            [m.upper() for m in _CVE_RE.findall(raw)]
            if raw.lower() not in _EMPTY
            else []
        )
        cwes = (
            [m.upper() for m in _CWE_RE.findall(raw)]
            if raw.lower() not in _EMPTY
            else []
        )
    parts = cves + cwes
    return ", ".join(parts) if parts else "N/A"


def _clean_status(status: str) -> str:
    """Normalise status to uppercase OPEN / CLOSED / IN PROGRESS."""
    _MAP = {"open": "OPEN", "closed": "CLOSED", "in progress": "IN PROGRESS"}
    return _MAP.get(str(status or "").strip().lower(), str(status or "OPEN").upper())


def _set_tc_text(tc, text: str):
    text = str(text) if text else ""
    paras = tc.findall(".//" + qn("w:p"))
    if not paras:
        return
    p = paras[0]
    for r in p.findall(qn("w:r")):
        p.remove(r)
    r_new = OxmlElement("w:r")
    t_new = OxmlElement("w:t")
    t_new.text = text
    if text and (text != text.strip() or "  " in text):
        t_new.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    r_new.append(t_new)
    p.append(r_new)


def _ensure_tc_pr(tc) -> OxmlElement:
    tc_pr = tc.find(qn("w:tcPr"))
    if tc_pr is None:
        tc_pr = OxmlElement("w:tcPr")
        tc.insert(0, tc_pr)
    return tc_pr


def _set_tc_fill(tc, fill: str):
    tc_pr = _ensure_tc_pr(tc)
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), str(fill or "FFFFFF"))


def _set_tc_run_color(tc, color: str):
    for run in tc.findall(".//" + qn("w:r")):
        r_pr = run.find(qn("w:rPr"))
        if r_pr is None:
            r_pr = OxmlElement("w:rPr")
            run.insert(0, r_pr)
        color_el = r_pr.find(qn("w:color"))
        if color_el is None:
            color_el = OxmlElement("w:color")
            r_pr.append(color_el)
        color_el.set(qn("w:val"), str(color or "000000"))
        for highlight in list(r_pr.findall(qn("w:highlight"))):
            r_pr.remove(highlight)


def _set_cell_text(cell, text: str):
    text = str(text) if text else ""
    paras = cell.paragraphs
    if not paras:
        cell.add_paragraph(text)
        return
    p = paras[0]
    runs = p.runs
    if runs:
        runs[0].text = text
        for r in runs[1:]:
            r.text = ""
    else:
        p.add_run(text)
    for extra in paras[1:]:
        for r in extra.runs:
            r.text = ""


def _replace_para_text(para, text: str):
    runs = para.runs
    if runs:
        runs[0].text = str(text) if text else ""
        for r in runs[1:]:
            r.text = ""
    else:
        para.add_run(str(text) if text else "")


def _make_tc(
    w: int,
    text: str,
    *,
    bold: bool = False,
    white_text: bool = False,
    bg: str = None,
    font_size: int = 24,
    span: int = 1,
    indent: int = None,
    spacing_before: int = 44,
) -> OxmlElement:
    """Build a w:tc element with a single paragraph + run, matching handmade style."""
    tc = OxmlElement("w:tc")
    tcPr = OxmlElement("w:tcPr")
    tcW = OxmlElement("w:tcW")
    tcW.set(qn("w:w"), str(w))
    tcW.set(qn("w:type"), "dxa")
    tcPr.append(tcW)
    if span > 1:
        gs = OxmlElement("w:gridSpan")
        gs.set(qn("w:val"), str(span))
        tcPr.append(gs)
    if bg:
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), bg)
        tcPr.append(shd)
    tc.append(tcPr)

    p = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr")
    pStyle = OxmlElement("w:pStyle")
    pStyle.set(qn("w:val"), "TableParagraph")
    pPr.append(pStyle)
    spac = OxmlElement("w:spacing")
    spac.set(qn("w:before"), str(spacing_before))
    pPr.append(spac)
    if indent is not None:
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), str(indent))
        pPr.append(ind)
    rPr_p = OxmlElement("w:rPr")
    sz_p = OxmlElement("w:sz")
    sz_p.set(qn("w:val"), str(font_size))
    rPr_p.append(sz_p)
    pPr.append(rPr_p)
    p.append(pPr)

    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    if bold:
        r_b = OxmlElement("w:b")
        rPr.append(r_b)
    if white_text:
        col = OxmlElement("w:color")
        col.set(qn("w:val"), "FFFFFF")
        rPr.append(col)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(font_size))
    rPr.append(sz)
    r.append(rPr)
    t = OxmlElement("w:t")
    t.text = str(text) if text else ""
    if t.text and (t.text != t.text.strip() or "  " in t.text):
        t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    r.append(t)
    p.append(r)
    tc.append(p)
    return tc


def _make_tr(height: int = None) -> OxmlElement:
    """Build a w:tr element with optional row height."""
    tr = OxmlElement("w:tr")
    if height:
        trPr = OxmlElement("w:trPr")
        trH = OxmlElement("w:trHeight")
        trH.set(qn("w:val"), str(height))
        trPr.append(trH)
        tr.append(trPr)
    return tr


def _make_label_value_row(label: str, value: str, height: int = 361) -> OxmlElement:
    """Build a 2-cell merged-label row (cols 1+2 merged for label, col 3 for value)."""
    tr = _make_tr(height)
    # Label cell spans col 1+2 (w = _COL_W_NUM + _COL_W_LABEL)
    tc_label = _make_tc(_COL_W_NUM + _COL_W_LABEL, label, span=2, indent=107)
    tr.append(tc_label)
    # Value cell
    tc_val = _make_tc(_COL_W_VALUE, value)
    tr.append(tc_val)
    return tr


def _build_finding_table(finding: dict, seq_num: int) -> OxmlElement:
    """
    Build a w:tbl element that exactly matches the handmade report's
    12-row finding card structure.

    Row layout:
      0  [seq_num | Affected URL /IP | <assets>]   — blue header cell
      1  [Vulnerability title / Observation | <name>]  — bold value
      2  [Severity | <sev>]
      3  [Status | OPEN/CLOSED]
      4  [Vulnerability point /Impact | <desc + impact>]
      5  [CVE /CWE | <cve_cvss>]
      6  [Control Objective | <obj>]
      7  [Control Name | <name>]
      8  [Audit Requirement | <req>]
      9  [Recommendation | <remed>]
     10  [Reference | <ref>]
     11  [New or Repeat Observation | <obs>]
    """
    tbl = OxmlElement("w:tbl")

    # ── Table properties (match handmade exactly) ────────────────────────────
    tblPr = OxmlElement("w:tblPr")
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), "0")
    tblW.set(qn("w:type"), "auto")
    tblPr.append(tblW)
    tblInd = OxmlElement("w:tblInd")
    tblInd.set(qn("w:w"), "182")
    tblInd.set(qn("w:type"), "dxa")
    tblPr.append(tblInd)
    tblBorders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "12")
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), "000000")
        tblBorders.append(b)
    tblPr.append(tblBorders)
    tblLayout = OxmlElement("w:tblLayout")
    tblLayout.set(qn("w:type"), "fixed")
    tblPr.append(tblLayout)
    tblCellMar = OxmlElement("w:tblCellMar")
    for side in ("left", "right"):
        m = OxmlElement(f"w:{side}")
        m.set(qn("w:w"), "0")
        m.set(qn("w:type"), "dxa")
        tblCellMar.append(m)
    tblPr.append(tblCellMar)
    tbl.append(tblPr)

    # ── Grid columns ─────────────────────────────────────────────────────────
    tblGrid = OxmlElement("w:tblGrid")
    for w in (_COL_W_NUM, _COL_W_LABEL, _COL_W_VALUE):
        gc = OxmlElement("w:gridCol")
        gc.set(qn("w:w"), str(w))
        tblGrid.append(gc)
    tbl.append(tblGrid)

    # ── Prepare values ───────────────────────────────────────────────────────
    sev = finding.get("severity", "Medium")
    asset_trace = str(finding.get("asset_trace_block") or "")
    assets = (
        asset_trace
        if _has_meaningful_asset_trace_block(asset_trace)
        else _truncate_assets(finding.get("affected_assets", ""))
    )
    name = str(finding.get("display_title") or finding.get("name") or "")
    finding_id = str(finding.get("id") or "").strip()
    if _has_meaningful_text(finding_id) and finding_id not in name:
        name = f"{name} ({finding_id})" if _has_meaningful_text(name) else finding_id
    status = _clean_status(finding.get("remediation_status", "Open"))
    obs = finding.get("observation", "New")
    cve_str = _clean_cve_cvss(
        finding.get("cve", ""), finding.get("cvss", ""), finding=finding
    )
    ctrl_obj = finding.get("control_objective") or ""
    ctrl_name = finding.get("control_name") or ""
    ref = _format_reference_cell(finding.get("reference", ""))

    audit_req = _prepare_audit_requirement(finding)

    _EMPTY_V = {"[placeholder]", "[insufficient data]", "nan", "", None}

    # Description + business impact merged. Strip leading severity labels
    # ("Critical:", "High:", etc.) that LLMs sometimes prepend.
    _SEV_PFX = re.compile(
        r"^\s*(Critical|High|Medium|Low|Informational|Info)\s*[:\-]\s*", re.IGNORECASE
    )
    desc = _SEV_PFX.sub("", str(finding.get("description", "") or ""))
    impact = _SEV_PFX.sub("", str(finding.get("business_impact", "") or ""))
    if impact and impact.strip().lower() not in _EMPTY_V:
        desc_combined = desc + "\n" + impact if desc else impact
    else:
        desc_combined = desc

    remediation = _format_recommendation_cell(finding.get("remediation", ""))

    # ── Row 0: seq_num | Affected URL /IP | assets (blue header) ────────────
    tr0 = _make_tr(height=630)
    tr0.append(
        _make_tc(_COL_W_NUM, str(seq_num), bg=_BG_HEADER, white_text=True, indent=278)
    )
    tr0.append(_make_tc(_COL_W_LABEL, "Affected URL /IP", indent=None))
    tc_assets = _make_tc(_COL_W_VALUE, assets)
    # Give asset cell a little top spacing like handmade
    _p = tc_assets.find(qn("w:p"))
    if _p is not None:
        _pPr = _p.find(qn("w:pPr"))
        if _pPr is not None:
            _sp = _pPr.find(qn("w:spacing"))
            if _sp is None:
                _sp = OxmlElement("w:spacing")
                _pPr.append(_sp)
            _sp.set(qn("w:before"), "24")
            _sp.set(qn("w:line"), "290")
            _sp.set(qn("w:lineRule"), "atLeast")
    tr0.append(tc_assets)
    tbl.append(tr0)

    # ── Row 1: Vulnerability title / Observation | <name> (bold value) ───────
    tr1 = _make_tr(height=361)
    tr1.append(
        _make_tc(
            _COL_W_NUM + _COL_W_LABEL,
            "Vulnerability title / Observation",
            span=2,
            indent=107,
        )
    )
    tc_name = _make_tc(_COL_W_VALUE, name, bold=True)
    tr1.append(tc_name)
    tbl.append(tr1)

    # ── Row 2: Severity ──────────────────────────────────────────────────────
    tr2 = _make_tr(height=361)
    tr2.append(_make_tc(_COL_W_NUM + _COL_W_LABEL, "Severity", span=2, indent=107))
    sev_bg = _BG_SEV.get(sev, "FFFFFF")
    tr2.append(_make_tc(_COL_W_VALUE, sev, bg=sev_bg))
    tbl.append(tr2)

    # ── Row 3: Status ────────────────────────────
    tr3 = _make_tr(height=361)
    tr3.append(_make_tc(_COL_W_NUM + _COL_W_LABEL, "Status", span=2, indent=107))
    tr3.append(_make_tc(_COL_W_VALUE, status))
    tbl.append(tr3)

    # ── Row 4: Vulnerability point /Impact ────────────────────
    tr4 = _make_tr()
    tr4.append(
        _make_tc(
            _COL_W_NUM + _COL_W_LABEL, "Vulnerability point /Impact", span=2, indent=107
        )
    )
    tr4.append(_make_tc(_COL_W_VALUE, desc_combined))
    tbl.append(tr4)

    # ── Row 5: CVE /CWE ───────────────────────────────────
    tr5 = _make_tr(height=361)
    tr5.append(_make_tc(_COL_W_NUM + _COL_W_LABEL, "CVE /CWE", span=2, indent=107))
    tr5.append(_make_tc(_COL_W_VALUE, cve_str))
    tbl.append(tr5)

    # ── Row 6: Control Objective ──────────────────────────────
    tr6 = _make_tr()
    tr6.append(
        _make_tc(_COL_W_NUM + _COL_W_LABEL, "Control Objective", span=2, indent=107)
    )
    tr6.append(_make_tc(_COL_W_VALUE, ctrl_obj))
    tbl.append(tr6)

    # ── Row 7: Control Name ─────────────────────────────────
    tr7 = _make_tr(height=361)
    tr7.append(_make_tc(_COL_W_NUM + _COL_W_LABEL, "Control Name", span=2, indent=107))
    tr7.append(_make_tc(_COL_W_VALUE, ctrl_name))
    tbl.append(tr7)

    # ── Row 8: Audit Requirement ──────────────────────────────
    tr8 = _make_tr()
    tr8.append(
        _make_tc(_COL_W_NUM + _COL_W_LABEL, "Audit Requirement", span=2, indent=107)
    )
    tr8.append(_make_tc(_COL_W_VALUE, audit_req))
    tbl.append(tr8)

    # ── Row 9: Recommendation ────────────────────────────────
    tr9 = _make_tr()
    tr9.append(
        _make_tc(_COL_W_NUM + _COL_W_LABEL, "Recommendation", span=2, indent=107)
    )
    tr9.append(_make_tc(_COL_W_VALUE, remediation))
    tbl.append(tr9)

    # ── Row 10: Reference ─────────────────────────────────────
    tr10 = _make_tr()
    tr10.append(_make_tc(_COL_W_NUM + _COL_W_LABEL, "Reference", span=2, indent=107))
    tr10.append(_make_tc(_COL_W_VALUE, ref))
    tbl.append(tr10)

    # ── Row 11: New or Repeat Observation ───────────────────────
    tr11 = _make_tr(height=519)
    tr11.append(
        _make_tc(
            _COL_W_NUM + _COL_W_LABEL,
            "New or Repeat Observation",
            span=2,
            indent=107,
            spacing_before=47,
        )
    )
    tr11.append(_make_tc(_COL_W_VALUE, obs, spacing_before=47))
    tbl.append(tr11)

    return tbl


# Legacy helper kept for backward compatibility
def _clone_finding_table(finding: dict, seq_num: int = 0):
    """Deprecated - builds a fresh table matching handmade schema."""
    return _build_finding_table(finding, seq_num)


def _update_chart_severity_counts(doc, findings: list) -> None:
    """Update the bar-chart severity counts in the embedded chart XML."""
    from lxml import etree as _et

    sev_order = ["Critical", "High", "Medium", "Low"]
    counts = dict.fromkeys(sev_order, 0)
    for f in findings:
        sev = str(f.get("severity") or "").strip()
        # case-insensitive match
        matched = next((k for k in counts if k.lower() == sev.lower()), None)
        if matched:
            counts[matched] += 1

    try:
        chart_part = None
        for rel in doc.part.rels.values():
            if hasattr(rel, "reltype") and "chart" in rel.reltype.lower():
                chart_part = rel.target_part
                break
        if chart_part is None:
            return

        NS = "http://schemas.openxmlformats.org/drawingml/2006/chart"
        tree = _et.fromstring(chart_part.blob)

        chart_categories = []
        for ser in tree.findall(f".//{{{NS}}}ser"):
            categories_by_idx = {}
            for cat_pt in ser.findall(f"./{{{NS}}}cat//{{{NS}}}pt"):
                try:
                    idx = int(cat_pt.get("idx", len(categories_by_idx)))
                except ValueError:
                    idx = len(categories_by_idx)
                label = cat_pt.findtext(f"./{{{NS}}}v", default="").strip()
                categories_by_idx[idx] = label

            if not categories_by_idx:
                categories_by_idx = {idx: sev for idx, sev in enumerate(sev_order)}

            if not chart_categories:
                chart_categories = [
                    categories_by_idx[idx] for idx in sorted(categories_by_idx)
                ]

            for val_pt in ser.findall(f"./{{{NS}}}val//{{{NS}}}pt"):
                try:
                    idx = int(val_pt.get("idx", 0))
                except ValueError:
                    idx = 0
                label = categories_by_idx.get(idx, "")
                sev_key = next(
                    (sev for sev in counts if sev.lower() == str(label).strip().lower()),
                    None,
                )
                value_el = val_pt.find(f"./{{{NS}}}v")
                if value_el is not None:
                    value_el.text = str(counts.get(sev_key, 0) if sev_key else 0)

        chart_part._blob = _et.tostring(
            tree, xml_declaration=True, encoding="UTF-8", standalone=True
        )

        try:
            from openpyxl import load_workbook

            workbook_part = None
            for rel in chart_part.rels.values():
                partname = str(getattr(rel.target_part, "partname", "")).lower()
                if "relationships/package" in rel.reltype.lower() and partname.endswith(
                    ".xlsx"
                ):
                    workbook_part = rel.target_part
                    break
            if workbook_part is not None:
                workbook = load_workbook(BytesIO(workbook_part.blob))
                sheet = workbook.active
                categories = chart_categories or list(sev_order)
                for row_idx, label in enumerate(categories, start=2):
                    sev_key = next(
                        (
                            sev
                            for sev in counts
                            if sev.lower() == str(label).strip().lower()
                        ),
                        None,
                    )
                    sheet.cell(row=row_idx, column=1).value = label
                    sheet.cell(row=row_idx, column=2).value = (
                        counts.get(sev_key, 0) if sev_key else 0
                    )
                output = BytesIO()
                workbook.save(output)
                workbook_part._blob = output.getvalue()
        except Exception as workbook_exc:
            import sys

            print(f"⚠️  Chart workbook update: {workbook_exc}", file=sys.stderr)
    except Exception as e:
        import sys

        print(f"⚠️  Chart update: {e}", file=sys.stderr)


def render_report(
    data: dict,
    template_path: str,
    output_path: str,
    include_summary_table: bool = False,
) -> str:
    """
    Render the final .docx using the user's actual template.

    Structural mapping (mirrors Internal_VAPT_Report_AEG_Vision.docx exactly):
      - Cover: Title paragraphs → client name
      - Table 0: Report Release Date / metadata
      - Table 1: Company info block
      - Body para after Heading1 'Introduction': scope_summary
      - Body para after Heading1 'Executive Summary': executive_summary
      - Detailed Observation: H2 per category + 12-row finding tables + PoC paragraphs
      - Summary table: only if include_summary_table=True (removed otherwise)
      - Conclusion: NOT modified (pre-written in template)
    """
    doc = Document(template_path)
    findings = _dedupe_findings(data.get("findings", []))
    data = {**data, "findings": findings}
    front_matter = _resolve_front_matter_text(data)
    body = doc.element.body
    children = list(body)

    # ── Helper: find child index by heading text ──────────────────────────────
    def _find_heading_idx(text_fragment: str, style_hint: str = "Heading") -> int:
        for i, child in enumerate(children):
            tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
            if tag != "p":
                continue
            sty = child.find("./" + qn("w:pPr") + "/" + qn("w:pStyle"))
            if sty is None:
                continue
            if style_hint.lower() not in sty.get(qn("w:val"), "").lower():
                continue
            txt = "".join(r.text or "" for r in child.findall(".//" + qn("w:t")))
            if text_fragment.lower() in txt.lower():
                return i
        return -1

    def _first_body_para_after(child_idx: int) -> int:
        for i, child in enumerate(children):
            if i <= child_idx:
                continue
            tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
            if tag != "p":
                continue
            sty = child.find("./" + qn("w:pPr") + "/" + qn("w:pStyle"))
            style_val = sty.get(qn("w:val"), "").lower() if sty is not None else ""
            if "heading" in style_val:
                break
            if "bodytext" in style_val or "body" in style_val or not style_val:
                return i
        return -1

    def _set_child_para_text(child_idx: int, text: str):
        child = children[child_idx]
        for r in child.findall(qn("w:r")):
            child.remove(r)
        r_new = OxmlElement("w:r")
        t_new = OxmlElement("w:t")
        t_new.text = str(text) if text else ""
        if t_new.text and t_new.text != t_new.text.strip():
            t_new.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        r_new.append(t_new)
        child.append(r_new)

    def _child_para_text(child_idx: int) -> str:
        child = children[child_idx]
        text = "".join(r.text or "" for r in child.findall(".//" + qn("w:t")))
        return _normalize_report_text(text)

    # ── Cover: client name ────────────────────────────────────────────────────
    client_name = data.get("client_name", "[Client Name]")
    title_paragraphs = []
    found_for = False
    for para in doc.paragraphs:
        if para.style.name.lower() != "title":
            if title_paragraphs:
                break
            continue
        title_paragraphs.append(para)
        txt = para.text.strip()
        if txt.lower() == "for":
            found_for = True
            continue
        if found_for and not txt:
            _replace_para_text(para, client_name)
            break

    # Fallback to the last blank title paragraph on the cover, not a body paragraph.
    try:
        if not found_for:
            for para in reversed(title_paragraphs):
                if not para.text.strip():
                    _replace_para_text(para, client_name)
                    break
    except Exception:
        pass

    # ── Table 0: Report Release Date ─────────────────────────────────────────
    try:
        _set_cell_text(doc.tables[0].rows[0].cells[1], data.get("report_date", ""))
    except Exception:
        pass

    # ── Table 1: Company name + date ─────────────────────────────────────────
    try:
        tbl1 = doc.tables[1]
        if tbl1.rows and len(tbl1.rows[0].cells) > 1:
            _set_cell_text(tbl1.rows[0].cells[1], client_name)
        if len(tbl1.rows) > 7 and len(tbl1.rows[7].cells) > 1:
            _set_cell_text(tbl1.rows[7].cells[1], data.get("report_date", ""))
    except Exception:
        pass

    # ── Introduction body: scope_summary ─────────────────────────────────────
    try:
        intro_idx = _find_heading_idx("Introduction", "Heading")
        if intro_idx >= 0:
            bp_idx = _first_body_para_after(intro_idx)
            if bp_idx >= 0:
                intro_text = front_matter["introduction"]
                if intro_text or _is_placeholder_text(_child_para_text(bp_idx)):
                    _set_child_para_text(bp_idx, intro_text)
    except Exception as e:
        print(f"⚠️  Introduction section: {e}", file=sys.stderr)

    # ── Executive Summary body ────────────────────────────────────────────────
    try:
        exec_idx = _find_heading_idx("Executive Summary", "Heading")
        if exec_idx >= 0:
            bp_idx = _first_body_para_after(exec_idx)
            if bp_idx >= 0:
                executive_text = front_matter["executive_summary"]
                if executive_text or _is_placeholder_text(_child_para_text(bp_idx)):
                    _set_child_para_text(bp_idx, executive_text)
    except Exception as e:
        print(f"⚠️  Executive Summary section: {e}", file=sys.stderr)

    # ── Objectives body ───────────────────────────────────────────────────────
    try:
        objectives_idx = _find_heading_idx("Objectives", "Heading")
        if objectives_idx >= 0:
            bp_idx = _first_body_para_after(objectives_idx)
            if bp_idx >= 0:
                objectives_text = front_matter["objectives"]
                if objectives_text or _is_placeholder_text(_child_para_text(bp_idx)):
                    _set_child_para_text(bp_idx, objectives_text)
    except Exception as e:
        print(f"⚠️  Objectives section: {e}", file=sys.stderr)

    # ── Engagement Scope table — expand to 9 columns ─────────────────────────
    try:
        scope_tbl_idx = None
        for tidx, tbl in enumerate(doc.tables):
            if not tbl.rows:
                continue
            hdr_texts = [c.text.strip() for c in tbl.rows[0].cells]
            if "Asset Description" in hdr_texts or (
                len(hdr_texts) >= 2
                and "Sr. No" in hdr_texts[0]
                and "Asset" in hdr_texts[1]
            ):
                scope_tbl_idx = tidx
                break
        if scope_tbl_idx is not None:
            tbl = doc.tables[scope_tbl_idx]
            tbl_xml = tbl._tbl
            rows = tbl_xml.findall(qn("w:tr"))
            for row_el in rows[1:]:
                tbl_xml.remove(row_el)
            # Parse subnets from finding categories
            import re as _re2

            subnets = []
            seen_subnets = set()
            for f in findings:
                cat = str(f.get("category", "") or "")
                cidrs = _re2.findall(r"\b(?:\d{1,3}\.){3}\d{1,3}(?:/\d{1,2})?\b", cat)
                for cidr in cidrs:
                    if cidr not in seen_subnets:
                        seen_subnets.add(cidr)
                        subnets.append(cidr)
            if not subnets:
                subnets = ["[INSUFFICIENT DATA]"]
            for sn_idx, subnet in enumerate(subnets, start=1):
                new_row = copy.deepcopy(rows[0])
                cells = new_row.findall(qn("w:tc"))
                data_vals = [
                    str(sn_idx),
                    "Internal Network Subnet",
                    "",
                    subnet,
                    "-",
                    "-",
                    "-",
                    "-",
                    "-",
                ]
                for ci, (tc_el, val) in enumerate(zip(cells, data_vals)):
                    _set_tc_text(tc_el, val)
                tbl_xml.append(new_row)
    except Exception as e:
        print(f"⚠️  Engagement Scope table: {e}", file=sys.stderr)

    # ── Tools table — remove Tool Description column ──────────────────────────
    try:
        for tbl in doc.tables:
            if not tbl.rows:
                continue
            hdr = " ".join(c.text.strip() for c in tbl.rows[0].cells)
            if "Tool" in hdr and "Software" in hdr and len(tbl.rows[0].cells) >= 5:
                tbl_xml = tbl._tbl
                for row_el in tbl_xml.findall(qn("w:tr")):
                    cells = row_el.findall(qn("w:tc"))
                    if len(cells) >= 5:
                        row_el.remove(cells[2])
                break
    except Exception as e:
        print(f"⚠️  Tools table: {e}", file=sys.stderr)

    # ── Summary of Vulnerabilities table ─────────────────────────────────────
    summary_tbl_xml = None
    for tbl in doc.tables:
        if not tbl.rows or len(tbl.rows) < 2:
            continue
        hdr_cells = [c.text.strip() for c in tbl.rows[0].cells[:4]]
        if "Severity" in hdr_cells and (
            "Vulnerability" in hdr_cells or "Vulnerability" in " ".join(hdr_cells)
        ):
            summary_tbl_xml = tbl._tbl
            break

    if summary_tbl_xml is not None:
        if include_summary_table:
            summary_findings = _order_findings_for_outline_groups(
                findings, data.get("outline_groups")
            )
            rows = summary_tbl_xml.findall(qn("w:tr"))
            header_row_xml = rows[0]
            for row_el in rows[1:]:
                summary_tbl_xml.remove(row_el)
            for i, finding in enumerate(summary_findings, start=1):
                new_row = copy.deepcopy(header_row_xml)
                cells = new_row.findall(qn("w:tc"))
                if len(cells) >= 7:
                    for cell_el in cells:
                        _set_tc_fill(cell_el, "FFFFFF")
                        _set_tc_run_color(cell_el, "000000")
                    summary_title = finding.get("display_title")
                    if not _has_meaningful_text(summary_title):
                        summary_title = finding.get("name", "")
                    summary_assets = finding.get("affected_assets_short")
                    if not _has_meaningful_text(summary_assets):
                        summary_assets = _truncate_assets(
                            finding.get("affected_assets", "")
                        )
                    sev = finding.get("severity", "")
                    summary_values = [
                        str(i),
                        _compose_summary_group_label(finding),
                        sev,
                        summary_title,
                        _clean_status(finding.get("remediation_status", "Open")),
                        finding.get("observation", "New"),
                        summary_assets,
                    ]
                    for cell_index, cell_el in enumerate(cells):
                        value = (
                            summary_values[cell_index]
                            if cell_index < len(summary_values)
                            else ""
                        )
                        _set_tc_text(cell_el, value)
                    _set_tc_fill(cells[2], _BG_SEV.get(sev, "FFFFFF"))
                    _set_tc_run_color(
                        cells[2],
                        "FFFFFF" if sev in {"Critical", "High"} else "000000",
                    )
                summary_tbl_xml.append(new_row)
        else:
            # Remove summary table entirely
            parent = summary_tbl_xml.getparent()
            if parent is not None:
                parent.remove(summary_tbl_xml)

    # ── Detailed Observation: H2 per category + finding tables + PoC ─────────
    try:
        # Find 'Detailed Observation' heading child index
        detail_idx = _find_heading_idx("Detailed Observation", "Heading")
        # Find template finding table (first table after heading)
        tpl_tbl_xml = None
        for i, child in enumerate(children):
            if i <= detail_idx:
                continue
            tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
            if tag == "tbl":
                tpl_tbl_xml = child
                break

        if findings:
            from collections import OrderedDict

            _SEV_ORDER = {"Critical": 0, "High": 1, "Medium": 2, "Low": 3}

            def _build_fallback_render_groups(
                source_findings: list[dict],
            ) -> list[dict]:
                findings_sorted = sorted(
                    source_findings,
                    key=lambda f: (
                        _SEV_ORDER.get(f.get("severity", "Low"), 3),
                        str(f.get("category", "")),
                        str(f.get("display_title") or f.get("name") or ""),
                        str(f.get("id") or ""),
                    ),
                )

                grouped: OrderedDict = OrderedDict()
                for finding in findings_sorted:
                    taxonomy_label = str(finding.get("taxonomy_label") or "").strip()
                    subnet_label = str(finding.get("subnet_label") or "").strip()
                    has_taxonomy_label = _has_meaningful_text(taxonomy_label)
                    has_subnet_label = _has_meaningful_text(subnet_label)
                    if has_taxonomy_label or has_subnet_label:
                        primary_label = (
                            taxonomy_label if has_taxonomy_label else subnet_label
                        )
                        secondary_label = (
                            subnet_label
                            if has_subnet_label and subnet_label != primary_label
                            else ""
                        )
                        group_key = (
                            str(
                                finding.get("taxonomy_family") or primary_label
                            ).strip(),
                            str(
                                finding.get("subnet_key")
                                or secondary_label
                                or primary_label
                            ).strip(),
                        )
                        group = grouped.setdefault(
                            group_key,
                            {
                                "taxonomy_label": primary_label,
                                "subnet_label": secondary_label,
                                "fallback_heading": primary_label,
                                "findings": [],
                            },
                        )
                    else:
                        raw_heading = (
                            str(finding.get("category") or "").strip()
                            or "General Findings"
                        )
                        group = grouped.setdefault(
                            ("legacy", raw_heading),
                            {
                                "taxonomy_label": raw_heading,
                                "subnet_label": "",
                                "fallback_heading": raw_heading,
                                "findings": [],
                            },
                        )
                    group["findings"].append(finding)

                return list(grouped.values())

            def _build_render_groups(source_findings: list[dict]) -> list[dict]:
                outline_groups = data.get("outline_groups")
                if not isinstance(outline_groups, list) or not outline_groups:
                    return _build_fallback_render_groups(source_findings)

                indexed_findings = list(enumerate(source_findings))
                findings_by_id = {
                    str(finding.get("id") or "").strip(): (finding_index, finding)
                    for finding_index, finding in indexed_findings
                    if _has_meaningful_text(finding.get("id"))
                }
                used_indexes = set()
                render_groups = []

                for group in outline_groups:
                    matched = []
                    for finding_id in group.get("finding_ids") or []:
                        finding_key = str(finding_id or "").strip()
                        pair = findings_by_id.get(finding_key)
                        if pair is None:
                            continue
                        finding_index, finding = pair
                        if finding_index in used_indexes:
                            continue
                        matched.append((finding_index, finding))
                        used_indexes.add(finding_index)

                    if not matched:
                        taxonomy_family = str(
                            group.get("taxonomy_family") or ""
                        ).strip()
                        subnet_key = str(group.get("subnet_key") or "").strip()
                        for finding_index, finding in indexed_findings:
                            if finding_index in used_indexes:
                                continue
                            if (
                                taxonomy_family
                                and str(finding.get("taxonomy_family") or "").strip()
                                != taxonomy_family
                            ):
                                continue
                            if (
                                subnet_key
                                and str(finding.get("subnet_key") or "").strip()
                                != subnet_key
                            ):
                                continue
                            matched.append((finding_index, finding))
                            used_indexes.add(finding_index)

                    if not matched:
                        continue

                    group_findings = sorted(
                        [finding for _, finding in matched],
                        key=_finding_outline_sort_key,
                    )
                    first_finding = group_findings[0]
                    taxonomy_label = str(group.get("taxonomy_label") or "").strip()
                    subnet_label = str(group.get("subnet_label") or "").strip()
                    fallback_heading = (
                        str(first_finding.get("category") or "").strip()
                        or "General Findings"
                    )
                    render_groups.append(
                        {
                            "taxonomy_label": (
                                taxonomy_label
                                if _has_meaningful_text(taxonomy_label)
                                else str(
                                    first_finding.get("taxonomy_label") or ""
                                ).strip()
                            ),
                            "subnet_label": (
                                subnet_label
                                if _has_meaningful_text(subnet_label)
                                else str(
                                    first_finding.get("subnet_label") or ""
                                ).strip()
                            ),
                            "fallback_heading": fallback_heading,
                            "findings": group_findings,
                        }
                    )

                remaining_findings = [
                    finding
                    for finding_index, finding in indexed_findings
                    if finding_index not in used_indexes
                ]
                if remaining_findings:
                    render_groups.extend(
                        _build_fallback_render_groups(remaining_findings)
                    )

                return render_groups

            render_groups = _build_render_groups(findings)

            anchor = tpl_tbl_xml
            if anchor is None:
                if detail_idx >= 0 and detail_idx + 1 < len(children):
                    anchor = children[detail_idx + 1]
                elif children:
                    anchor = children[-1]
            if anchor is None:
                raise ValueError(
                    "Unable to resolve Detailed Observation insertion anchor"
                )

            def _add_heading_before(text: str, anchor_el, style_value: str):
                heading_text = str(text or "").strip()
                if not heading_text:
                    return
                p = OxmlElement("w:p")
                pPr = OxmlElement("w:pPr")
                ps = OxmlElement("w:pStyle")
                ps.set(qn("w:val"), style_value)
                pPr.append(ps)
                p.append(pPr)
                r = OxmlElement("w:r")
                t = OxmlElement("w:t")
                t.text = heading_text
                r.append(t)
                p.append(r)
                anchor_el.addprevious(p)

            def _add_page_break_before(anchor_el):
                pb_p = OxmlElement("w:p")
                pb_r = OxmlElement("w:r")
                pb_br = OxmlElement("w:br")
                pb_br.set(qn("w:type"), "page")
                pb_r.append(pb_br)
                pb_p.append(pb_r)
                anchor_el.addprevious(pb_p)

            global_seq = 1
            first_finding_overall = True
            current_taxonomy_heading = None
            for render_group in render_groups:
                group_findings = render_group.get("findings") or []
                if not group_findings:
                    continue

                if not first_finding_overall:
                    _add_page_break_before(anchor)

                taxonomy_heading = str(render_group.get("taxonomy_label") or "").strip()
                subnet_heading = str(render_group.get("subnet_label") or "").strip()
                fallback_heading = (
                    str(render_group.get("fallback_heading") or "").strip()
                    or "General Findings"
                )

                if _has_meaningful_text(taxonomy_heading):
                    if taxonomy_heading != current_taxonomy_heading:
                        _add_heading_before(taxonomy_heading, anchor, "Heading2")
                        current_taxonomy_heading = taxonomy_heading
                    if (
                        _has_meaningful_text(subnet_heading)
                        and subnet_heading != taxonomy_heading
                    ):
                        _add_heading_before(subnet_heading, anchor, "Heading3")
                else:
                    _add_heading_before(fallback_heading, anchor, "Heading2")
                    current_taxonomy_heading = fallback_heading

                for finding_index, finding in enumerate(group_findings):
                    if finding_index > 0:
                        _add_page_break_before(anchor)

                    new_tbl = _build_finding_table(finding, global_seq)
                    anchor.addprevious(new_tbl)

                    anchor.addprevious(OxmlElement("w:p"))

                    poc_val = _prepare_proof_of_concept(finding)
                    if poc_val:
                        poc_p = OxmlElement("w:p")
                        poc_r = OxmlElement("w:r")
                        poc_rPr = OxmlElement("w:rPr")
                        poc_rPr.append(OxmlElement("w:b"))
                        poc_r.append(poc_rPr)
                        poc_t = OxmlElement("w:t")
                        poc_t.text = "Proof of Concept:"
                        poc_r.append(poc_t)
                        poc_p.append(poc_r)
                        anchor.addprevious(poc_p)

                        poc_c_p = OxmlElement("w:p")
                        poc_c_r = OxmlElement("w:r")
                        poc_c_t = OxmlElement("w:t")
                        poc_c_t.text = poc_val
                        if poc_val != poc_val.strip():
                            poc_c_t.set(
                                "{http://www.w3.org/XML/1998/namespace}space", "preserve"
                            )
                        poc_c_r.append(poc_c_t)
                        poc_c_p.append(poc_c_r)
                        anchor.addprevious(poc_c_p)

                        anchor.addprevious(OxmlElement("w:p"))

                    first_finding_overall = False
                    global_seq += 1

            # Remove original placeholder finding table
            if tpl_tbl_xml is not None:
                parent_el = tpl_tbl_xml.getparent()
                if parent_el is not None:
                    parent_el.remove(tpl_tbl_xml)

    except Exception as e:
        import traceback

        print(f"⚠️  Detailed Observation section: {e}", file=sys.stderr)
        print(traceback.format_exc(), file=sys.stderr)

    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    _update_chart_severity_counts(doc, findings)
    doc.save(output_path)
    return output_path


def generate(
    scan_filepath: str, config_path: str = "config.json", cancel_event=None
) -> tuple:
    """
    Run pipeline: config → read scan → LLM → parse → validate → split FPs.
    Returns (data_dict, raw_llm_text, false_positive_list).
    """
    config = load_config(config_path)
    ensure_run_log(config, pipeline="single_shot")
    scan_text = read_scan_input(scan_filepath, config)
    raw_llm_text = call_llm(scan_text, config, cancel_event=cancel_event)
    data = safe_parse_json(raw_llm_text)
    data = validate_json_schema(data)
    data, fps = split_false_positives(data)
    data = _apply_local_normalization_pass(data, config)
    data = _apply_reference_validation_pass(data, config)
    data = _apply_quality_gate(data, config)
    data = compose_report_narrative(data)
    return data, raw_llm_text, fps


# ── Structured Extraction (CSV/Excel) & Web Search ────────────────────────────

STRUCTURED_LOOKUP_PROMPT = f"""You are a senior penetration tester writing findings for a professional VAPT report. Your output appears verbatim in a client-facing document — match the quality of a $30,000 manual pentest engagement.

YOUR TASK: Complete all fields for this finding and return a single JSON object.

ABSOLUTE RULES:
1. STRICT PRESERVATION: Do NOT alter 'name', 'severity', 'cvss', 'cve', 'affected_assets', 'observation', 'remediation_status', or 'risk_status'.
2. SPECIFIC OVER GENERIC: Every field must be specific to THIS vulnerability. Generic boilerplate is unacceptable.
2.1 TREAT GENERIC DEFAULTS AS INCOMPLETE: Replace ALL generic defaults with vulnerability-specific content.
3. BANNED GENERIC VALUES: Never output these values — they are placeholders, not real content:
   - control_objective: "{_GENERIC_CONTROL_OBJECTIVE_FALLBACK}"
   - control_objective: "Ensure that either the remote server is running the most updated version..."  (use this ONLY for actual outdated/EOL software, NOT for network config, SNMP, cipher issues, etc.)
   - control_name: "{_GENERIC_CONTROL_NAME_FALLBACK}"
   - audit_requirement: "{_GENERIC_AUDIT_REQUIREMENT_FALLBACK}"
   - audit_requirement: "Verify that the recommended control is implemented and periodically reviewed."
   If you cannot determine a specific value, return "[INSUFFICIENT DATA]" — do NOT invent generic filler.
4. OUTPUT FORMAT: RAW JSON only. No markdown, no code fences, no commentary.

FIELD REQUIREMENTS:

BUSINESS IMPACT (1-2 sentences):
  Describe the concrete attacker outcome and business consequence in this environment.
  Good: "Successful exploitation would allow an attacker to execute arbitrary code on the Jenkins controller, tamper with CI/CD jobs, access stored credentials, and pivot into connected internal systems."
  Bad: "This can impact the business."

CONTROL OBJECTIVE (1-2 sentences):
  The specific security goal for THIS vulnerability type. Tied to the threat.
  Good examples:
  - SSL/TLS ciphers: "Ensure that only strong cipher suites (AES-GCM, CHACHA20) are enabled in SSL/TLS configurations to protect data confidentiality and prevent SWEET32 birthday attacks and downgrade attacks."
  - SNMP default community: "Ensure SNMP community strings are changed from default values to non-guessable strings, access is restricted to authorised management hosts via ACL, and SNMPv3 with authentication and encryption is deployed."
  - RCE vulnerability: "Prevent attackers from executing arbitrary code on the server by applying the vendor security patch within the defined SLA, restricting the vulnerable service to trusted network segments, and monitoring for exploitation attempts."
  - SQL injection: "Prevent unauthorised database access and data manipulation by ensuring all database queries use parameterised statements and server-side input validation, eliminating direct inclusion of user-controlled data in SQL commands."
  Bad (banned): "Identify and remediate the vulnerability to reduce the attack surface."

CONTROL NAME (2-5 words, noun phrase, reusable):
  Good: "SSL/TLS Cipher Suite Configuration", "SNMP Community String Hardening", "Input Validation and Parameterised Queries", "Patch Management and Service Exposure Control", "SMB Service Configuration"
  Bad: "Vulnerability Remediation", "Security Control"

AUDIT REQUIREMENT (1-2 sentences — HOW to TEST, not HOW to fix):
  Good: "Verify SSL/TLS configuration using a TLS scanner (e.g., testssl.sh). Confirm only TLS 1.2 and 1.3 are accepted and 64-bit block ciphers (3DES) are rejected."
  Good: "Verify SNMP community strings have been changed from default values ('public', 'private', 'cisco'). Confirm ACL rules restrict SNMP access to authorised management hosts only."
  Bad: "Verify that the recommended control is implemented."

REMEDIATION (exactly 3 lines, newline-separated — no JSON arrays, no bullet characters):
  Line 1 — Primary fix: the definitive remediation. Version-specific if CVE data exists. 2-3 sentences.
  Line 2 — Alternative: compensating control when primary is not immediately feasible. DIFFERENT from primary. Think: vendor extended support, WAF rule, network isolation, ACL restriction, feature disable. Start with context: "If [primary] is not feasible due to [reason]:"
  Line 3 — Defensive: detection and monitoring ONLY — not a fix. SIEM/IDS/log review. Start with "Isolate..." or "Monitor..." or "Enable..."

  Example (unsupported OS):
  Upgrade to a supported version either 22.04 or 24.04 and apply all available security patches through the vendor update channel.
  Purchase Extended Security Maintenance (ESM) from Canonical via the Ubuntu Pro Subscription if an immediate upgrade is not operationally feasible, to continue receiving security patches on the legacy release.
  Isolate and harden the system: set up a firewall to restrict access to trusted hosts only, disable all non-essential services, and actively monitor system activity for anomalous behaviour.

  Example (SSL medium cipher):
  Reconfigure the affected service to disable medium-strength and weak cipher suites (3DES, RC4, EXPORT) and accept only strong AEAD ciphers (AES-128-GCM, AES-256-GCM, CHACHA20-POLY1305) under TLS 1.2 and TLS 1.3.
  If configurations cannot be updated immediately due to legacy client compatibility dependencies, implement a WAF rule or network-layer filter to block connections negotiating weak cipher suites, and restrict access to the affected service to trusted internal hosts only via firewall ACL.
  Enable logging and monitoring for unusual SSL/TLS traffic patterns. Implement network-level encryption solutions such as VPN tunnels as a temporary mitigation. Educate stakeholders about risks related to medium-strength cipher suites and planned remediation timelines.

REFERENCE: Return 2-4 authoritative references only. Prefer NVD, MITRE CWE, OWASP, and vendor advisories. Do not return placeholders or generic labels.

CLIENT CONTEXT: {{client_context}}

RAW JSON ONLY. NO MARKDOWN. NO COMMENTARY.
/no_think"""


def _build_metadata_from_findings(findings: list, client_context: str = "") -> dict:
    """Build report metadata programmatically from findings — no LLM call."""
    counts = {"Critical": 0, "High": 0, "Medium": 0, "Low": 0}
    for f in findings:
        sev = f.get("severity", "")
        if sev in counts:
            counts[sev] += 1
    total = sum(counts.values())

    # Extract client name from context if provided
    client_name = "[Client Name]"
    if client_context:
        # Try to extract company/client name from context
        for line in client_context.splitlines():
            line = line.strip()
            if line and len(line) < 100 and not line.startswith("#"):
                client_name = line.split(":")[1].strip() if ":" in line else line
                break

    # Collect unique subnets from affected_assets
    subnets = set()
    for f in findings:
        assets = f.get("affected_assets", "")
        for m in re.finditer(r"(\d{1,3}\.\d{1,3}\.\d{1,3})\.", str(assets)):
            subnets.add(m.group(1) + ".0/24")

    scope_nets = ", ".join(sorted(subnets)) if subnets else "the target network"

    exec_summary = (
        f"A vulnerability assessment and penetration test was conducted on the internal network infrastructure. "
        f"The assessment identified a total of {total} unique vulnerabilities: "
        f"{counts['Critical']} Critical, {counts['High']} High, {counts['Medium']} Medium, and {counts['Low']} Low. "
        f"Several critical and high-severity findings require immediate remediation to mitigate risk of unauthorized access and data compromise."
    )

    scope_summary = (
        f"As a part of the internal VAPT engagement, the assessment covered the network subnets {scope_nets}. "
        f"The testing approach included automated vulnerability scanning and manual verification of findings. "
        f"The objective was to identify exploitable vulnerabilities in the internal infrastructure and provide actionable remediation guidance."
    )

    return {
        "client_name": client_name,
        "report_date": datetime.now().strftime("%d-%b-%Y"),
        "engagement_type": "Internal VAPT",
        "assessment_phase": "Phase 1",
        "assessor_firm": "[Assessor Firm]",
        "scope_summary": scope_summary,
        "executive_summary": exec_summary,
        "total_critical": counts["Critical"],
        "total_high": counts["High"],
        "total_medium": counts["Medium"],
        "total_low": counts["Low"],
        "total_findings": total,
        "conclusion": "[PRESERVE_ORIGINAL]",
        "methodology": "OWASP / PTES / NIST",
    }


def _generate_from_structured_file(
    filepath, config: dict, client_context: str, cancel_event, progress_callback
):
    """
    Groups vulnerability rows from Nessus CSV/Excel, enriches missing detail,
    and runs a single detail-lookup LLM pass over the combined structured representation.

    filepath can be a single path string or a list of paths for multi-file consolidation.
    """
    assert_clean_client_context_for_cloud(client_context, config)
    ensure_run_log(config, pipeline="structured_file")

    def _progress(stage, current, total, message, detail=None):
        _emit_progress_callback(
            progress_callback, stage, current, total, message, detail
        )

    _progress("read", 0, 1, "Parsing structured file(s)…")

    # Support single file or list of files
    if isinstance(filepath, str):
        filepaths = [filepath]
    else:
        filepaths = list(filepath)

    frames = []
    for fp in filepaths:
        ext = Path(fp).suffix.lower()
        if ext == ".csv":
            part = pd.read_csv(fp)
        else:
            part = pd.read_excel(fp)
        frames.append(part)
    df = pd.concat(frames, ignore_index=True) if len(frames) > 1 else frames[0]

    # Standardize column names (lowercase and strip)
    orig_cols = list(df.columns)
    col_map = {c: str(c).strip().lower() for c in orig_cols}
    df = df.rename(columns=col_map)

    # Identify key columns from common Nessus variations
    _get_col = lambda possible_names: next(
        (c for c in df.columns if c in possible_names), None
    )

    c_name = _get_col(["name", "plugin name", "vulnerability", "title"])
    c_sev = _get_col(["risk", "severity", "risk factor"])
    c_plugin = _get_col(["plugin id", "plugin_id", "pluginid"])
    c_cve = _get_col(["cve", "cves"])
    c_cvss = _get_col(
        [
            "cvss",
            "cvss v3.0 base score",
            "cvss v2.0 base score",
            "cvss v2.0",
            "cvss base score",
            "cvss score",
        ]
    )
    c_host = _get_col(["host", "ip address", "ip"])
    c_port = _get_col(["port"])
    c_desc = _get_col(["description", "synopsis"])
    c_sol = _get_col(["solution", "remediation"])
    c_poc = _get_col(["plugin output", "proof of concept", "exploit"])
    c_ref = _get_col(["see also", "references", "ref", "links", "external references"])
    c_observation = _get_col(["observation", "finding observation", "repeat status"])
    c_rem_status = _get_col(["remediation status", "finding status", "closure status"])
    c_risk_status = _get_col(
        ["risk status", "disposition", "review status", "false positive check"]
    )

    if not c_name:
        raise ValueError(
            "Could not find a 'Name' or 'Plugin Name' column in the structured file."
        )

    findings_map = {}

    _progress("initial_extract", 0, 1, "Grouping structured findings…")

    rows = df.to_dict("records")
    for row in rows:
        raw_name = _strip_leading_plugin_id(_structured_cell_text(row.get(c_name, "")))
        if not raw_name:
            continue

        group_name, display_name = _structured_group_and_title_for_row(row, raw_name)

        host = _structured_cell_text(row.get(c_host, "")) if c_host else ""
        port = _structured_cell_text(row.get(c_port, "")) if c_port else ""

        asset = host
        if host and port and port != "0":
            asset = f"{host}:{port}"

        sev_val = _structured_normalize_severity(row.get(c_sev, "Informational"))
        if sev_val == "Informational" and not _structured_keep_informational(
            row, raw_name
        ):
            continue
        if sev_val == "Informational":
            sev_val = "Low"

        observation = (
            _structured_cell_text(row.get(c_observation, "")) if c_observation else ""
        )
        remediation_status = (
            _structured_cell_text(row.get(c_rem_status, "")) if c_rem_status else ""
        )
        risk_status = (
            _structured_cell_text(row.get(c_risk_status, "")) if c_risk_status else ""
        )

        if not observation:
            observation = "New"
        if not remediation_status:
            remediation_status = "Open"
        if not risk_status:
            risk_status = "Open"

        status_family = _structured_status_family(
            observation, remediation_status, risk_status
        )
        group_key = (group_name, status_family)

        candidate = {
            "name": display_name,
            "severity": sev_val,
            "cve": _structured_cell_text(row.get(c_cve, "[INSUFFICIENT DATA]"))
            or "[INSUFFICIENT DATA]",
            "cvss": _structured_cell_text(row.get(c_cvss, "[INSUFFICIENT DATA]"))
            or "[INSUFFICIENT DATA]",
            "description": _structured_cell_text(row.get(c_desc, "")),
            "remediation": _structured_cell_text(row.get(c_sol, "")),
            "affected_assets": set(),
            "observation": observation,
            "remediation_status": remediation_status,
            "risk_status": risk_status,
            "business_impact": "[INSUFFICIENT DATA]",
            "proof_of_concept": _structured_cell_text(
                row.get(c_poc, "[INSUFFICIENT DATA]")
            )
            or "[INSUFFICIENT DATA]",
            "control_name": "[INSUFFICIENT DATA]",
            "control_objective": "[INSUFFICIENT DATA]",
            "audit_requirement": "[INSUFFICIENT DATA]",
            "reference": (
                _structured_cell_text(row.get(c_ref, "[INSUFFICIENT DATA]"))
                if c_ref
                else "[INSUFFICIENT DATA]"
            ),
            "vuln_id": _structured_cell_text(row.get(c_plugin, "")) if c_plugin else "",
            "_raw_name": raw_name,
        }

        if group_key not in findings_map:
            findings_map[group_key] = candidate
        else:
            existing = findings_map[group_key]
            preferred_title = existing["name"]
            if _structured_find_score(
                candidate, preferred_title
            ) > _structured_find_score(existing, preferred_title):
                candidate["affected_assets"] = existing["affected_assets"]
                findings_map[group_key] = candidate
                existing = findings_map[group_key]

            if not (
                _structured_is_preferred_title(existing, preferred_title)
                and not _structured_is_preferred_title(candidate, preferred_title)
            ):
                existing["description"] = _structured_merge_text(
                    existing.get("description"), candidate.get("description")
                )
                existing["remediation"] = _structured_merge_text(
                    existing.get("remediation"), candidate.get("remediation")
                )
                existing["proof_of_concept"] = _structured_merge_text(
                    existing.get("proof_of_concept"), candidate.get("proof_of_concept")
                )
                existing["cve"] = _structured_merge_text(
                    existing.get("cve"), candidate.get("cve")
                )
                existing["cvss"] = _structured_merge_text(
                    existing.get("cvss"), candidate.get("cvss")
                )
                existing["reference"] = _structured_merge_text(
                    existing.get("reference"), candidate.get("reference")
                )

        if asset:
            findings_map[group_key]["affected_assets"].add(asset)

    findings_list = []
    for idx, (_, f_data) in enumerate(findings_map.items()):
        f_data["affected_assets"] = (
            ", ".join(sorted(list(f_data["affected_assets"])))
            if f_data["affected_assets"]
            else "[INSUFFICIENT DATA]"
        )
        f_data["id"] = f"VAPT-{(idx+1):03d}"
        if f_data["name"] in _STRUCTURED_SEVERITY_OVERRIDES:
            f_data["severity"] = _STRUCTURED_SEVERITY_OVERRIDES[f_data["name"]]
        f_data.pop("_raw_name", None)
        findings_list.append(f_data)

    total_findings = len(findings_list)

    lookup_findings = []
    raw_responses = []
    lookup_context = client_context
    system_prompt = STRUCTURED_LOOKUP_PROMPT.format(client_context=lookup_context)
    structured_config = config.copy()
    structured_config["llm"] = {**config["llm"], "max_tokens": 2048}
    cloud_lookup_stats = {}
    second_stage_skipped = False

    append_run_log_event(
        config,
        "nessus_findings_parsed",
        {
            "source_files": [str(fp) for fp in filepaths],
            "input_row_count": len(rows),
            "finding_count": total_findings,
            "findings": copy.deepcopy(findings_list),
        },
    )

    _preflight_structured_cloud_prompts(
        structured_config,
        system_prompt,
        lookup_context,
    )

    # ── Phase 0: cloud_enrich primary pass ───────────────────────────────────
    # Run NVD/EPSS/CWE/CAPEC/framework lookup + OpenRouter LLM lookup before
    # local LLM lookup. Findings already enriched by cloud_enrich skip the
    # per-finding local LLM call entirely. Cloud failures abort the run for
    # cloud providers so the UI surfaces a real error instead of degrading.
    _cloud_cfg = _resolve_structured_cloud_lookup_cfg(config)
    if not _cloud_cfg:
        _progress(
            "lookup",
            0,
            total_findings,
            f"Cloud lookup skipped — local pipeline will process {total_findings} findings.",
        )
    else:
        from report_tool.lookup.cloud_enrich import (
            lookup_report as _cloud_enrich_report,
        )
        _cloud_cfg = {
            **_cloud_cfg,
            "_run_log_path": config.get("_run_log_path"),
            "paths": config.get("paths", {}),
        }

        _progress(
            "lookup",
            0,
            total_findings,
            f"Cloud lookup ({_cloud_cfg.get('model')}) starting on {total_findings} findings…",
        )

        def _cloud_on_progress(current, total, *args):
            msg = args[0] if args else None
            detail = args[1] if len(args) > 1 and isinstance(args[1], dict) else None
            try:
                _progress(
                    "lookup",
                    int(current) if current is not None else 0,
                    int(total) if total is not None else max(1, total_findings),
                    str(msg) if msg else f"Cloud lookup {current}/{total}",
                    detail=detail,
                )
            except Exception:
                pass

        try:
            _cloud_data = _cloud_enrich_report(
                {"findings": findings_list}, _cloud_cfg, on_progress=_cloud_on_progress
            )
        except Exception as _ce:
            _raise_if_client_data_egress_blocked(_ce)
            if is_cloud_provider(config):
                raise RuntimeError(
                    f"Cloud lookup failed for {total_findings} findings: {_ce}"
                ) from _ce
            _progress(
                "lookup",
                0,
                total_findings,
                f"Cloud lookup failed (continuing with local LLM): {_ce}",
            )
        else:
            findings_list = _cloud_data.get("findings", findings_list)
            cloud_lookup_stats = _cloud_data.get("_lookup_stats", {})

    findings_missing_data = []
    derived_findings_snapshot = []
    for f in findings_list:
        fill_missing_fields(f)

        business_impact = str(f.get("business_impact", "") or "").strip()
        if not business_impact or business_impact in ("[INSUFFICIENT DATA]", "nan"):
            f["business_impact"] = _derive_business_impact_from_finding(f)

        desc = f.get("description", "").strip()
        rem = f.get("remediation", "").strip()

        # ── Derive description from plugin_output if missing ─────────────────
        # Plugin output often contains useful technical details
        if not desc or desc in ("[INSUFFICIENT DATA]", "nan"):
            poc = f.get("proof_of_concept", "").strip()
            if poc and poc not in ("[INSUFFICIENT DATA]", "nan") and len(poc) > 50:
                # Extract first 2-3 meaningful sentences from plugin output
                sentences = re.split(r"(?<=[.!?])\s+", poc)[:3]
                derived_desc = " ".join(s for s in sentences if len(s) > 20)
                if derived_desc:
                    f["description"] = derived_desc
                    desc = derived_desc

        needs_llm = _structured_finding_needs_llm_lookup(f)
        snapshot = copy.deepcopy(f)
        snapshot["needs_llm_lookup"] = bool(needs_llm)
        derived_findings_snapshot.append(snapshot)

        if needs_llm:
            findings_missing_data.append(f)
        else:
            lookup_findings.append(f)

    append_run_log_event(
        config,
        "structured_findings_derived",
        {
            "source_files": [str(fp) for fp in filepaths],
            "input_row_count": len(rows),
            "original_columns": orig_cols,
            "normalized_columns": list(df.columns),
            "selected_columns": {
                "name": c_name,
                "severity": c_sev,
                "plugin_id": c_plugin,
                "cve": c_cve,
                "cvss": c_cvss,
                "host": c_host,
                "port": c_port,
                "description": c_desc,
                "solution": c_sol,
                "proof_of_concept": c_poc,
                "reference": c_ref,
                "observation": c_observation,
                "remediation_status": c_rem_status,
                "risk_status": c_risk_status,
            },
            "finding_count": total_findings,
            "findings_requiring_llm": len(findings_missing_data),
            "findings_ready_without_llm": len(lookup_findings),
            "findings": derived_findings_snapshot,
        },
    )

    if findings_missing_data:
        run_second_stage = _should_run_structured_second_stage(
            structured_config,
            cloud_lookup_stats,
        )
        cloud_eligible = int(cloud_lookup_stats.get("cloud_eligible_findings") or 0)
        if not run_second_stage:
            second_stage_skipped = True
            lookup_findings.extend(findings_missing_data)
            append_run_log_event(
                config,
                "structured_second_stage_skipped",
                {
                    "missing_after_cloud": len(findings_missing_data),
                    "total_findings": total_findings,
                    "cloud_eligible_findings": cloud_eligible,
                    "reason": "cloud_lookup_already_ran",
                },
            )
            _progress(
                "enrich",
                len(findings_missing_data),
                len(findings_missing_data),
                (
                    f"Second-stage LLM lookup skipped for {len(findings_missing_data)} findings"
                    + (
                        f" after cloud enrichment ({cloud_eligible} were cloud-eligible)."
                        if cloud_eligible
                        else "."
                    )
                ),
            )
        else:
            _progress(
                "enrich",
                0,
                len(findings_missing_data),
                (
                    f"LLM lookup: {len(findings_missing_data)} of {total_findings} total findings still missing data"
                    + (
                        f" after cloud enrichment ({cloud_eligible} were cloud-eligible)."
                        if cloud_eligible
                        else "."
                    )
                ),
            )

            for i, finding in enumerate(findings_missing_data):
                if cancel_event and cancel_event.is_set():
                    break

                send_finding = _build_structured_lookup_prompt_finding(finding)
                restore_map = {}
                if is_cloud_provider(config):
                    send_finding, restore_map = sanitize_finding(send_finding)

                payload = json.dumps(send_finding, indent=2, ensure_ascii=False)
                user_prompt = _build_structured_lookup_user_prompt(
                    lookup_context,
                    payload,
                )

                try:
                    raw_response = _call_llm_generic(
                        system_prompt,
                        user_prompt,
                        structured_config,
                        cancel_event,
                        log_label=f"structured_lookup_{i+1}",
                        task_type="lookup",
                    )
                    try:
                        parsed = safe_parse_json(raw_response)
                    except ValueError:
                        print(
                            f"⚠️ Structured lookup JSON parse failed for {finding['name']}; retrying once.",
                            file=sys.stderr,
                        )
                        retry_config = copy.deepcopy(structured_config)
                        cur_max = retry_config.get("llm", {}).get("max_tokens", 2048)
                        retry_config.setdefault("llm", {})["max_tokens"] = max(
                            cur_max + 512, int(cur_max * 1.25)
                        )
                        raw_response = _call_llm_generic(
                            system_prompt,
                            user_prompt
                            + "\n\nIMPORTANT: Your previous reply was not valid JSON. Return exactly one valid JSON object. No markdown, no code fences.",
                            retry_config,
                            cancel_event,
                            log_label=f"structured_lookup_{i+1}_retry",
                            task_type="lookup",
                        )
                        parsed = safe_parse_json(raw_response)
                    if isinstance(parsed, dict):
                        if restore_map:
                            parsed = restore_finding(parsed, restore_map)
                        parsed = _merge_structured_lookup_result(finding, parsed)
                        lookup_findings.append(parsed)
                    else:
                        lookup_findings.append(finding)
                    if raw_response:
                        raw_responses.append(raw_response)
                except Exception as e:
                    _raise_if_client_data_egress_blocked(e)
                    if is_cloud_provider(config):
                        raise RuntimeError(
                            f"Structured lookup failed for {finding['name']}: {e}"
                        ) from e
                    print(f"⚠️ Lookup failed for {finding['name']}: {e}", file=sys.stderr)
                    lookup_findings.append(finding)

                _progress(
                    "enrich",
                    i + 1,
                    len(findings_missing_data),
                    f"Enriched {i + 1}/{len(findings_missing_data)} findings (of {total_findings} total)…",
                )
    else:
        cloud_eligible = int(cloud_lookup_stats.get("cloud_eligible_findings") or 0)
        enrich_message = (
            (
                f"All {total_findings} findings were enriched after cloud lookup — "
                "no additional second-stage LLM lookup required."
            )
            if cloud_eligible
            else (
                f"All {total_findings} findings were resolved locally — "
                "no LLM lookup required."
            )
        )
        _progress(
            "enrich",
            total_findings,
            total_findings,
            enrich_message,
        )

    try:
        _validate_structured_source_fields(lookup_findings)
    except ValueError as exc:
        if not second_stage_skipped:
            raise
        append_run_log_event(
            config,
            "structured_incomplete_fields_warning",
            {
                "warning": str(exc),
                "second_stage_skipped": True,
            },
        )
        print(f"⚠️  {exc}", file=sys.stderr)

    # Metadata — build programmatically from findings (no LLM call needed)
    _progress("metadata", 0, 1, "Building metadata from findings…")
    metadata = _build_metadata_from_findings(lookup_findings, client_context)
    _progress("metadata", 1, 1, "Metadata built")

    # Cleanup and re-number final IDs
    for idx, f in enumerate(lookup_findings, start=1):
        f["id"] = f"VAPT-{idx:03d}"
        if "_web_search_context" in f:
            del f["_web_search_context"]

    # Deduplicate and validate
    data = metadata if isinstance(metadata, dict) else {}
    data["findings"] = dedupe_findings(lookup_findings)
    data = validate_json_schema(data)
    data, fps = split_false_positives(data)
    data = _apply_local_normalization_pass(data, config)
    data = _apply_reference_validation_pass(data, config)
    data = _apply_quality_gate(data, config)
    data = compose_report_narrative(data)

    _progress(
        "done", 1, 1, f"Lookup complete — {data.get('total_findings',0)} findings."
    )
    return data, raw_responses, fps
